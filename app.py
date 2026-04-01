"""
수석 리서치 품질 검수관 - 보고서 분석기
Streamlit Cloud 배포용 메인 앱
"""
import streamlit as st
import concurrent.futures
import datetime
import config
from config import (
    APP_TITLE, APP_ICON, APP_DESCRIPTION,
    AVAILABLE_MODELS, DEFAULT_MODEL, MAX_TEXT_CHARS,
    MODEL_DISPLAY_NAMES, AUTO_MODE_LABEL, AUTO_MODEL_PRIORITY,
)
from file_processor import extract_text, truncate_text
from analyzer import (
    run_analysis_stream, run_analysis, _run_single,
    STEP_PROMPTS, STEP_LABELS,
    FULL_ANALYSIS_PROMPT, get_api_key, get_api_keys,
    run_step_with_chunks, check_key_quotas
)
from rfp_prompts import RFP_SECTIONS
import rfp_utils
import io
import pandas as pd
import numpy as np
from data_cleaner import DataImputer, WeightCalculator, DataAugmentor
from codebook_utils import CodebookParser
import plotly.express as px

from api.constants import DATA_GO_KR_SERVICE_KEY as SERVICE_KEY, OPEN_DART_API_KEY as DART_API_KEY
from api_utils import (
    detect_and_load_mois_excel, get_mois_region_levels,
    parse_mois_excel_with_gender, format_sample_pivot_table,
    SIDO_LIST
)
from api.nhis_api import (
    download_nhis_dataset, NHIS_SELECTABLE_FIELDS, NHIS_FIELD_LABELS, NHIS_FIELD_MAP,
    NHIS_ENDPOINTS
)
from api.nps_api import search_and_match_nps, NPS_SELECTABLE_FIELDS, NPS_FIELD_LABELS, estimate_avg_salary
from api.fss_api import (
    search_corp_by_name, search_financial_by_crno, validate_crno,
    FSS_CORP_SELECTABLE_FIELDS, FSS_CORP_FIELD_LABELS,
    FSS_FINA_SELECTABLE_FIELDS, FSS_FINA_FIELD_LABELS,
)
from api.dart_api import get_dart_corp_info
from api.g2b_api import get_g2b_corp_info, G2B_SELECTABLE_FIELDS, G2B_FIELD_LABELS
from api.nts_api import get_nts_business_status
from api.biz_search_api import (
    get_ai_industry_suggestions, 
    batch_search_and_consolidate,
    SIDO_LIST as BIZ_SIDO_LIST
)
from utils.excel_handler import load_excel, export_result_excel
from utils.matcher import normalize_brn, clean_company_names_bulk, clean_addresses_bulk, clean_address, split_address
from utils.stats_utils import get_association
from usage_tracker import UsageTracker

# ── 이용 통계 트래커 초기화
tracker = UsageTracker()
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None

# ── 페이지 설정
st.set_page_config(
    page_title="보고서 검수 AI Tools",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Google Fonts (style 블록과 반드시 분리)
st.markdown(
    '<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">',
    unsafe_allow_html=True,
)

# ── Qualtrics SaaS 스타일 CSS
st.markdown("""
<style>
html, body, [class*="css"] {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
}
[data-testid="stAppViewContainer"] { background-color: #F4F6F9; }
[data-testid="block-container"] { padding-top: 0 !important; }

[data-testid="stSidebar"] { background-color: #1B2437; border-right: none; }
[data-testid="stSidebar"] * { color: #C8D0DC !important; }
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] strong { color: #FFFFFF !important; }
[data-testid="stSidebar"] .stButton > button {
    background: transparent !important;
    border: 1px solid #374151 !important;
    color: #9CA3AF !important;
    border-radius: 6px !important;
    font-size: 0.82rem !important;
    padding: 0.4rem 1rem !important;
    box-shadow: none !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
    border-color: #6B7280 !important;
    color: #D1D5DB !important;
    background: #1F2A3D !important;
}
[data-testid="stSidebar"] hr { border-color: #2D3A50 !important; }
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] input::placeholder,
[data-testid="stSidebar"] [data-baseweb="select"] *,
[data-testid="stSidebar"] [data-baseweb="input"] *,
[data-testid="stSidebar"] [class*="st-"] { color: #FFFFFF !important; }
[data-testid="stSidebar"] [data-baseweb="select"] > div,
[data-testid="stSidebar"] [data-baseweb="input"] > div { background-color: #243048 !important; border-color: #374151 !important; }

/* 네비게이션 라디오 버튼 텍스트 크기 및 줄바꿈 방지 */
[data-testid="stSidebar"] .stRadio div[data-testid="stMarkdownContainer"] p {
    font-size: 0.82rem !important;
    white-space: nowrap !important;
    overflow: hidden;
    text-overflow: ellipsis;
}


.qx-topbar {
    background: #FFFFFF;
    border-bottom: 1px solid #E5E9F0;
    padding: 0.9rem 2rem;
    display: flex;
    align-items: center;
    gap: 0.75rem;
    margin-bottom: 1.5rem;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
}
.qx-topbar-logo { font-size: 1.35rem; font-weight: 700; color: #0F6CBD; letter-spacing: -0.5px; }
.qx-topbar-sep  { width: 1px; height: 22px; background: #D1D9E6; margin: 0 0.5rem; }
.qx-topbar-title { font-size: 0.88rem; font-weight: 500; color: #4A5568; }
.qx-topbar-badge {
    margin-left: auto;
    background: #EEF4FD; color: #0F6CBD;
    border: 1px solid #BDD7F5; border-radius: 20px;
    padding: 2px 12px; font-size: 0.75rem; font-weight: 600;
}

.qx-section-label {
    font-size: 0.72rem; font-weight: 600;
    letter-spacing: 0.8px; text-transform: uppercase;
    color: #8B96A9; margin-bottom: 0.5rem;
}

.qx-card {
    background: #FFFFFF;
    border: 1px solid #E5E9F0;
    border-radius: 10px;
    padding: 1.5rem 1.75rem;
    box-shadow: 0 1px 4px rgba(15,107,189,0.05);
    margin-bottom: 1rem;
}
.qx-card-title { font-size: 0.95rem; font-weight: 600; color: #1A2237; margin-bottom: 0.3rem; }

.qx-upload-zone {
    background: #FAFBFE;
    border: 2px dashed #C5D5EE;
    border-radius: 10px;
    padding: 2.5rem 1.5rem;
    text-align: center;
    transition: border-color 0.25s, background 0.25s;
}
.qx-upload-zone:hover { border-color: #0F6CBD; background: #EEF4FD; }
.qx-upload-icon  { font-size: 2.2rem; margin-bottom: 0.5rem; }
.qx-upload-text  { font-size: 0.9rem; font-weight: 500; color: #3D4F6B; }
.qx-upload-hint  { font-size: 0.78rem; color: #A0AABB; margin-top: 0.3rem; }

[data-baseweb="tab-list"] {
    background: #FFFFFF !important;
    border-bottom: 2px solid #E5E9F0 !important;
    border-radius: 0 !important;
    padding: 0 !important; gap: 0 !important;
}
[data-baseweb="tab"] {
    font-size: 0.85rem !important; font-weight: 500 !important;
    color: #64748B !important; border-radius: 0 !important;
    padding: 0.7rem 1.25rem !important;
    border-bottom: 2px solid transparent !important;
    margin-bottom: -2px !important;
}
[data-baseweb="tab"]:hover { color: #0F6CBD !important; background: #F0F7FF !important; }
[aria-selected="true"] {
    color: #0F6CBD !important;
    border-bottom: 2px solid #0F6CBD !important;
    background: transparent !important;
    font-weight: 600 !important;
}

.badge-ok    { background:#EEF4FD; color:#0F6CBD; border:1px solid #BDD7F5; border-radius:20px; padding:3px 11px; font-size:0.76rem; font-weight:600; }
.badge-warn  { background:#FFF8E6; color:#B45309; border:1px solid #FCD34D; border-radius:20px; padding:3px 11px; font-size:0.76rem; font-weight:600; }
.badge-error { background:#FEF2F2; color:#DC2626; border:1px solid #FCA5A5; border-radius:20px; padding:3px 11px; font-size:0.76rem; font-weight:600; }

.stButton > button {
    background: #0F6CBD !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    padding: 0.55rem 1.25rem !important;
    box-shadow: 0 1px 3px rgba(15,107,189,0.2) !important;
    transition: background 0.18s !important;
}
.stButton > button:hover { background: #0A5AA0 !important; }

[data-testid="stFileUploader"] {
    background: #FAFBFE;
    border: 1.5px dashed #C5D5EE;
    border-radius: 8px; padding: 0.5rem;
}

[data-testid="stDownloadButton"] > button {
    background: #FFFFFF !important;
    color: #0F6CBD !important;
    border: 1.5px solid #0F6CBD !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    box-shadow: none !important;
}
[data-testid="stDownloadButton"] > button:hover { background: #EEF4FD !important; }

.sb-file-card {
    background: #1F2A3D;
    border: 1px solid #2D3A50;
    border-radius: 8px;
    padding: 0.85rem 1rem; margin-bottom: 0.75rem;
}
.sb-file-name { font-size: 0.82rem; font-weight: 600; color: #E2E8F0 !important; word-break: break-all; }
.sb-file-meta { font-size: 0.75rem; color: #8B96A9 !important; margin-top: 0.25rem; }

/* AI 모델 수동 선택 영역 강조 (회색톤으로 변경) */
.sb-highlight-container {
    background: #2D3A50 !important; /* 회색톤 배경 */
    border: 1px solid #4A5568 !important; /* 차분한 회색 테두리 */
    border-radius: 10px;
    padding: 1.2rem;
    margin-top: 0.8rem;
    box-shadow: 0 4px 8px rgba(0,0,0,0.3);
}
.sb-model-badge {
    background: #4A5568; /* 회색톤 배지 */
    color: #FFFFFF !important;
    padding: 2px 10px;
    border-radius: 4px;
    font-size: 0.7rem;
    font-weight: 700;
    text-transform: uppercase;
    margin-bottom: 0.8rem;
    display: inline-block;
    letter-spacing: 0.5px;
}

hr { border-color: #E5E9F0 !important; margin: 1rem 0 !important; }
[data-testid="stAlert"] { border-radius: 8px !important; font-size: 0.875rem !important; }
#MainMenu { visibility: hidden; }
footer    { visibility: hidden; }
header    { visibility: hidden; }

/* 표 내부 줄바꿈 허용 및 가독성 개선 */
table { width: 100% !important; border-collapse: collapse !important; }
th, td { 
    white-space: normal !important; 
    word-break: keep-all !important; 
    line-height: 1.6 !important; 
    padding: 10px 15px !important;
    vertical-align: top !important;
}
td br { content: ""; display: block; margin-bottom: 0.5rem; }

@media print {
    [data-testid="stSidebar"], .stButton, hr, .qx-topbar-badge, .badge-ok, .badge-warn { display: none !important; }
    [data-testid="block-container"] { padding: 0 !important; background-color: white !important; }
    .qx-card { border: none !important; box-shadow: none !important; padding: 0 !important; margin-bottom: 2rem !important; page-break-inside: avoid; }
    .qx-topbar { border-bottom: 2px solid #0F6CBD !important; }
    body { background-color: white !important; }
}
</style>
""", unsafe_allow_html=True)


# ── 세션 상태 초기화
def init_session():
    defaults = {
        "is_logged_in": False,
        "logged_in_user": "",
        "login_error": "",
        "report_text": "",
        "file_name": "",
        "file_pages": 0,
        "step_results": {1: "", 2: "", 3: ""},
        "full_result": "",
        "selected_model": DEFAULT_MODEL,
        "auto_mode": True,    # 자동 최적화 모드 기본값
        "menu_selection": "과업 내용 체크 리스트",
        "rfp_curr_text": "",
        "rfp_prev_text": "",
        "rfp_results": {},
        "rfp_project_name": "",
        "rfp_curr_name": "",
        "rfp_prev_name": "",
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val

init_session()


# ── 로그인 헬퍼
def get_allowed_users() -> list:
    try:
        users = st.secrets["ALLOWED_USERS"]
        if isinstance(users, str):
            return [u.strip() for u in users.split(",") if u.strip()]
        return [str(u).strip() for u in users if u]
    except Exception:
        return []


def do_login(username: str):
    allowed = get_allowed_users()
    uid = username.strip()
    if not uid:
        st.session_state["login_error"] = "아이디를 입력해 주세요."
    elif uid in allowed:
        st.session_state["is_logged_in"] = True
        st.session_state["logged_in_user"] = uid
        st.session_state["login_error"] = ""
        # [v4.12] 로그인 로그 기록
        tracker.log_event(uid, "LOGIN")
    else:
        st.session_state["login_error"] = f"'{uid}'은(는) 등록되지 않은 아이디입니다."


def do_logout():
    st.session_state["is_logged_in"] = False
    st.session_state["logged_in_user"] = ""
    st.session_state["login_error"] = ""


def show_data_roadmap(current_step=1):
    """RFP 분석 프로세스 도식 (시각적 도식)"""
    steps = [
        {"id": 1, "icon": "📄", "label": "RFP 업로드"},
        {"id": 2, "icon": "🔎", "label": "과업 요약 분석"},
        {"id": 3, "icon": "⚖️", "label": "참가자격 검토"},
        {"id": 4, "icon": "🎯", "label": "핵심요구 추출"},
        {"id": 5, "icon": "📋", "label": "전략 리포트"}
    ]
    
    html = '<div style="display:flex; align-items:center; justify-content:space-between; margin: 1.5rem 0; padding:1rem; background:white; border-radius:12px; border:1px solid #E5E9F0;">'
    for i, s in enumerate(steps):
        is_active = s['id'] == current_step
        color = "#0F6CBD" if is_active else "#64748B"
        bg = "#EEF4FD" if is_active else "#F8FAFC"
        border = "2px solid #0F6CBD" if is_active else "1px solid #E2E8F0"
        
        html += f'''
        <div style="flex:1; display:flex; flex-direction:column; align-items:center; position:relative;">
            <div style="width:40px; height:40px; border-radius:50%; background:{bg}; border:{border}; display:flex; align-items:center; justify-content:center; font-size:1.2rem; margin-bottom:0.5rem; color:{color}; z-index:2;">
                {s['icon']}
            </div>
            <div style="font-size:0.7rem; font-weight:{'700' if is_active else '500'}; color:{color}; text-align:center;">{s['label']}</div>
        '''
        if i < len(steps) - 1:
            html += f'<div style="position:absolute; top:20px; left:calc(50% + 20px); width:calc(100% - 40px); height:2px; background:#E2E8F0; z-index:1;"></div>'
        html += '</div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


def show_security_notice():
    """데이터 보안 및 로컬 저장 안내 (v4.13)"""
    st.warning("""
    🛡️ **데이터 보안 안내:** 본 시스템은 사용자의 소중한 데이터 보안을 최우선으로 고려하며, 일체의 분석 결과를 서버에 저장하지 않습니다. 
    귀중한 분석 자료의 유실을 방지하기 위해 생성된 결과물은 반드시 개인 단말기에 실시간으로 다운로드하여 보관해 주시기 바랍니다.
    """)


def show_admin_dashboard():
    """관리자 전용 사용량 통계 대시보드"""
    st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">ADMIN DASHBOARD</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">시스템 이용 통계 분석</span>
        <span class="qx-topbar-badge">Administrator Only</span>
    </div>
    """, unsafe_allow_html=True)

    df_login, df_section, df_all = tracker.get_summary_data()

    if df_login.empty and df_section.empty:
        st.info("📊 아직 수집된 이용 통계 데이터가 없습니다.")
        return

    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="qx-section-label">📅 일별 로그인 추이 (최근 7일)</div>', unsafe_allow_html=True)
        if not df_login.empty:
            fig_login = px.line(df_login, x="date", y="count", markers=True, 
                                title="일별 로그인 수 상세", template="plotly_white")
            fig_login.update_traces(line_color="#0F6CBD")
            st.plotly_chart(fig_login, use_container_width=True)
        else:
            st.caption("로그인 데이터 없음")

    with col2:
        st.markdown('<div class="qx-section-label">🍕 당일 섹션별 사용 점유율</div>', unsafe_allow_html=True)
        if not df_section.empty:
            fig_section = px.pie(df_section, values="count", names="section_name", 
                                 title="오늘의 기능별 사용 비중", color_discrete_sequence=px.colors.qualitative.Pastel)
            st.plotly_chart(fig_section, use_container_width=True)
        else:
            st.caption("오늘의 사용 데이터 없음")

    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown('<div class="qx-section-label">📋 최근 시스템 접근 로그 (Last 100)</div>', unsafe_allow_html=True)
    st.dataframe(df_all, use_container_width=True, hide_index=True)

def show_win_strategy_section():
    st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">과업 내용 체크 리스트</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">RFP 심층 분석 솔루션</span>
        <span class="qx-topbar-badge">Winning RFP Analysis</span>
    </div>
    """, unsafe_allow_html=True)

    # [v4.8] 도식 및 상세 안내
    with st.expander("📘 RFP 심층 분석 - 이용 방법 및 데이터 프로세스 가이드", expanded=False):
        st.markdown("### 🗺️ RFP 분석 프로세스")
        show_data_roadmap(1)
        st.markdown("""
        ### 🛠️ 이용 단계 및 상세 가이드
        1. **RFP 업로드:** 금년도 공고된 제안요청서(RFP)를 업로드합니다. (PDF, Word, TXT 지원)
        2. **차이 분석 (선택):** 작년 도 RFP가 있다면 함께 업로드하여 가감된 과업을 자동으로 비교합니다.
        3. **항목별 정밀 진단:** 
           - **참가 자격:** 입찰 자격 제한 요소 및 필수 실적 요건 여부 판별
           - **과업 범위:** 제안서 작성 시 누락되어서는 안 될 핵심 요구사항 추출
           - **기술적 평가:** 정보보안, 개인정보 처리 등 기술 점수 감점 요인 점검
        4. **전략 리포트 추출:** 분석 결과를 엑셀 및 워드 리포트로 저장하여 제안서 기초 자료로 활용합니다.
        """)
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown('<div class="qx-section-label">1. 금년도 RFP (필수)</div>', unsafe_allow_html=True)
        curr_file = st.file_uploader("올해 제안요청서", type=["pdf", "docx", "txt"], key="rfp_curr_up")
    with col2:
        st.markdown('<div class="qx-section-label">2. 직전 RFP (선택)</div>', unsafe_allow_html=True)
        prev_file = st.file_uploader("직전 회차 제안요청서", type=["pdf", "docx", "txt"], key="rfp_prev_up")

    # [v4.15] 초기 진입 화면 가이드 (Landing Card)
    if not curr_file and not st.session_state["rfp_results"]:
        st.markdown("""
<div class="qx-card" style="text-align:center; padding:2rem 2rem 2.5rem 2rem; margin-top: 1rem; margin-bottom: 2rem; min-height: 320px;">
    <div style="font-size:3rem; margin-bottom:0.5rem;">📝</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">
        RFP 심층 분석을 위해 문서를 업로드하세요
    </div>
    <div style="font-size:0.87rem; color:#8B96A9; margin-bottom:1.5rem;">
        올해 공고된 제안요청서를 업로드하면 과업 범위와 핵심 요구사항을 인공지능이 자동으로 분석합니다.
    </div>
    <div style="display:flex; justify-content:center; gap:1.5rem; flex-wrap:wrap;">
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">📋</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">과업 범위 추출</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">⚖️</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">참가 자격 검토</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🎯</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">수주 전략 정립</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
        st.markdown("<div style='margin-bottom: 5rem;'></div>", unsafe_allow_html=True)

    # 텍스트 추출 가이드
    if curr_file and curr_file.name != st.session_state["rfp_curr_name"]:
        with st.spinner("금년도 RFP 분석 준비 중..."):
            text, _ = extract_text(curr_file)
            st.session_state["rfp_curr_text"] = text
            st.session_state["rfp_curr_name"] = curr_file.name
    
    if prev_file and prev_file.name != st.session_state["rfp_prev_name"]:
        with st.spinner("직전 RFP 분석 준비 중..."):
            text, _ = extract_text(prev_file)
            st.session_state["rfp_prev_text"] = text
            st.session_state["rfp_prev_name"] = prev_file.name

    st.markdown("<hr>", unsafe_allow_html=True)
    
    if st.button("🚀 RFP 심층 분석 시작", type="primary", use_container_width=True):
        if not st.session_state["rfp_curr_text"]:
            st.error("금년도 RFP 문서를 먼저 업로드해 주세요.")
        else:
            perform_rfp_analysis()

    # 결과 표시 영역
    if st.session_state["rfp_results"]:
        show_security_notice()
        st.success(f"✅ **[{st.session_state['rfp_project_name']}]** 분석 완료")
        
        # 탭으로 결과 표시
        tab_names = [s["title"] for s in RFP_SECTIONS]
        tabs = st.tabs(tab_names)
        
        for i, sec in enumerate(RFP_SECTIONS):
            with tabs[i]:
                res = st.session_state["rfp_results"].get(sec["id"], "분석 결과가 없습니다.")
                # [v2.10] &nbsp; HTML 엔티티가 원문 텍스트로 표시되는 문제 보정
                res = res.replace("&nbsp;", " ")
                st.markdown(res, unsafe_allow_html=True)
        
        # 워드 다운로드 (RFP용) - 한 번 클릭으로 바로 다운로드
        st.markdown("<hr>", unsafe_allow_html=True)
        rfp_md = f"# RFP 분석 보고서: {st.session_state['rfp_project_name']}\n\n"
        for sec in RFP_SECTIONS:
            sec_text = st.session_state['rfp_results'].get(sec['id'], '').replace("&nbsp;", " ")
            rfp_md += f"## {sec['title']}\n\n{sec_text}\n\n"
        
        docx_file = export_to_docx(rfp_md)
        st.download_button(
            label="📝 RFP 분석 결과 워드 다운로드",
            data=docx_file,
            file_name=f"{st.session_state['rfp_project_name']}_RFP분석.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


def perform_rfp_analysis():
    curr_text = st.session_state["rfp_curr_text"]
    prev_text = st.session_state["rfp_prev_text"]
    
    # 컨텍스트 구성
    user_content = f"[금년도 문서]\n{rfp_utils.get_balanced_context(curr_text, 25000)}\n\n[직전 회차 문서]\n{rfp_utils.get_balanced_context(prev_text, 10000) if prev_text else '없음'}"
    
    # 사업명 감지
    project_name = rfp_utils.detect_project_name(curr_text)
    st.session_state["rfp_project_name"] = project_name
    
    progress_bar = st.progress(0, text="RFP 심층 분석 시작...")
    status_text = st.empty()
    
    total = len(RFP_SECTIONS)
    for i, sec in enumerate(RFP_SECTIONS):
        status_text.info(f"⏳ **{sec['title']}** 분석 중... ({i+1}/{total})")
        
        # analyzer의 run_analysis 활용
        result, err = run_analysis(
            sec["prompt"], 
            user_content, # report_text 인자로 보내야 프롬프트의 {report_text}가 치환됨
            model_name=st.session_state["selected_model"],
            auto_mode=st.session_state["auto_mode"]
        )
        
        if err:
            st.error(f"Error in {sec['title']}: {err}")
            st.session_state["rfp_results"][sec["id"]] = f"⚠️ 분석 실패: {err}"
        else:
            import re
            # 후처리: 모든 세미콜론(; , ；) 뒤의 공백을 포함하여 <br>로 변환 (줄바꿈 구현)
            processed_result = re.sub(r'[;；]\s*', '<br>', result)
            st.session_state["rfp_results"][sec["id"]] = processed_result
        
        progress_bar.progress((i + 1) / total)
    
    status_text.empty()
    progress_bar.empty()
    st.rerun()


# ── 사업체 명부 추출 UI ──
def show_unified_business_search():
    """AI 추천 및 공공데이터 연동 사업체 명부 추출 시스템 UI (4-Step Pipeline)"""
    st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">사업체 명부 추출</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">단계별 업체 발굴 솔루션</span>
        <span class="qx-topbar-badge">G2B · NPS · NHIS 통합</span>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("📘 단계별 이용 안내 (4-Step Pipeline)", expanded=False):
        st.markdown("""
        ### 🗺️ 파이프라인 단계
        - **1단계 (지역 선택):** 시도 및 시군구를 지정하여 해당 지역 전체 사업장 기초 데이터를 수집합니다.
        - **2단계 (현황 집계):** 선택한 지역 내 업체들의 업종분포와 근로자 수 규모를 한눈에 파악합니다.
        - **3단계 (추출 기준 설정):** 집계된 통계를 바탕으로 타겟할 업종이나 근로자 수 스펙을 필터링합니다.
        - **4단계 (상세 통합 추출):** 최종 선별된 타겟 업체들에 대해 공공망(NPS/G2B)을 찔러 실시간 전화번호, 상세 기업 정보를 조립해 냅니다.
        """)

    # --- Step 1 ---
    st.markdown('<div class="qx-section-label">STEP 1. 지역 탐색</div>', unsafe_allow_html=True)
    col_s1, col_sigg, col_btn1 = st.columns([1.5, 1.5, 1])
    with col_s1:
        sido = st.selectbox("시도 선택", options=BIZ_SIDO_LIST)
    with col_sigg:
        SIGG_MAP = {
            "전체": ["전체"],
            "서울특별시": ["전체", "강남구", "강동구", "강북구", "강서구", "관악구", "광진구", "구로구", "금천구", "노원구", "도봉구", "동대문구", "동작구", "마포구", "서대문구", "서초구", "성동구", "성북구", "송파구", "양천구", "영등포구", "용산구", "은평구", "종로구", "중구", "중랑구"],
            "부산광역시": ["전체", "강서구", "금정구", "기장군", "남구", "동구", "동래구", "부산진구", "북구", "사상구", "사하구", "서구", "수영구", "연제구", "영도구", "중구", "해운대구"],
            "대구광역시": ["전체", "군위군", "남구", "달서구", "달성군", "동구", "북구", "서구", "수성구", "중구"],
            "인천광역시": ["전체", "강화군", "계양구", "남동구", "동구", "미추홀구", "부평구", "서구", "연수구", "옹진군", "중구"],
            "광주광역시": ["전체", "광산구", "남구", "동구", "북구", "서구"],
            "대전광역시": ["전체", "대덕구", "동구", "서구", "유성구", "중구"],
            "울산광역시": ["전체", "남구", "동구", "북구", "울주군", "중구"],
            "세종특별자치시": ["전체"],
            "경기도": ["전체", "가평군", "고양시", "과천시", "광명시", "광주시", "구리시", "군포시", "김포시", "남양주시", "동두천시", "부천시", "성남시", "수원시", "시흥시", "안산시", "안성시", "안양시", "양주시", "양평군", "여주시", "연천군", "오산시", "용인시", "의왕시", "의정부시", "이천시", "파주시", "평택시", "포천시", "하남시", "화성시"],
            "강원특별자치도": ["전체", "강릉시", "고성군", "동해시", "삼척시", "속초시", "양구군", "양양군", "영월군", "원주시", "인제군", "정선군", "철원군", "춘천시", "태백시", "평창군", "홍천군", "화천군", "횡성군"],
            "충청북도": ["전체", "괴산군", "단양군", "보은군", "영동군", "옥천군", "음성군", "제천시", "증평군", "진천군", "청주시", "충주시"],
            "충청남도": ["전체", "계룡시", "공주시", "금산군", "논산시", "당진시", "보령시", "부여군", "서산시", "서천군", "아산시", "예산군", "천안시", "청양군", "태안군", "홍성군"],
            "전북특별자치도": ["전체", "고창군", "군산시", "김제시", "남원시", "무주군", "부안군", "순창군", "완주군", "익산시", "임실군", "장수군", "전주시", "정읍시", "진안군"],
            "전라남도": ["전체", "강진군", "고흥군", "곡성군", "광양시", "구례군", "나주시", "담양군", "목포시", "무안군", "보성군", "순천시", "신안군", "여수시", "영광군", "영암군", "완도군", "장성군", "장흥군", "진도군", "함평군", "해남군", "화순군"],
            "경상북도": ["전체", "경산시", "경주시", "고령군", "구미시", "김천시", "문경시", "봉화군", "상주시", "성주군", "안동시", "영덕군", "영양군", "영주시", "영천시", "예천군", "울릉군", "울진군", "의성군", "청도군", "청송군", "칠곡군", "포항시"],
            "경상남도": ["전체", "거제시", "거창군", "고성군", "김해시", "남해군", "밀양시", "사천시", "산청군", "양산시", "의령군", "진주시", "창녕군", "창원시", "통영시", "하동군", "함안군", "함양군", "합천군"],
            "제주특별자치도": ["전체", "서귀포시", "제주시"]
        }
        sigg_options = SIGG_MAP.get(sido, ["전체"])
        sigg = st.selectbox("시/군/구 선택", options=sigg_options)
        
    with col_btn1:
        st.markdown("<div style='margin-top: 28px'></div>", unsafe_allow_html=True)
        if st.button("📊 1단계: 지역 현황 집계", use_container_width=True, type="secondary"):
            if sido == "전체":
                st.warning("전체 지역을 집계하기에는 데이터가 방대하므로 특정 시/도를 선택해 주시기 바랍니다.")
            else:
                with st.spinner("건강보험 기초 데이터를 동기화하고 지역별 현황을 집계 중입니다..."):
                    nhis_df = st.session_state.get("biz_nhis_dataset")
                    if nhis_df is None or nhis_df.empty:
                        st.toast("건강보험 전국망 기초 데이터를 1회 동기화합니다 (약 15초 소요).", icon="⏳")
                        from api.nhis_api import download_nhis_dataset
                        nhis_df = download_nhis_dataset(SERVICE_KEY)
                        st.session_state["biz_nhis_dataset"] = nhis_df
                    
                    regional_df = nhis_df.copy()
                    addr_cols = [c for c in regional_df.columns if "주소" in c or "addr" in c.lower()]
                    
                    if addr_cols:
                        addr_col = addr_cols[0]
                        regional_df = regional_df[regional_df[addr_col].str.contains(sido, na=False)]
                        if sigg != "전체":
                            regional_df = regional_df[regional_df[addr_col].str.contains(sigg, na=False)]
                    
                    st.session_state["biz_step1_df"] = regional_df
                    # 새로운 지역으로 집계했으므로, 4단계 조회 결과는 파기
                    st.session_state.pop("biz_step4_results", None)

    st.markdown("<hr>", unsafe_allow_html=True)

    # --- Step 2 & 3 ---
    if "biz_step1_df" in st.session_state:
        regional_df = st.session_state["biz_step1_df"]
        total_comps = len(regional_df)
        
        st.markdown(f'<div class="qx-section-label">STEP 2. 표준산업분류별 사업체수 집계 (선택 지역 총 {total_comps:,}개 업체)</div>', unsafe_allow_html=True)
        
        # 근로자수 수치화
        if "직장가입자수" in regional_df.columns:
            regional_df["근로자수_num"] = pd.to_numeric(regional_df["직장가입자수"], errors="coerce").fillna(0)
        else:
            regional_df["근로자수_num"] = 0
        # 업종이 없는 경우 '정보없음' 처리
        if "업종코드" in regional_df.columns:
            regional_df["업종코드_fill"] = regional_df["업종코드"].fillna("미기재")
        else:
            regional_df["업종코드_fill"] = "미기재"
        
        # 집계: 업체 수는 언제나 집계하고, 근로자수는 컴럼 존재 시에만
        agg_group = regional_df.groupby("업종코드_fill")
        agg_count = agg_group.size().rename("사업체수")
        if "근로자수_num" in regional_df.columns:
            agg_sum = agg_group["근로자수_num"].sum().rename("총근로자수")
            agg_df = pd.concat([agg_count, agg_sum], axis=1).reset_index()
        else:
            agg_df = agg_count.reset_index()
            agg_df["총근로자수"] = 0
        agg_df = agg_df.rename(columns={"업종코드_fill": "공단 업종코드"})
        agg_df = agg_df.sort_values("사업체수", ascending=False)
        
        # 3열 구성 (가운데 표, 우측 필터)
        col_t1, col_t2 = st.columns([1.5, 2.5])
        
        with col_t1:
            st.dataframe(agg_df, use_container_width=True, hide_index=True, height=250)
            st.caption("※ 출처: 국민건강보험공단 사업장관리 현황 (단위: 명, 개)")
            
        with col_t2:
            st.markdown('<div style="padding: 1rem; background: #FAFBFE; border: 1px solid #E5E9F0; border-radius: 8px;">', unsafe_allow_html=True)
            st.markdown('<div class="qx-section-label">STEP 3. 추출 기준 설정</div>', unsafe_allow_html=True)
            
            filter_type = st.radio("추출 기준 (라디오 버튼)", ["특정 산업분류 한정", "전체 업종 (근로자수만 지정)"], horizontal=True)
            
            target_inds = []
            min_workers = 0
            
            subcol1, subcol2 = st.columns(2)
            with subcol1:
                if filter_type == "특정 산업분류 한정":
                    target_inds = st.multiselect("대상 업종코드 선택", options=agg_df["공단 업종코드"].tolist(), placeholder="복수 선택 가능")
            with subcol2:
                min_workers = st.number_input("최소 근로자 수 (명 이상)", min_value=0, value=1)
                
            st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("<hr>", unsafe_allow_html=True)
        
        # --- Step 4 ---
        if st.button("🚀 STEP 4. 최종 전체 추출 (NPS, G2B 상세 연락처 연동)", type="primary", use_container_width=True):
            filtered_df = regional_df.copy()
            
            if filter_type == "특정 산업분류 한정" and target_inds:
                filtered_df = filtered_df[filtered_df["업종코드_fill"].isin(target_inds)]
                
            filtered_df = filtered_df[filtered_df["근로자수_num"] >= min_workers]
            # _brn 컴럼 안전하게 추출 (난 경우 사업자등록번호 컴럼 활용)
            if "_brn" in filtered_df.columns:
                brn_list = filtered_df["_brn"].dropna().tolist()
            elif "사업자등록번호" in filtered_df.columns:
                brn_list = (
                    filtered_df["사업자등록번호"]
                    .astype(str).str.replace("-", "", regex=False).str.zfill(10)
                    .dropna().tolist()
                )
            else:
                brn_list = []
            
            if not brn_list:
                st.warning("조건을 만족하는 업체가 0건입니다. 추출 기준을 완화해 보세요.")
            else:
                st.success(f"조건을 만족하는 총 {len(brn_list):,}개 업체 중, 최대 100건의 상세(연락처 등) 데이터를 정밀 취합합니다.")
                with st.spinner("NPS, G2B 공공망과 통신하여 연락처 및 심층 데이터를 수집 중입니다..."):
                    from api.biz_search_api import batch_fetch_by_brns
                    results = batch_fetch_by_brns(brn_list, SERVICE_KEY, nhis_df=st.session_state["biz_nhis_dataset"])
                    st.session_state["biz_step4_results"] = results
                    
    # 결과 표시
    if "biz_step4_results" in st.session_state:
        results = st.session_state["biz_step4_results"]
        
        st.markdown('<div class="qx-section-label">FINAL EXTRACTION RESULTS</div>', unsafe_allow_html=True)
        
        if not results:
            st.info("조건을 만족하는 업체의 상세 정보를 API에서 불러오지 못했습니다.")
            return

        col_ctrl1, col_ctrl2 = st.columns([1, 4])
        with col_ctrl1:
            if st.button("✅ 전체 선택", use_container_width=True):
                st.session_state["biz_selected_indices"] = list(range(len(results)))
                st.rerun()
        with col_ctrl2:
            if st.button("❌ 전체 해제", use_container_width=True):
                st.session_state["biz_selected_indices"] = []
                st.rerun()

        if "biz_selected_indices" not in st.session_state:
            st.session_state["biz_selected_indices"] = list(range(len(results)))

        display_data = []
        for i, res in enumerate(results):
            display_data.append({
                "선택": i in st.session_state["biz_selected_indices"],
                "업체명": res["corp_name"],
                "업종": res["industry"],
                "주소": res["address"],
                "가입자(NPS)": f"{res['nps_subscriber']:,}명",
                "가입자(NHIS)": f"{res['nhis_subscriber']:,}명",
                "전화번호": res["tel"],
                "기업규모": res["corp_size"],
                "데이터출처": ", ".join(res["source"])
            })
        
        df_display = pd.DataFrame(display_data)
        
        edited_df = st.data_editor(
            df_display, hide_index=True, use_container_width=True,
            column_config={"선택": st.column_config.CheckboxColumn("선택", default=True)},
            key="biz_results_editor"
        )
        
        selected_rows = edited_df[edited_df["선택"] == True]
        
        if not selected_rows.empty:
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                selected_rows.drop(columns=["선택"]).to_excel(writer, index=False, sheet_name='Search_Results')
            output.seek(0)
            
            st.download_button(
                label=f"📥 선택된 {len(selected_rows)}건 최종 리스트 다운로드 (Excel)", data=output,
                file_name=f"업체명부_4단계추출결과_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True, type="primary"
            )
        else:
            st.info("다운로드할 업체를 선택해 주세요.")


# ── 기업체 일반 현황 행정자료 비교 UI ──
def show_business_info_crawling():
    """공공데이터포털 API를 활용한 기업체 일반 현황 행정자료 비교 시스템 UI"""
    import time

    # 상단 헤더
    st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">행정자료 비교</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">기업체 일반 현황</span>
        <span class="qx-topbar-badge">공공데이터 API 연동</span>
    </div>
    """, unsafe_allow_html=True)

    # ── 사용방법 안내
    with st.expander("📘 사용 방법 안내", expanded=False):
        st.markdown("""
        ### 🛠️ 이용 방법
        1. **API 키 설정:** 공공데이터포털 API 서비스 키가 Streamlit Secrets에 등록되어 있어야 합니다.
        2. **엑셀 파일 업로드:** 조회할 사업체 목록이 담긴 엑셀 파일(.xlsx, .xls)을 업로드합니다.
        3. **컬럼 매핑:** 업로드된 파일의 컬럼 중 사업자등록번호, 회사명, 대표자명, 주소에 해당하는 컬럼을 지정합니다.
        4. **조회 항목 선택:** 아래 4개 기관의 데이터 중 필요한 항목을 선택합니다.
        5. **조회 실행:** '조회 시작' 버튼을 클릭하여 API 조회를 수행합니다.
        6. **결과 다운로드:** 조회 결과를 엑셀 파일로 다운로드합니다.

        | 출처 기관 | 데이터 항목 | 조회 키 |
        | :--- | :--- | :--- |
        | **국민연금 (NPS)** | 가입자수, 당월고지금액, 신규/상실 가입자수, 평균소득(추정) | 회사명 + 사업자번호/주소 |
        | **건강보험 (NHIS)** | 가입자수, 당월고지금액 | 사업자등록번호 (10자리) |
        | **금융위 기업기본** | 법인번호, 1인평균급여, 종업원수, 업종명, 설립일 | 회사명 + 사업자번호/주소 |
        | **금융위 기업재무** | 매출액, 영업이익, 당기순이익, 총자산, 총부채, 자본금 | 법인등록번호 (연계) |
        | **국세청 (NTS)** | 사업자 상태(계속/휴업/폐업), 과세유형, 폐업일자 | 사업자등록번호 (10자리) |
        | **전자공시 (DART)** | 마스킹된 사업자번호(BRN) 복원, 대표자명, 주소 정합성 | 법인명/법인번호 |
        | **나라장터 (G2B)** | 업종명, 전화번호, 기업구분(대/중/소) | 사업자등록번호 (10자리) |

        ### 🛡️ 데이터 매치 및 유사도 산정 방식
        본 시스템은 사업자등록번호가 없거나 불완전한 자료의 정확한 매칭을 위해 다음과 같은 지능형 매칭 엔진을 사용합니다.

        1. **다면적 검색 로직 (회사명 + 주소 조합)**
           - 회사명으로 1차 검색 후, 결과가 여러 개일 경우 사용자가 입력한 **'주소'** 유사도를 분석하여 가장 일치하는 사업장을 자동으로 선택합니다.
        
        2. **유사도(%) 가중치 산정 방식**
           최종 유사도는 아래 3가지 항목의 가중 합산으로 결정됩니다 (총합 100점).
           - **매칭 성공 (30점):** 제공된 키로 공공데이터 조회 성공 시 기본 점수 부여
           - **회사명 일치도 (40점):** 엑셀의 회사명과 API 검색 결과 이름의 텍스트 유사도
           - **주소 일치도 (30점):** 엑셀의 주소와 API 사업장 주소의 텍스트 유사도
           - ※ **필터링:** 설정된 유사도 임계값(기본 50%) 미만인 결과는 오매칭 방지를 위해 자동 제외됩니다.

        3. **항목별 개별 매핑 (Index-based Mapping)**
           - 업로드된 엑셀의 각 행(Row)에 고유 번호를 부여하여 조회합니다.
           - 동일한 이름을 가진 업체가 여러 개 있거나 사업자번호가 누락되어도, 데이터 부딪힘(Collision) 없이 각 업체에 맞는 정보를 정확히 매치합니다.

        ### 🔑 API 키 발급 방법
        1. [공공데이터포털(data.go.kr)](https://data.go.kr) 회원가입
        2. 아래 **4개 API** 활용 신청 (동일한 서비스키 사용 가능):
           - 국민연금 가입 사업장 내역
           - 건강보험공단 사업장관리 현황
           - **금융위원회 기업기본정보** (`GetCorpBasicInfoService_V2`)
           - **금융위원회 기업재무정보** (`GetFinaStatInfoService_V2`)
           - **국세청 사업자등록정보 진위확인 및 상태조회** (`v1/status`)
        3. 발급받은 **Decoding Key**를 Streamlit Secrets에 등록

        ### ⚠️ 주의사항
        - 국민연금 API는 **가입자 3인 이상 법인사업장**, **10인 이상 개인사업장**만 조회 가능합니다.
        - 개발 계정 기준 **일 10,000건** 조회 제한이 있습니다.
        - 건강보험공단 데이터는 **정기 갱신** 기반이며, 실시간 데이터가 아닙니다.
        - **국세청 데이터**는 실시간 휴/폐업 상태 및 과세유형을 제공합니다.
        - 금융위원회 기업정보는 **공시 대상 법인(상장사·외감법인 등)** 위주로 조회되며, 소규모 개인사업장은 결과가 없을 수 있습니다.
        - 재무정보는 **직전 사업연도** 기준이며, 미공시 시 2년 전 데이터를 자동 조회합니다.
        - 대표자명은 개인정보보호를 위해 API 조회 키로 사용하지 않습니다.
        """)

    # ── 데이터 수집 흐름도
    st.markdown("""
<div class="qx-card" style="padding: 1.5rem 2rem;">
<div class="qx-card-title" style="margin-bottom: 1.2rem; font-size: 1.1rem;">📊 기업체 행정정보 수집 및 매칭 프로세스</div>
<div style="display:flex; flex-direction:column; gap:1rem;">
<!-- 상단: 입력 및 엔진 -->
<div style="display:flex; align-items:center; justify-content:center; gap:1rem; margin-bottom:0.5rem;">
<div style="background:#0F6CBD; color:white; padding:0.8rem 1.4rem; border-radius:10px; font-weight:600; font-size:0.9rem; min-width:160px; text-align:center; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
📋 엑셀 업로드<br><span style="font-size:0.75rem; font-weight:400; opacity:0.9;">사업자번호 · 회사명 · 주소</span>
</div>
<div style="color:#8B96A9; font-size:1.5rem; font-weight:bold;">→</div>
<div style="background:#F4F6F9; border:2px solid #0F6CBD; padding:0.8rem 1.4rem; border-radius:10px; font-size:0.9rem; color:#0F6CBD; font-weight:700; text-align:center; min-width:160px;">
⚙️ AI 지능형 매칭 엔진<br><span style="font-size:0.75rem; font-weight:400; color:#3D4F6B;">데이터 정제 및 식별자 해결</span>
</div>
<div style="color:#8B96A9; font-size:1.5rem; font-weight:bold;">→</div>
<div style="background:#059669; color:white; padding:0.8rem 1.4rem; border-radius:10px; font-weight:600; font-size:0.9rem; min-width:160px; text-align:center; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
📥 결과 엑셀 다운로드<br><span style="font-size:0.75rem; font-weight:400; opacity:0.9;">행정통합 리포트 생성</span>
</div>
</div>
<!-- 중앙: 연결선 (화살표) -->
<div style="display:flex; justify-content:center; font-size:1.2rem; color:#8B96A9; margin:-0.5rem 0;">
<div style="width:160px;"></div>
<div style="text-align:center; padding-left:20px;">⇅</div>
<div style="width:160px;"></div>
</div>
<!-- 하단: 4개 채널 (6개 기관 연동) -->
<div style="display:flex; gap:0.7rem; flex-wrap:wrap;">
<div style="flex:1; min-width:140px; background:#EEF4FD; border:1px solid #BDD7F5; border-radius:12px; padding:0.8rem; text-align:center; transition: transform 0.2s;">
<div style="font-size:1.3rem; margin-bottom:0.3rem;">🏢</div>
<div style="font-size:0.8rem; font-weight:700; color:#0F6CBD; margin-bottom:0.3rem;">NPS / NHIS</div>
<div style="font-size:0.7rem; color:#475569; line-height:1.5;"><b>복지·고용</b><br>인원수 · 고지금액<br>업종 · 가입상태</div>
</div>
<div style="flex:1; min-width:140px; background:#FEF3E2; border:1px solid #FCD34D; border-radius:12px; padding:0.8rem; text-align:center; border-left: 4px solid #F59E0B;">
<div style="font-size:1.3rem; margin-bottom:0.3rem;">⚖️</div>
<div style="font-size:0.8rem; font-weight:700; color:#B45309; margin-bottom:0.3rem;">NTS (국세청)</div>
<div style="font-size:0.7rem; color:#475569; line-height:1.5;"><b>실시간 휴/폐업</b><br>계속/휴업/폐업<br>과세유형 · 폐업일</div>
</div>
<div style="flex:1; min-width:140px; background:#F0FDF4; border:1px solid #86EFAC; border-radius:12px; padding:0.8rem; text-align:center;">
<div style="font-size:1.3rem; margin-bottom:0.3rem;">🏛️</div>
<div style="font-size:0.8rem; font-weight:700; color:#059669; margin-bottom:0.3rem;">FSS (금융위)</div>
<div style="font-size:0.7rem; color:#475569; line-height:1.5;"><b>재무·경영</b><br>재무제표 · 매출액<br>평균연봉 · 종업원</div>
</div>
<div style="flex:1; min-width:140px; background:#F5F3FF; border:1px solid #C4B5FD; border-radius:12px; padding:0.8rem; text-align:center; border-left: 4px solid #8B5CF6;">
<div style="font-size:1.3rem; margin-bottom:0.3rem;">🔍</div>
<div style="font-size:0.8rem; font-weight:700; color:#7C3AED; margin-bottom:0.3rem;">DART / G2B</div>
<div style="font-size:0.7rem; color:#475569; line-height:1.5;"><b>기업 식별·공공</b><br>BRN 복원(DART)<br>산업군·연락처(G2B)</div>
</div>
</div>
</div>
</div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── API 키 로드 (Streamlit Cloud secrets 우선, api_utils.SERVICE_KEY 차선)
    api_service_key = ""
    try:
        api_service_key = st.secrets.get("api_keys", {}).get("DATA_GO_KR_KEY", "")
    except Exception:
        pass
    
    if not api_service_key:
        api_service_key = SERVICE_KEY

    if api_service_key:
        st.success("✅ 공공데이터포털 API 키가 연결되어 있습니다.", icon="🔑")
    else:
        st.error(
            "⚠️ **API 키가 설정되지 않았습니다.**\n\n"
            "Streamlit Cloud의 **Secrets** 설정에서 아래 형식으로 API 키를 등록해주세요:\n\n"
            "```\n[api_keys]\nDATA_GO_KR_KEY = \"발급받은_서비스키\"\n```\n\n"
            "공공데이터포털(data.go.kr)에서 API 활용 신청 후 발급받을 수 있습니다.",
            icon="🔑",
        )
        st.stop()

    st.markdown("<hr>", unsafe_allow_html=True)

    # ── Step 2: 엑셀 파일 업로드
    st.markdown('<div class="qx-section-label">STEP 2 · 사업체 목록 업로드</div>', unsafe_allow_html=True)
    st.caption("📋 **사업체 목록**이 담긴 엑셀 파일(.xlsx, .xls)을 아래에 업로드하세요.")

    uploaded_excel = st.file_uploader(
        "사업체 목록 엑셀 파일 선택 (XLSX, XLS)",
        type=["xlsx", "xls"],
        label_visibility="collapsed",
        key="biz_excel_uploader",
    )

    # 표준 양식 다운로드
    import io as _io
    _tpl_df = pd.DataFrame({
        "사업자등록번호": ["123-45-67890"],
        "회사명": ["(주)예시기업"],
        "대표자명": ["홍길동"],
        "주소": ["서울특별시 종로구 세종대로 209"],
    })
    _tpl_buf = _io.BytesIO()
    with pd.ExcelWriter(_tpl_buf, engine="openpyxl") as _w:
        _tpl_df.to_excel(_w, index=False, sheet_name="사업체목록")
    _tpl_buf.seek(0)
    st.download_button(
        label="📋 표준 엑셀 양식 다운로드",
        data=_tpl_buf,
        file_name="사업체정보_표준양식.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="btn_biz_template_dl",
    )

    df = None
    if uploaded_excel is not None:
        try:
            df = load_excel(uploaded_excel)
            st.success(f"✅ 파일 로드 완료: **{uploaded_excel.name}** ({len(df):,}건)", icon="📊")

            # 미리보기
            st.markdown("**📋 데이터 미리보기 (상위 5행)**")
            st.dataframe(df.head(5), use_container_width=True, hide_index=True)
        except Exception as e:
            st.error(f"엑셀 파일 읽기 실패: {e}")
            df = None

    if df is not None:
        st.markdown("<hr>", unsafe_allow_html=True)

        # ── 컬럼 매핑
        st.markdown('<div class="qx-section-label">컬럼 매핑</div>', unsafe_allow_html=True)
        st.caption("업로드된 엑셀 파일의 컬럼을 아래 항목에 매핑하세요.")

        col_options = ["(선택 안 함)"] + list(df.columns)

        col_m1, col_m2 = st.columns(2)
        with col_m1:
            brn_col = st.selectbox(
                "사업자등록번호 *",
                col_options,
                index=_auto_detect_col(col_options, ["사업자", "등록번호", "사업자번호"]),
                key="biz_col_brn",
            )
        with col_m2:
            name_col = st.selectbox(
                "회사명",
                col_options,
                index=_auto_detect_col(col_options, ["회사명", "사업장명", "업체명", "상호"]),
                key="biz_col_name",
            )
        col_m3, col_m4 = st.columns(2)
        with col_m3:
            ceo_col = st.selectbox(
                "대표자명",
                col_options,
                index=_auto_detect_col(col_options, ["대표자", "대표"]),
                key="biz_col_ceo",
            )
        with col_m4:
            addr_col = st.selectbox(
                "주소",
                col_options,
                index=_auto_detect_col(col_options, ["주소", "소재지"]),
                key="biz_col_addr",
            )
        
        # [v9.5] 법인등록번호(CRNO) 컬럼 매핑 추가 (데이터 추출 확률 향상용)
        crno_col = st.selectbox(
            "법인등록번호 (선택 시 추출 확률 향상)",
            col_options,
            index=_auto_detect_col(col_options, ["법인번호", "법인등록번호"]),
            key="biz_col_crno",
        )

        if brn_col == "(선택 안 함)":
            st.warning("⚠️ **사업자등록번호** 컬럼은 필수입니다. 매핑해주세요.")
        else:
            # ── 회사명 정제
            cleaned_name_col = None
            if name_col != "(선택 안 함)" and name_col in df.columns:
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown('<div class="qx-section-label">회사명 자동 정제</div>', unsafe_allow_html=True)

                if "회사명_정제" not in df.columns:
                    df, clean_stats = clean_company_names_bulk(df, name_col)
                    st.session_state["biz_cleaned_df"] = df
                    st.session_state["biz_clean_stats"] = clean_stats
                else:
                    clean_stats = st.session_state.get("biz_clean_stats", {})

                if clean_stats:
                    col_c1, col_c2, col_c3 = st.columns(3)
                    with col_c1:
                        st.metric("전체 건수", f"{clean_stats['total']:,}")
                    with col_c2:
                        st.metric("정제 변경", f"{clean_stats['changed']:,}건")
                    with col_c3:
                        st.metric("변경 없음", f"{clean_stats['unchanged']:,}건")

                    # 유형별 카운트
                    if clean_stats.get('type_counts'):
                        type_str = " · ".join(f"{k}: {v:,}건" for k, v in clean_stats['type_counts'].items())
                        st.caption(f"제거된 유형별 건수: {type_str}")

                    # 변경 샘플
                    if clean_stats.get('samples'):
                        with st.expander(f"정제 전·후 비교 샘플 ({len(clean_stats['samples'])}건)", expanded=False):
                            import pandas as _pd
                            sample_df = _pd.DataFrame(
                                clean_stats['samples'],
                                columns=["원본 회사명", "정제된 회사명"]
                            )
                            st.dataframe(sample_df, use_container_width=True, hide_index=True)

                cleaned_name_col = "회사명_정제"
                st.success(f"✅ '{name_col}' → '회사명_정제' 컬럼이 자동 생성되었습니다. 정제된 이름으로 매칭합니다.")

            # ── 주소 정제 (v8.0)
            cleaned_addr_col = None
            if addr_col != "(선택 안 함)" and addr_col in df.columns:
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown('<div class="qx-section-label">주소 정제 및 지역 분리</div>', unsafe_allow_html=True)

                if "주소_정제" not in df.columns:
                    df = clean_addresses_bulk(df, addr_col)
                    st.session_state["biz_cleaned_df"] = df
                
                # 정제 결과 샘플 표시
                c_c1, c_c2, c_c3 = st.columns(3)
                with c_c1:
                    st.metric("시도 구분", f"{df['시도'].nunique()}종")
                with c_c2:
                    st.metric("시군구 구분", f"{df['시군구'].nunique()}종")
                
                with st.expander("주소 정제 및 분리 결과 샘플 (상위 5건)", expanded=False):
                    st.dataframe(df[[addr_col, "주소_정제", "시도", "시군구"]].head(5), use_container_width=True, hide_index=True)

                cleaned_addr_col = "주소_정제"
                st.info("💡 우편번호 제거 및 시도/시군구 분리가 완료되었습니다. 행정데이터와 비교 시 활용됩니다.")

            st.markdown("<hr>", unsafe_allow_html=True)

            # ── Step 3: 조회 항목 및 기준 선택
            st.markdown('<div class="qx-section-label">STEP 3 · 조회 항목 및 추출 기준 설정</div>', unsafe_allow_html=True)
            
            # [v8.6] 추출 기준 선택 버튼을 더 넓고 명확하게 배치
            st.markdown("""
<div style="background:#F8FAFC; border:1px solid #E2E8F0; border-radius:8px; padding:1rem; margin-bottom:1rem;">
    <div style="font-size:0.85rem; font-weight:600; color:#475569; margin-bottom:0.5rem;">📊 데이터 집계 및 추출 기준 선택</div>
</div>
            """, unsafe_allow_html=True)
            
            agg_mode = st.radio(
                "추출 기준 선택",
                ["사업체 기준 (Establishment)", "기업체 기준 (Enterprise)"],
                index=0,
                horizontal=True,
                help="사업체 기준은 개별 지점별로 결과를 보여주며, 기업체 기준은 동일 법인번호를 가진 지점들을 하나로 합산하여 요약해 줍니다.",
                key="extraction_criteria_selector",
                label_visibility="collapsed"
            )
            
            if agg_mode == "기업체 기준 (Enterprise)":
                st.info("🏢 **기업체 기준 집계:** 동일한 법인번호를 가진 여러 사업체(지점)의 데이터를 합산하여 하나의 법인 단위로 결과를 제공합니다.")
            else:
                st.caption("📑 **사업체 기준 조회:** 각 행별로 독립적인 행정자료 매칭 결과를 제공합니다.")
                
            st.markdown("<br>", unsafe_allow_html=True)
            
            # 선택 항목 초기화
            selected_nps, selected_nhis, selected_nts, selected_g2b = [], [], [], []

            col_nps, col_nhis, col_nts, col_g2b = st.columns(4)

            with col_nps:
                st.markdown("""
<div class="qx-card">
    <div class="qx-card-title">🏢 국민연금공단</div>
    <div style="font-size:0.78rem; color:#8B96A9; margin-bottom:0.5rem;">국민연금 가입 사업장 내역</div>
</div>
""", unsafe_allow_html=True)
                for field in NPS_SELECTABLE_FIELDS:
                    label = NPS_FIELD_LABELS.get(field, field)
                    if field == "jnngpCnt":
                        label += " 💡" # 종사자수 강조
                    if field == "avgBasSalary":
                        label += " 💰" # 매출/소득 관련
                    if st.checkbox(label, value=(field != "avgBasSalary"), key=f"nps_{field}"):
                        selected_nps.append(field)

            with col_nhis:
                st.markdown("""
<div class="qx-card">
    <div class="qx-card-title">🏥 건강보험공단</div>
    <div style="font-size:0.78rem; color:#8B96A9; margin-bottom:0.1rem;">사업장관리 현황</div>
</div>
""", unsafe_allow_html=True)
                if brn_col == "(선택 안 함)":
                    st.info("💡 **사업자번호가 없다면**, '회사명'과 '주소' 컬럼을 모두 매핑해 주세요.")
                
                nhis_year = st.selectbox(
                    "데이터 시점 선택", 
                    options=list(NHIS_ENDPOINTS.keys()),
                    index=0,
                    key="nhis_year_selector"
                )
                selected_nhis_uddi = NHIS_ENDPOINTS[nhis_year]
                for field in NHIS_SELECTABLE_FIELDS:
                    label = NHIS_FIELD_LABELS.get(field, field)
                    if field == "직장가입자수":
                        label += " 💡" # 종업원수 강조
                    if st.checkbox(label, value=True, key=f"nhis_{field}"):
                        selected_nhis.append(field)

            with col_nts:
                st.markdown("""
<div class="qx-card">
    <div class="qx-card-title">⚖️ 국세청 (NTS)</div>
    <div style="font-size:0.78rem; color:#8B96A9; margin-bottom:0.1rem;">진위확인 및 상태조회</div>
</div>
""", unsafe_allow_html=True)
                nts_fields = [
                    ("status", "사업자상태 🚩"), # 휴폐업정보 강조
                    ("tax_type", "과세유형 (일반/간이 등)"),
                    ("end_dt", "폐업일자 📅"),
                ]
                for field_id, label in nts_fields:
                    if st.checkbox(label, value=True, key=f"nts_{field_id}"):
                        selected_nts.append(field_id)

            with col_g2b:
                st.markdown("""
<div class="qx-card">
    <div class="qx-card-title">🔍 나라장터 (G2B)</div>
    <div style="font-size:0.78rem; color:#8B96A9; margin-bottom:0.5rem;">조달청 업체정보 및 업종</div>
</div>
""", unsafe_allow_html=True)
                for field in G2B_SELECTABLE_FIELDS:
                    label = G2B_FIELD_LABELS.get(field, field)
                    # 강조는 이미 label에 포함되어 있을 수 있음 (전화번호 📞 등)
                    if st.checkbox(label, value=True, key=f"g2b_{field}"):
                        selected_g2b.append(field)

            # ── 금융위원회 기업정보 (3번째 컬럼)
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown('<div class="qx-section-label">금융위원회 기업정보 (자동 조회)</div>', unsafe_allow_html=True)

            fss_available = (name_col != "(선택 안 함)")
            if not fss_available:
                st.warning("⚠️ 금융위원회 기업정보를 조회하려면 **회사명** 컬럼을 매핑해야 합니다.")

            col_fss_corp, col_fss_fina, col_empty = st.columns(3)

            with col_fss_corp:
                st.markdown("""
<div class="qx-card">
    <div class="qx-card-title">🏛️ 금융위원회 (기업정보)</div>
    <div style="font-size:0.78rem; color:#8B96A9; margin-bottom:0.5rem;">회사명 → 종업원수, 평균급여 등 자동 조회</div>
</div>
""", unsafe_allow_html=True)
                selected_fss_corp = []
                for field in FSS_CORP_SELECTABLE_FIELDS:
                    label = FSS_CORP_FIELD_LABELS.get(field, field)
                    if field == "enpEmpeCnt":
                        label += " 💡" # 종업원수 강조
                    checked = st.checkbox(label, value=fss_available, key=f"fss_corp_{field}", disabled=not fss_available)
                    if checked and fss_available:
                        selected_fss_corp.append(field)

            with col_fss_fina:
                st.markdown("""
<div class="qx-card">
    <div class="qx-card-title">📊 금융위원회 (재무정보)</div>
    <div style="font-size:0.78rem; color:#8B96A9; margin-bottom:0.5rem;">법인번호 연동 → 매출액, 영업이익 등</div>
</div>
""", unsafe_allow_html=True)
                selected_fss_fina = []
                for field in FSS_FINA_SELECTABLE_FIELDS:
                    label = FSS_FINA_FIELD_LABELS.get(field, field)
                    if field == "enpSaleAmt":
                        label += " 💰" # 매출액 강조
                    checked = st.checkbox(label, value=fss_available, key=f"fss_fina_{field}", disabled=not fss_available)
                    if checked and fss_available:
                        selected_fss_fina.append(field)

            has_any_selection = selected_nps or selected_nhis or selected_fss_corp or selected_fss_fina or selected_nts or selected_g2b
            if not has_any_selection:
                st.info("조회할 항목을 하나 이상 선택하세요.")
            else:
                st.markdown("<hr>", unsafe_allow_html=True)

                # ── Step 4: 조회 실행
                st.markdown('<div class="qx-section-label">STEP 4 · 데이터 매칭 실행</div>', unsafe_allow_html=True)

                # 매핑된 컬럼 수 확인
                mapped_count = sum(1 for c in [brn_col, name_col, ceo_col, addr_col] if c != "(선택 안 함)")

                similarity_threshold = 0

                st.markdown("<br>", unsafe_allow_html=True)
                total_rows = len(df)
                st.caption(f"총 {total_rows:,}건의 사업체를 매칭합니다.")

                if st.button("🔍 조회 시작", use_container_width=True, type="primary", key="btn_biz_start"):
                    progress = st.progress(0, text="조회 준비 중...")
                    status_area = st.empty()

                    try:
                        _match_name = cleaned_name_col if cleaned_name_col and cleaned_name_col in df.columns else (name_col if name_col != "(선택 안 함)" else "")

                        # ── 조회 대상 데이터 준비
                        query_items = []  # [(idx, brn, name, addr), ...]
                        for idx, row in df.iterrows():
                            brn = normalize_brn(row[brn_col])
                            search_name = str(row.get(_match_name, "")).strip() if _match_name else ""
                            if not search_name:
                                search_name = str(row.get(name_col, "")).strip() if name_col != "(선택 안 함)" else ""
                            
                            # 주소 데이터 획득
                            search_addr = str(row.get(addr_col, "")).strip() if addr_col != "(선택 안 함)" else ""
                            
                            # 법인등록번호 획득
                            search_crno = str(row.get(crno_col, "")).strip() if crno_col != "(선택 안 함)" else ""
                            
                            query_items.append((idx, brn, search_name, search_addr, search_crno))

                        # ── NPS 캐시 확인
                        cached_nps = st.session_state.get("biz_nps_cache", {})

                        # ── 1단계: 식별자 해결 (Identity Resolution)
                        from concurrent.futures import ThreadPoolExecutor, as_completed
                        import threading
                        
                        status_area.caption("🆔 기업 식별자 조회 중 (1/2단계)...")
                        resolved_identities = {} # {idx: {brn, crno, api_name, api_addr, match_score}}

                        def _is_masked(b):
                            b_str = str(b).strip()
                            if "*" in b_str or not b_str or b_str == "0000000000":
                                return True
                            return False

                        def _are_brns_consistent(b1, b2):
                            """마스킹을 고려하여 두 사업자번호가 일치하는지 확인 (불일치 시 False)"""
                            b1s = str(b1 or "").replace("-", "").replace(" ", "").zfill(10)
                            b2s = str(b2 or "").replace("-", "").replace(" ", "").zfill(10)
                            if not b1s or not b2s or (_is_masked(b1s) and _is_masked(b2s)):
                                return True
                            if len(b1s) != 10 or len(b2s) != 10: return True
                            for i in range(10):
                                if b1s[i] != "*" and b2s[i] != "*" and b1s[i] != b2s[i]:
                                    return False
                            return True

                        def _update_brn(current, new_candidate, res_id_ref, initial_input=None):
                            """마스킹되지 않은 유효한 10자리 번호를 우선시하여 업데이트하며 충돌 시 플래그 설정"""
                            c_clean = str(current or "").replace("-", "").replace(" ", "")
                            n_clean = str(new_candidate or "").replace("-", "").replace(" ", "")
                            i_clean = str(initial_input or "").replace("-", "").replace(" ", "") if initial_input else ""
                            
                            if not n_clean or n_clean == "0000000000" or n_clean == "None":
                                return c_clean

                            # [v12.0] BRN 불일치 확인
                            if c_clean and not _is_masked(c_clean) and not _is_masked(n_clean):
                                if not _are_brns_consistent(c_clean, n_clean):
                                    if res_id_ref is not None: res_id_ref["brn_mismatch"] = True
                                    return c_clean

                            # [v12.1] 추가: 입력된 사업자번호와도 대조
                            if i_clean and not _is_masked(i_clean) and not _is_masked(n_clean):
                                if not _are_brns_consistent(i_clean, n_clean):
                                    if res_id_ref is not None: res_id_ref["brn_mismatch"] = True
                                    return i_clean

                            if _is_masked(c_clean) and not _is_masked(n_clean) and len(n_clean) >= 10:
                                return n_clean
                            if not c_clean or c_clean == "0000000000":
                                return n_clean
                            if not _is_masked(c_clean) and len(c_clean) >= 10:
                                return c_clean
                            return n_clean

                        def _resolve_identity_task(idx, row_brn, row_name, row_addr, row_crno):
                            """단건에 대해 BRN 및 CRNO를 확정하는 로직"""
                            res_id = {
                                "brn": normalize_brn(row_brn) if not _is_masked(row_brn) else "", 
                                "input_brn": row_brn,
                                "crno": "", "api_name": "", 
                                "api_addr": "", "api_ceo": "",
                                "api_tel": "", "api_biz_type": "",
                                "match_score": 0.0,
                                "brn_mismatch": False,
                                "nps_name": "", "g2b_name": "", "fss_name": "", "dart_name": ""
                            }
                            def _log_debug(msg):
                                try:
                                    import os
                                    log_path = os.path.join(os.getcwd(), "api_debug.log")
                                    with open(log_path, "a", encoding="utf-8") as f:
                                        f.write(f"[{datetime.datetime.now()}] {msg}\n")
                                except: pass

                            # (Helpers already moved to parent scope)

                            # [v10.0] BRN-First Discovery Strategy 적용
                            # 모든 정보의 결합 목적은 "정확한(마스킹 없는) BRN 확보"

                            # 1) DART API를 통한 언마스킹 시도 (가장 강력한 언마스킹 수단)
                            dart_info = {}
                            if DART_API_KEY and row_name:
                                # [v16.0] 이름 매칭 강화: (주) 접두사 시도
                                dart_names_to_try = [row_name, f"(주){row_name}", f"{row_name}리서치", f"{row_name}코리아"]
                                for target_nm in dart_names_to_try:
                                    dart_info = get_dart_corp_info(target_nm, DART_API_KEY, brn=row_brn)
                                    if dart_info:
                                        api_brn = normalize_brn(dart_info.get("brn", ""))
                                        res_id["brn"] = _update_brn(res_id["brn"], api_brn, res_id, initial_input=row_brn)
                                        res_id["crno"] = res_id["crno"] or dart_info.get("crno", "")
                                        res_id["dart_name"] = dart_info.get("corp_name", "")
                                        res_id["api_name"] = res_id["api_name"] or res_id["dart_name"]
                                        res_id["api_addr"] = dart_info.get("addr", "")
                                        _log_debug(f"DART Unmasked/Enriched ({target_nm}): {res_id['brn']}, CRNO: {res_id['crno']}")
                                        break

                            # 2) G2B(나라장터)를 통한 보완 (BRN 기반 직접 조회)
                            current_brn = normalize_brn(res_id["brn"] or row_brn)
                            if api_service_key and current_brn and not _is_masked(current_brn):
                                g2b_info = get_g2b_corp_info(current_brn, api_service_key)
                                if g2b_info:
                                    # G2B info available
                                    res_id["g2b_name"] = g2b_info.get("corp_name", "") # [v13.5]
                                    res_id["api_name"] = res_id["api_name"] or res_id["g2b_name"]
                                    res_id["api_ceo"] = res_id["api_ceo"] or g2b_info.get("ceo_nm", "")
                                    res_id["api_tel"] = res_id["api_tel"] or g2b_info.get("telno", "")
                                    res_id["api_biz_type"] = res_id["api_biz_type"] or g2b_info.get("bizType", "")
                                    # [v13.5] G2B 주소가 있고 기존 DART 주소만 있는 경우, G2B 주소로 업데이트 (DART보단 실무 데이터인 G2B/NPS 선호)
                                    res_id["api_addr"] = g2b_info.get("addr", "") or res_id["api_addr"]
                                    res_id["corpSizeNm"] = g2b_info.get("corpSizeNm", "")
                                    res_id["main_product"] = g2b_info.get("main_product", "")
                                    res_id["restriction"] = g2b_info.get("restriction", "")
                                    _log_debug(f"G2B Enriched: {current_brn}")

                            # 3) FSS API 최종 확인 및 상세정보 확보
                            try:
                                target_brn_for_fss = res_id["brn"] if not _is_masked(res_id["brn"]) else row_brn
                                target_crno_for_fss = res_id["crno"] or row_crno
                                
                                # [v16.0] FSS 또한 이름 매칭 시도 (주식회사 접두어 등)
                                fss_names_to_try = [row_name, f"(주){row_name}", f"주식회사 {row_name}"]
                                fss_id_res = None
                                
                                for f_nm in fss_names_to_try:
                                    fss_id_res = search_corp_by_name(f_nm, api_service_key, brn=target_brn_for_fss, address=row_addr, crno=target_crno_for_fss)
                                    if fss_id_res and "_error" not in fss_id_res:
                                        break
                                
                                if fss_id_res and "_error" not in fss_id_res:
                                    api_brn = normalize_brn(fss_id_res.get("bzno", ""))
                                    api_name = str(fss_id_res.get("corpNm", "")).strip()
                                    api_addr = str(fss_id_res.get("enpAddr", "")).strip()
                                    
                                    # [v11.0] 상호명 70% 이상 + 시도 일치 조건 (BRN 부재 시)
                                    is_input_brn_missing = not row_brn or _is_masked(row_brn)
                                    if is_input_brn_missing:
                                        name_sim = text_similarity(row_name, api_name)
                                        u_sido, _, _ = split_address(row_addr)
                                        a_sido, _, _ = split_address(api_addr)
                                        
                                        # 조건을 만족할 때만 BRN으로 채택
                                        if name_sim >= 0.7 and u_sido == a_sido and u_sido:
                                            res_id["brn"] = _update_brn(res_id["brn"], api_brn, res_id, initial_input=row_brn)
                                            res_id["match_score"] = name_sim * 100
                                        else:
                                            # 조건 미충족 시 FSS 결과 무시 (BRN 확보 실패)
                                            api_brn = ""
                                    
                                    if api_brn:
                                        if not is_input_brn_missing:
                                            res_id["brn"] = _update_brn(res_id["brn"], api_brn, res_id, initial_input=row_brn)
                                            
                                        res_id["crno"] = res_id["crno"] or str(fss_id_res.get("crno", "")).strip()
                                        res_id["fss_name"] = api_name # [v13.5]
                                        res_id["api_name"] = res_id["api_name"] or res_id["fss_name"]
                                        # [v13.5] 금융위 주소는 공시 정보이므로 우선순위 높임 (DART와 유사하나 개별 재무 공시 기준)
                                        res_id["api_addr"] = api_addr or res_id["api_addr"]
                                        res_id["api_ceo"] = res_id["api_ceo"] or str(fss_id_res.get("ceoNm", "")).strip()
                                        
                                        if not is_input_brn_missing:
                                            is_brn_match = not _is_masked(res_id["brn"]) and (res_id["brn"] == normalize_brn(row_brn) or len(res_id["brn"]) == 10)
                                            res_id["match_score"] = 100.0 if (is_brn_match) else 80.0
                                else:
                                    res_id["fss_error"] = fss_id_res.get("_error") if fss_id_res else "결과없음"
                            except Exception as e:
                                res_id["fss_error"] = str(e)

                            # 4) NPS API 시도 (마지막 수단 또는 추가 확인)
                            try:
                                current_brn = normalize_brn(res_id["brn"] or row_brn)
                                nps_id_res = search_and_match_nps(row_name, current_brn, api_service_key, address=row_addr)
                                if nps_id_res and "_error" not in nps_id_res:
                                    api_brn = normalize_brn(nps_id_res.get("bzowrRgstNo", ""))
                                    res_id["brn"] = _update_brn(res_id["brn"], api_brn, res_id, initial_input=row_brn)
                                    res_id["nps_name"] = str(nps_id_res.get("wkplNm", "")).strip() # [v13.5]
                                    res_id["api_name"] = res_id["api_name"] or res_id["nps_name"]
                                    # [v13.5] 국민연금 주소는 매달 갱신되므로 가장 최신일 확률이 높음 → 최우선 적용
                                    nps_addr = str(nps_id_res.get("wkplRoadNmAddr", "") or nps_id_res.get("wkplRoadNmDtlAddr", "")).strip()
                                    if nps_addr:
                                        res_id["api_addr"] = nps_addr
                                    
                                    if res_id["match_score"] < 70:
                                        res_id["match_score"] = 70.0
                            except Exception as e:
                                res_id["nps_error"] = str(e)
                            
                            # [v15.3] G2B industry 필드 이중 체크 (api_biz_type이 비어있는 경우)
                            if not res_id.get("api_biz_type") and res_id.get("g2b_name"):
                                # G2B가 이미 1단계에서 실행되었으므로 g2b_info는 로컬에 없지만 res_id에는 데이터가 있어야 함
                                # 만약 1단계에서 g2b_info["bizType"]을 못 가져왔다면 여기서 다시 한번 로깅하거나 보완 가능
                                pass

                            return idx, res_id

                        # 식별자 해결 실행
                        with ThreadPoolExecutor(max_workers=5) as executor:
                            futures = {
                                executor.submit(_resolve_identity_task, idx, brn, name, addr, crno): idx
                                for idx, brn, name, addr, crno in query_items
                            }
                            for i, future in enumerate(as_completed(futures), 1):
                                idx_res, data = future.result()
                                resolved_identities[idx_res] = data
                                progress.progress(min(i / total_rows * 0.2, 0.19), text=f"기업 식별 중... ({i}/{total_rows})")

                        # ── 3단계: 국세청 (NTS) 상태 조회
                        if selected_nts:
                            status_area.caption("⚖️ 국세청 사업자 상태 조회 중...")
                            # 유효한(마스킹되지 않은) BRN 목록 추출
                            nts_query_brns = [ident.get("brn") for ident in resolved_identities.values() if ident.get("brn") and not "*" in ident.get("brn")]
                            if nts_query_brns:
                                nts_results = get_nts_business_status(nts_query_brns, api_service_key)
                                # 결과 매핑
                                for idx, ident in resolved_identities.items():
                                    b_no = ident.get("brn")
                                    if b_no in nts_results:
                                        status_val = nts_results[b_no].get("status", "")
                                        # [v15.5] 국세청 상태 아이콘 적용
                                        if "계속" in status_val: status_val = f"🟢 {status_val}"
                                        elif "폐업" in status_val: status_val = f"🔴 {status_val}"
                                        elif "휴업" in status_val: status_val = f"🟡 {status_val}"
                                        
                                        resolved_identities[idx].update({
                                            "nts_status": status_val,
                                            "nts_tax_type": nts_results[b_no].get("tax_type", ""),
                                            "nts_end_dt": nts_results[b_no].get("end_dt", "")
                                        })
                            progress.progress(0.4, text="국세청 상태 조회 완료")

                        # ── 4단계: 개별 API 데이터 수집 (NHIS, NPS, FSS)
                        nhis_df = pd.DataFrame()
                        if selected_nhis:
                            # 캐시가 있고, 선택된 UDDI와 동일한 경우만 재사용
                            if "biz_nhis_dataset" in st.session_state and st.session_state.get("biz_nhis_uddi_cache") == selected_nhis_uddi:
                                nhis_df = st.session_state["biz_nhis_dataset"]
                                status_area.caption(f"✅ 건강보험 데이터셋 캐시 사용 ({len(nhis_df):,}건)")
                            else:
                                def nhis_progress(page, total, msg):
                                    progress.progress(min(0.2 + page/total * 0.1, 0.29), text=msg)
                                    status_area.caption(msg)
                                    status_area.caption(msg)
                                nhis_df = download_nhis_dataset(api_service_key, uddi=selected_nhis_uddi, progress_callback=nhis_progress)
                                st.session_state["biz_nhis_dataset"] = nhis_df
                                st.session_state["biz_nhis_uddi_cache"] = selected_nhis_uddi
                        
                        # ── 2단계: 상세 데이터 수집 (Data Collection)
                        status_area.caption("📊 확정된 ID로 행정 데이터 수집 중 (2/2단계)...")
                        nps_results = {}
                        fss_results = {}
                        cached_nps = st.session_state.get("biz_nps_cache", {})

                        def _fetch_details_task(idx, resolved_id):
                            res_brn = resolved_id.get("brn") or resolved_id.get("input_brn")
                            res_crno = resolved_id.get("crno")
                            res_name = resolved_id.get("api_name") or query_items[idx][2]
                            
                            details = {"nps": {}, "fss": {"corp": {}, "fina": {}}}
                            
                            # NPS 수집
                            if selected_nps:
                                # 캐시 체크
                                if res_brn and res_brn in cached_nps:
                                    details["nps"] = cached_nps[res_brn]
                                else:
                                    res_sido = resolved_id.get("sido", "")
                                    details["nps"] = search_and_match_nps(res_name, res_brn, api_service_key, address=resolved_id.get("api_addr", ""), input_sido=res_sido)
                                    
                                    # [v15.2] 1차 검색 실패 시 원본 입력명으로 재시도 (api_name != 사용자 입력명인 경우)
                                    original_input_name = query_items[idx][2]
                                    if (not details["nps"] or "_error" in details["nps"]) and original_input_name and original_input_name != res_name:
                                        details["nps"] = search_and_match_nps(original_input_name, res_brn, api_service_key, address=resolved_id.get("api_addr", ""), input_sido=res_sido)
                                    
                                    if res_brn and details["nps"] and "_error" not in details["nps"]:
                                        cached_nps[res_brn] = details["nps"]
                                
                                # [v15.1] NPS를 찾은 경우 등록명을 resolved_identities에 역전파
                                if details["nps"] and "_error" not in details["nps"]:
                                    nps_wkpl_nm = details["nps"].get("wkplNm", "")
                                    if nps_wkpl_nm and not resolved_id.get("nps_name"):
                                        resolved_id["nps_name"] = nps_wkpl_nm
                                        resolved_id["api_name"] = resolved_id.get("api_name") or nps_wkpl_nm
                                        nps_addr = details["nps"].get("wkplRoadNmAddr") or details["nps"].get("wkplRoadNmDtlAddr") or ""
                                        if nps_addr and not resolved_id.get("api_addr"):
                                            resolved_id["api_addr"] = nps_addr
                            
                            # FSS 수집
                            if selected_fss_corp or selected_fss_fina:
                                biz_year = str(datetime.datetime.now().year - 1)
                                if res_crno:
                                    fina = search_financial_by_crno(res_crno, biz_year, api_service_key)
                                    corp = search_corp_by_name(res_name, api_service_key, brn=res_brn) # 기본정보 재조회
                                    details["fss"] = {"corp": corp, "fina": fina}
                                elif res_brn:
                                    corp = search_corp_by_name(res_name, api_service_key, brn=res_brn)
                                    details["fss"] = {"corp": corp, "fina": {"_error": "CRNO 없음"}}
                            return idx, details

                        with ThreadPoolExecutor(max_workers=5) as executor:
                            futures = {
                                executor.submit(_fetch_details_task, idx, resolved_identities[idx]): idx
                                for idx in resolved_identities
                            }
                            for i, future in enumerate(as_completed(futures), 1):
                                idx_res, data = future.result()
                                nps_results[idx_res] = data["nps"]
                                fss_results[idx_res] = data["fss"]
                                progress.progress(min(0.3 + i / total_rows * 0.5, 0.79), text=f"데이터 수집 중... ({i}/{total_rows})")

                        st.session_state["biz_nps_cache"] = cached_nps

                        def _extract_nps_field(results_dict, idx, field):
                            res = results_dict.get(idx)
                            if not res or "_error" in res: return "-"
                            val = res.get(field, "")
                            if field == "crrmmNtcAmt" and val:
                                try: return f"{int(float(val)):,}"
                                except: pass
                            return str(val) if val else "-"

                        def _extract_nhis_field(idx, field):
                            res_id = resolved_identities.get(idx, {})
                            brn = res_id.get("brn")
                            
                            # 1) BRN으로 우선 검색
                            if brn:
                                nrow = nhis_lookup.get(brn)
                                # [v15.3] 마스킹된 BRN인 경우 대조 처리하여 nhis_lookup 검색
                                if not nrow and _is_masked(brn):
                                    for target_brn, target_row in nhis_lookup.items():
                                        if _are_brns_consistent(brn, target_brn):
                                            nrow = target_row
                                            break
                                            
                                if nrow:
                                    val = nrow.get(field, "")
                                    return str(val) if val else "-"
                            
                            # 2) 상호명+주소로 대체 검색 (BRN 없거나 매칭 안된 경우)
                            row_name = query_items[idx][2] # 정제된 이름
                            row_addr = query_items[idx][3]
                            
                            matches = nhis_name_lookup.get(row_name, [])
                            if matches:
                                best_match = None
                                max_sim = 0
                                for m in matches:
                                    sim = text_similarity(row_addr, m.get("주소", ""))
                                    if sim > max_sim:
                                        max_sim = sim
                                        best_match = m
                                
                                # 주소가 어느 정도 유사하거나, 주소 정보가 둘 다 없는 경우 매칭
                                if best_match and (max_sim > 0.4 or (not row_addr and not best_match.get("주소"))):
                                    val = best_match.get(field, "")
                                    return str(val) if val else "-"
                                    
                            return "-"

                        # 4) 결과 병합 및 유사도 계산
                        progress.progress(0.8, text="결과 병합 중...")
                        result_df = df.copy()
                        
                        # NHIS 룩업 사전 구축 (속도 최적화)
                        nhis_lookup = {}
                        nhis_name_lookup = {} # 상호명 기반 루치업용
                        if not nhis_df.empty:
                            if "_brn" in nhis_df.columns:
                                nhis_lookup = nhis_df.set_index("_brn").to_dict('index')
                            
                            # 상호명 기반 룩업 사전 (정제된 상호명 사용)
                            if "사업장명" in nhis_df.columns:
                                for _, n_row in nhis_df.iterrows():
                                    nm = clean_company_name(n_row["사업장명"])
                                    if nm:
                                        if nm not in nhis_name_lookup:
                                            nhis_name_lookup[nm] = []
                                        nhis_name_lookup[nm].append(n_row.to_dict())

                        # 식별된 마스터 ID를 기반으로 신뢰도 및 기본 정보 세팅
                        result_df["회사명 정제"] = ""
                        result_df["주소 정제"] = ""
                        result_df["시도"] = ""
                        result_df["시군"] = ""
                        result_df["[입력] 주소(정제)"] = ""
                        result_df["[입력] 상호(정제)"] = ""
                        result_df["[공공데이터] 사업자등록번호"] = ""
                        result_df["[공공데이터] 법인등록번호"] = ""
                        result_df["[공공데이터] 대표자명"] = ""
                        result_df["[공공데이터] 주소(도로명)"] = ""
                        # [v13.5] 각 기관별 등록명 컬럼 추가
                        result_df["[국민연금] 등록명"] = ""
                        result_df["[건강보험] 등록명"] = ""
                        result_df["[나라장터] 등록명"] = ""
                        result_df["[금융위] 등록명"] = ""
                        result_df["[전자공시] 등록명"] = ""
                        result_df["[행정] 시도"] = ""
                        result_df["[행정] 시군구"] = ""
                        result_df["유사도(%)"] = 0.0

                        rows_with_data = 0
                        for idx, row in result_df.iterrows():
                            # 원본 주소 처리
                            uaddr = str(row.get(addr_col, "")).strip() if addr_col != "(선택 안 함)" else ""
                            if uaddr:
                                c_uaddr = clean_address(uaddr)
                                u_sido, u_sgg, _ = split_address(c_uaddr)
                                # [UI 개선] 사용자 요청 명칭 반영
                                result_df.at[idx, "회사명 정제"] = row_name
                                result_df.at[idx, "주소 정제"] = c_uaddr
                                result_df.at[idx, "시도"] = u_sido
                                result_df.at[idx, "시군"] = u_sgg
                                
                                # 히든용 원본 유지
                                result_df.at[idx, "[입력] 주소(정제)"] = c_uaddr
                                result_df.at[idx, "[입력] 상호(정제)"] = row_name

                            res_id = resolved_identities.get(idx, {})
                            master_brn = res_id.get("brn", "")
                            master_crno = res_id.get("crno", "")
                            result_df.at[idx, "[공공데이터] 사업자등록번호"] = master_brn
                            result_df.at[idx, "[공공데이터] 법인등록번호"] = master_crno
                            result_df.at[idx, "[공공데이터] 대표자명"] = res_id.get("api_ceo", "")
                            
                            # 행정데이터 주소 및 분리
                            api_addr = res_id.get("api_addr", "")
                            result_df.at[idx, "[공공데이터] 주소(도로명)"] = api_addr
                            if api_addr:
                                a_sido, a_sgg, _ = split_address(api_addr)
                                result_df.at[idx, "[행정] 시도"] = a_sido
                                result_df.at[idx, "[행정] 시군구"] = a_sgg

                            # [v13.5] 각 기관별 등록명 채우기
                            result_df.at[idx, "[국민연금] 등록명"] = res_id.get("nps_name", "")
                            result_df.at[idx, "[나라장터] 등록명"] = res_id.get("g2b_name", "")
                            result_df.at[idx, "[금융위] 등록명"] = res_id.get("fss_name", "")
                            result_df.at[idx, "[전자공시] 등록명"] = res_id.get("dart_name", "")

                            # result_df.at[idx, "[조달청] 등록업종"] = res_id.get("api_biz_type", "")
                            # result_df.at[idx, "[조달청] 전화번호"] = res_id.get("api_tel", "")
                            
                            has_any_data = False

                            # G2B 결과 추출 (v13.0)
                            if selected_g2b:
                                for field in selected_g2b:
                                    label = G2B_FIELD_LABELS.get(field, field).replace(" 📞", "")
                                    col_name = f"[나라장터] {label}"
                                    val = ""
                                    if field == "bizType": val = res_id.get("api_biz_type", "")
                                    elif field == "telno": val = res_id.get("api_tel", "")
                                    elif field == "corpSizeNm": val = res_id.get("corpSizeNm", "")
                                    elif field == "main_product": val = res_id.get("main_product", "")
                                    elif field == "restriction": val = res_id.get("restriction", "")
                                    
                                    result_df.at[idx, col_name] = str(val) if val else "-"
                                    if val: has_any_data = True

                            # NPS 결과 추출
                            if selected_nps:
                                nps_row = nps_results.get(idx, {})
                                
                                # 실질적인 데이터가 있는지 확인 (단순 에러나 빈 리스트 제외)
                                has_nps_data = False
                                if nps_row and "_error" not in nps_row:
                                    # 가입자수나 당월고지금액 등 핵심 필드가 있는지 확인
                                    if any(nps_row.get(f) for f in ["jnngpCnt", "crrmmNtcAmt", "bzowrRgstNo"]):
                                        has_nps_data = True
                                        has_any_data = True
                                        
                                for field in selected_nps:
                                    label = NPS_FIELD_LABELS.get(field, field)
                                    col_name = f"[국민연금] {label}"
                                    val = _extract_nps_field(nps_results, idx, field)
                                    result_df.at[idx, col_name] = val
                                
                                if "avgBasSalary" in selected_nps:
                                    result_df.at[idx, "[국민연금] 추정 평균 기준소득월액"] = estimate_avg_salary(nps_row)

                                # [v12.7] 매칭 신뢰도 필드 추가
                                result_df.at[idx, "[국민연금] 매칭도"] = nps_row.get("_match_score", 0.0) if has_nps_data else 0.0

                            if selected_nhis:
                                for field in selected_nhis:
                                    val = _extract_nhis_field(idx, field)
                                    if val and val != "-":
                                        has_any_data = True
                                    result_df.at[idx, f"[건강보험] {NHIS_FIELD_LABELS.get(field, field)}"] = val
                                    
                                    # [v13.5] 건강보험 등록명 별도 추출 (룩업 결과 활용)
                                    if field == selected_nhis[0]: # 한 번만 수행
                                        res_nhis_name = _extract_nhis_field(idx, "사업장명")
                                        if res_nhis_name not in ["조회불가", "미조회", "해당없음"]:
                                            result_df.at[idx, "[건강보험] 등록명"] = res_nhis_name

                            # NTS 결과 추출
                            if selected_nts:
                                nts_labels = {"status": "사업자상태", "tax_type": "과세유형", "end_dt": "폐업일자"}
                                for field in selected_nts:
                                    label = nts_labels.get(field, field)
                                    col_name = f"[국세청] {label}"
                                    val = res_id.get(f"nts_{field}", "")
                                    if not val: val = "-"
                                    result_df.at[idx, col_name] = val
                                    if val and val != "-":
                                        has_any_data = True

                            # FSS 결과 추출
                            fss_data = fss_results.get(idx, {})
                            if selected_fss_corp:
                                corp = fss_data.get("corp", {})
                                if corp and "_error" not in corp:
                                    # 하나라도 값이 있는지 확인
                                    if any(corp.get(f) for f in selected_fss_corp):
                                        has_any_data = True
                                for field in selected_fss_corp:
                                    label = FSS_CORP_FIELD_LABELS.get(field, field)
                                    col_name = f"[기업정보] {label}"
                                    val = corp.get(field, "") if "_error" not in corp else "-"
                                    if not val: val = "-"
                                    result_df.at[idx, col_name] = str(val)

                            if selected_fss_fina:
                                fina = fss_data.get("fina", {})
                                if fina and "_error" not in fina:
                                    # 하나라도 값이 있는지 확인
                                    if any(fina.get(f) for f in selected_fss_fina):
                                        has_any_data = True
                                for field in selected_fss_fina:
                                    label = FSS_FINA_FIELD_LABELS.get(field, field)
                                    col_name = f"[재무정보] {label}"
                                    val = fina.get(field, "")
                                    if not val: val = "-"
                                    elif str(val).replace("-", "").replace(".", "").isdigit():
                                        try: val = f"{int(float(val)):,}"
                                        except: pass
                                    result_df.at[idx, col_name] = str(val)
                            
                            if has_any_data:
                                rows_with_data += 1

                            # ── 유사도 계산 (Identity Resolution 결과 활용)
                            scores, weights = [], []
                            # 1. 식별 성공 여부 (30%) - API에서 확인된 경우만 만점
                            scores.append(1.0 if res_id.get("brn") else 0.0)
                            weights.append(30)

                            # 2. 회사명 유사도 (40%)
                            uname = str(row.get(name_col, "")).strip() if name_col != "(선택 안 함)" else ""
                            api_name = res_id.get("api_name", "")
                            if uname and api_name:
                                from utils.matcher import text_similarity
                                scores.append(text_similarity(uname, api_name))
                                weights.append(40)
                            
                            # 3. 주소 유사도 (30%)
                            uaddr = str(row.get(addr_col, "")).strip() if addr_col != "(선택 안 함)" else ""
                            api_addr = res_id.get("api_addr", "")
                            if uaddr and api_addr:
                                from utils.matcher import text_similarity
                                scores.append(text_similarity(uaddr, api_addr))
                                weights.append(30)

                            final_sim = round(sum(s*w for s, w in zip(scores, weights)) / sum(weights) * 100, 1) if weights else 0.0
                            result_df.at[idx, "유사도(%)"] = final_sim

                        # ── [v8.5] 기업체 기준(Enterprise) 집계 로직 적용
                        if agg_mode == "기업체 기준 (Enterprise)":
                            status_area.caption("🏢 기업체 단위로 데이터 집계 중...")
                            # 집계 기준 키 생성: 법인번호(우선) -> 사업자번호 -> 회사명
                            def _get_agg_key(r):
                                crno = str(r.get("[공공데이터] 법인등록번호", "")).strip()
                                if crno and crno != "해당없음": return crno
                                brn = str(r.get("[공공데이터] 사업자등록번호", "")).strip()
                                if brn and brn != "해당없음" and "*" not in brn: return brn
                                return str(r.get(name_col, "")).strip()

                            result_df["_agg_key"] = result_df.apply(_get_agg_key, axis=1)
                            
                            # 집계 규칙 정의
                            agg_rules = {}
                            num_cols = []
                            for c in result_df.columns:
                                if c in [brn_col, name_col, ceo_col, addr_col, "_agg_key", "유사도(%)"]: continue
                                # 수치형 데이터(가입자수, 금액 등)는 합산
                                if any(keyword in c for keyword in ["가입자수", "인원", "금액", "매출액", "영업이익", "순이익", "자본금", "자산", "부채", "평균소득"]):
                                    num_cols.append(c)
                                    agg_rules[c] = "sum"
                                else:
                                    # 텍스트 데이터는 첫 번째 값 유지
                                    agg_rules[c] = "first"
                            
                            # 기본 컬럼 유지 규칙
                            agg_rules[name_col] = "first"
                            agg_rules[addr_col] = "first"
                            agg_rules["유사도(%)"] = "max" # 유사도는 최대값 유지

                            # 집계 실행
                            # 수치형 컬럼 전처리 (콤마 제거 등)
                            for c in num_cols:
                                result_df[c] = result_df[c].apply(lambda x: str(x).replace(",", ""))
                                result_df[c] = pd.to_numeric(result_df[c], errors='coerce').fillna(0)

                            result_df = result_df.groupby("_agg_key").agg(agg_rules).reset_index(drop=True)
                            
                            # 수치형 다시 포맷팅
                            for c in num_cols:
                                if "인원" in c or "수" in c:
                                    result_df[c] = result_df[c].apply(lambda x: f"{int(x):,}")
                                else:
                                    result_df[c] = result_df[c].apply(lambda x: f"{int(x):,}" if x > 0 else "0")
                            
                            # 집계 후 건수 업데이트
                            rows_with_data = len(result_df) 
                            st.info(f"💡 {total_rows}개의 지점 데이터를 법인/기업 단위로 합산하여 {rows_with_data}개의 요약 행으로 정리했습니다.")

                        # 유사도 필터링 제거 (모든 결과 표시)
                        
                        # 통계 계산 (실제로 데이터를 가져온 행 기준)
                        # 통계 계산 (실제로 유의미한 데이터를 가져온 행 기준)
                        # 통계 계산
                        total_matched = rows_with_data
                        nps_matched = sum(1 for v in nps_results.values() if v and "_error" not in v)
                        
                        nhis_matched = 0
                        for idx in resolved_identities:
                            val = _extract_nhis_field(idx, selected_nhis[0]) if selected_nhis else "미조회"
                            if val not in ["조회불가", "미조회", "해당없음"]:
                                nhis_matched += 1
                        
                        fss_corp_matched = sum(1 for v in fss_results.values() if "_error" not in v.get("corp", {}))
                        fss_fina_matched = sum(1 for v in fss_results.values() if "_error" not in v.get("fina", {}))
                        
                        # 최종 성공 건수는 이 중 하나라도 데이터를 가진 행의 수
                        total_matched = rows_with_data

                        progress.progress(1.0, text="✅ 조회 완료!")
                        status_area.empty()

                        total_matched = max(nps_matched, nhis_matched, fss_corp_matched)
                        # 유사도 분포 계산 (10% 단위)
                        sim_dist = [0] * 11 # 0, 10, 20, ..., 100
                        for s in result_df["유사도(%)"]:
                            idx_bin = int(s // 10)
                            if idx_bin > 10: idx_bin = 10
                            sim_dist[idx_bin] += 1
                        
                        dist_labels = ["0-10%", "11-20%", "21-30%", "31-40%", "41-50%", "51-60%", "61-70%", "71-80%", "81-90%", "91-100%", "보정매칭"]
                        # 보정매칭은 사실상 90-100% 구간에 포함되나, 필요시 별도 표기 가능. 일단 10개 구간으로 정리
                        dist_bins = {}
                        for i in range(10):
                            start = i * 10
                            end = (i+1) * 10
                            count = sum(1 for s in result_df["유사도(%)"] if start < s <= end)
                            if i == 0: # 0점 포함
                                count = sum(1 for s in result_df["유사도(%)"] if start <= s <= end)
                            dist_bins[f"{start+1 if i>0 else 0}-{end}%"] = count

                        stats = {
                            "total": total_rows,
                            "matched": total_matched,
                            "unmatched": total_rows - total_matched,
                            "result_rows": len(result_df),
                            "nps_searched": len(nps_results),
                            "nps_matched": nps_matched,
                            "nhis_size": len(nhis_df),
                            "nhis_matched": nhis_matched,
                            "fss_corp_matched": fss_corp_matched,
                            "fss_fina_matched": fss_fina_matched,
                            "sim_dist": dist_bins
                        }

                        # [v16.0] 사용자 요청 컬럼 순서 재배치 및 가시성 조정
                        # 요청: 유사도, 회사명 정제, 기관 통합 등록명, 주소 정제, 시도, 시군
                        priority_cols = ["유사도(%)", "회사명 정제", "[행정] 기관 통합 등록명", "주소 정제", "시도", "시군"]
                        
                        # 1. 기관명 통합 (DART > G2B > NPS > NHIS/FSS)
                        for idx, ident in resolved_identities.items():
                            dart_nm = ident.get("dart_name")
                            g2b_nm = ident.get("g2b_name")
                            nps_nm = ident.get("nps_name")
                            fss_nm = ident.get("fss_name")
                            
                            unified_nm = dart_nm or g2b_nm or fss_nm or nps_nm or ""
                            if unified_nm:
                                result_df.at[idx, "[행정] 기관 통합 등록명"] = unified_nm
                        
                        # 2. 숨길 컬럼 정의 (원본 + 원본보조필드 + 특정 NPS 필드)
                        raw_input_cols = [c for c in [brn_col, name_col, ceo_col, addr_col] if c and c != "(선택 안 함)"]
                        redundant_cols = [
                            "[입력] 시도", "[입력] 시군구", "[전자공시] 등록명", "[건강보험] 등록명", 
                            "[입력] 주소(정제)", "[입력] 상호(정제)",
                            "[국민연금] 사업자등록번호(6자리)", "[국민연금] 도로명주소"
                        ]
                        hide_list = raw_input_cols + redundant_cols
                        
                        # 3. 나머지 컬럼 필터링 (원본 제외)
                        other_cols = [c for c in cols if c not in priority_cols and c not in hide_list and c != "_agg_key"]
                        
                        # 4. 최종 컬럼 구성 및 이름 변경
                        final_col_order = [c for c in priority_cols if c in result_df.columns] + other_cols
                        result_df = result_df[final_col_order]
                        
                        # 컬럼명 최종 매핑 (사용자 요청 명칭으로 간결화)
                        rename_map = {
                            "[행정] 기관 통합 등록명": "기관 통합 등록명"
                        }
                        result_df = result_df.rename(columns=rename_map)
                        
                        st.session_state["biz_crawl_result"] = result_df
                        st.session_state["biz_crawl_stats"] = stats
                        st.session_state["biz_crawl_selected_nps"] = selected_nps
                        st.session_state["biz_crawl_selected_nhis"] = selected_nhis
                        st.rerun()

                    except PermissionError as e:
                        progress.empty()
                        status_area.empty()
                        st.error(
                            "🔑 **API 키의 유효기간이 만료되었거나 인증에 실패했습니다.**\n\n"
                            "공공데이터포털(data.go.kr)에서 API 키를 갱신한 후, "
                            "Streamlit Cloud Secrets의 `DATA_GO_KR_KEY` 값을 업데이트해주세요.\n\n"
                            f"오류 상세: `{str(e)}`"
                        )
                        st.stop()
                    except Exception as e:
                        progress.empty()
                        status_area.empty()
                        st.error(f"데이터 조회 중 오류가 발생했습니다: {str(e)}")

    # ══════════════════════════════════════════════════════════════
    # ── Step 5: 결과 표시 (세션 상태 기반 — 중첩 외부에서 독립 렌더링)
    # ══════════════════════════════════════════════════════════════
    if "biz_crawl_result" in st.session_state and st.session_state["biz_crawl_result"] is not None:
        result_df = st.session_state["biz_crawl_result"]
        stats = st.session_state.get("biz_crawl_stats", {})

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="qx-section-label">STEP 5 · 조회 결과</div>', unsafe_allow_html=True)

        # 요약 통계 카드
        total = stats.get("total", len(result_df))
        matched = stats.get("matched", 0)
        unmatched = stats.get("unmatched", 0)
        success_rate = (matched / total * 100) if total > 0 else 0

        if matched == 0 and total > 0:
            st.warning(
                "⚠️ **데이터가 전혀 조회되지 않았습니다. 원인을 확인해 주세요.**\n\n"
                "- **조회 실패 유형:**\n"
                "  1. **API 매칭 실패:** 입력하신 '사업자번호'나 '회사명'이 공공데이터 DB와 일치하지 않을 수 있습니다.\n"
                "  2. **API 엔진 오류:** 최근 NPS API V2 전환 등에 따라 파라미터 규격이 변경된 경우 조회가 불가할 수 있습니다. (현재 최신 규격 반영 완료)\n"
                "  3. **인증 오류:** 서비스키가 만료되었거나 'Encoding' 키를 사용한 경우 발생합니다. 'Decoding' 키를 권장합니다.\n"
                "  4. **승인 대기:** API 신청 직후라면 실제 사용까지 **1~2시간(최대 24시간)**이 소요될 수 있습니다."
            )

        col_s1, col_s2, col_s3, col_s4 = st.columns(4)
        with col_s1:
            st.markdown(f"""
<div class="qx-card" style="text-align:center;">
    <div style="font-size:2rem; font-weight:700; color:#0F6CBD;">{total:,}</div>
    <div style="font-size:0.8rem; color:#8B96A9;">전체 건수</div>
</div>""", unsafe_allow_html=True)
        with col_s2:
            st.markdown(f"""
<div class="qx-card" style="text-align:center;">
    <div style="font-size:2rem; font-weight:700; color:#059669;">{matched:,}</div>
    <div style="font-size:0.8rem; color:#8B96A9;">매칭 성공</div>
</div>""", unsafe_allow_html=True)
        with col_s3:
            st.markdown(f"""
<div class="qx-card" style="text-align:center;">
    <div style="font-size:2rem; font-weight:700; color:#DC2626;">{unmatched:,}</div>
    <div style="font-size:0.8rem; color:#8B96A9;">매칭 실패</div>
</div>""", unsafe_allow_html=True)
        with col_s4:
            st.markdown(f"""
<div class="qx-card" style="text-align:center;">
    <div style="font-size:2rem; font-weight:700; color:#7C3AED;">{success_rate:.1f}%</div>
    <div style="font-size:0.8rem; color:#8B96A9;">성공률</div>
</div>""", unsafe_allow_html=True)

        # API 상태 정보 (디버깅/안내용)
        with st.expander("🌐 API 서비스 연결 상태 확인", expanded=(matched == 0)):
            api_cols = st.columns(3)
            s_nps = stats.get("nps_matched", 0)
            s_nhis = stats.get("nhis_matched", 0)
            s_fss_corp = stats.get("fss_corp_matched", 0)
            s_fss_fina = stats.get("fss_fina_matched", 0)

            with api_cols[0]:
                nps_ok = "✅ 정상" if s_nps > 0 else "⚠️ 미매칭/인증확인필요"
                st.write(f"**국민연금:** {nps_ok}")
            with api_cols[1]:
                nhis_ok = "✅ 정상" if s_nhis > 0 else "⚠️ 미매칭/인증확인필요"
                st.write(f"**건강보험:** {nhis_ok}")
            with api_cols[2]:
                fss_ok = "✅ 정상" if (s_fss_corp > 0 or s_fss_fina > 0) else "⚠️ 미매칭/인증확인필요"
                st.write(f"**금융위:** {fss_ok}")
            st.caption("※ '정상'은 유의미한 데이터가 1건 이상 추출되었음을 의미합니다.")

        # 유사도 분포 (10% 단위) 표시
        sim_dist = stats.get("sim_dist", {})
        if sim_dist:
            st.markdown("#### 📊 유사도 분포 (10% 단위)")
            cols_dist = st.columns(len(sim_dist))
            for i, (label, count) in enumerate(sim_dist.items()):
                with cols_dist[i]:
                    color = "#0F6CBD" if count > 0 else "#8B96A9"
                    st.markdown(f"""
                    <div style="text-align:center; padding:0.5rem; background:#F8FAFC; border-radius:6px; border:1px solid #E2E8F0;">
                        <div style="font-size:0.7rem; color:#64748B; margin-bottom:0.2rem;">{label}</div>
                        <div style="font-size:1rem; font-weight:700; color:{color};">{count}</div>
                    </div>
                    """, unsafe_allow_html=True)

        # 데이터셋 정보
        nps_size = stats.get("nps_size", 0)
        nhis_size = stats.get("nhis_size", 0)
        if nps_size or nhis_size:
            info_parts = []
            if nps_size:
                info_parts.append(f"국민연금 {nps_size:,}건")
            if nhis_size:
                info_parts.append(f"건강보험 {nhis_size:,}건")
            fss_corp_m = stats.get("fss_corp_matched", 0)
            fss_fina_m = stats.get("fss_fina_matched", 0)
            if fss_corp_m:
                info_parts.append(f"기업정보 {fss_corp_m:,}건 매칭")
            if fss_fina_m:
                info_parts.append(f"재무정보 {fss_fina_m:,}건 매칭")
            st.caption(f"📦 참조 데이터셋: {' · '.join(info_parts)}")

        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("💡 데이터 일치 여부 판단 가이드 (나비케어 등 주소 불일치 시)", expanded=True):
            st.info(
                "**주소가 달라도 같은 기업인지 판단하는 방법:**\n\n"
                "1. **대표자명 대조:** 상호명과 대표자명이 일치한다면 주소지가 지점이거나 이전된 경우일 가능성이 높습니다.\n"
                "2. **법인등록번호(13자리) 확인:** 사업자번호는 지점마다 다를 수 있지만, **법인번호**는 법인 전체가 동일합니다. 재무정보의 법인번호와 타 기관 자료를 대조해 보세요.\n"
                "3. **설립일/업종 확인:** [기업정보] 섹션의 **설립일**이나 **업종명**이 조사 대상 기업의 정보와 일치하는지 확인해 보세요.\n\n"
                "**[국세청 사업자상태 범례]:**  🟢 계속사업  |  🔴 폐업  |  🟡 휴업 / 정보없음\n\n"
                "※ 국민연금은 '사업장' 기준이므로 본사와 지점의 번호가 다를 수 있으며, 금감원 재무제표는 '법인' 전체 실적을 보여줍니다."
            )

        # 결과 DataFrame 표시
        column_config = {
            "유사도(%)": st.column_config.ProgressColumn(
                "유사도(%)",
                help="업로드 데이터와 API 조회 결과 간 유사도",
                min_value=0, max_value=100,
                format="%.1f%%",
            ),
        }
        for col in result_df.columns:
            if col.startswith("[국민연금]"):
                help_text = (
                    "국민연금 API(V2)는 개인정보보호를 위해 사업자번호 뒷자리를 마스킹(*) 처리하여 제공합니다. "
                    "첫 6자리와 사업장명/주소(시도 검증 포함)를 기반으로 매칭되었습니다."
                )
                if "매칭도" in col:
                    column_config[col] = st.column_config.ProgressColumn(col, help="NPS 매칭 신뢰도 (0.0~1.0)", min_value=0.0, max_value=1.0, format="%.2f")
                else:
                    column_config[col] = st.column_config.TextColumn(col, help=help_text)
            elif col.startswith("[건강보험]"):
                column_config[col] = st.column_config.TextColumn(col, help="건강보험공단 데이터")
            elif col.startswith("[기업정보]"):
                column_config[col] = st.column_config.TextColumn(col, help="금융위원회 기업기본정보")
            elif col.startswith("[재무정보]"):
                help_text = (
                    "금융감독원(DART) 공시 대상 기업이 아닌 경우 조회가 불가할 수 있습니다. "
                    "법인등록번호(13자리)를 기반으로 조회됩니다."
                )
                column_config[col] = st.column_config.TextColumn(col, help=help_text)

        def style_rows(row):
            """시도가 다른 경우 또는 NPS 매칭도가 낮은 경우 강조"""
            u_sido = str(row.get("[입력] 시도", ""))
            a_sido = str(row.get("[행정] 시도", ""))
            nps_score = row.get("[국민연금] 매칭도", 0.0)
            
            styles = [''] * len(row)
            
            # 1. 시도 불일치 (빨간 글씨 + 노란 배경)
            if u_sido and a_sido and u_sido != a_sido:
                styles = ['background-color: #FFFFE0; color: #E03131; font-weight: bold;'] * len(row)
            
            # 2. NPS 매칭도 낮음 (회색 배경 + 이탤릭) - 시도 불일치보다 우선순위 낮음
            elif 0 < nps_score < 0.7:
                styles = ['background-color: #F8F9FA; color: #868E96; border-left: 3px solid #FAB005;'] * len(row)
                
            return styles

        st.dataframe(
            result_df.style.apply(style_rows, axis=1),
            use_container_width=True,
            hide_index=True,
            column_config=column_config,
        )

        # 다운로드
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="qx-section-label">DOWNLOAD RESULTS</div>', unsafe_allow_html=True)

        from utils.excel_handler import export_result_excel
        col_dl, _ = st.columns([2, 2])
        with col_dl:
            try:
                excel_bytes = export_result_excel(result_df)
                st.download_button(
                    label="📥 조회 결과 엑셀 다운로드 (.xlsx)",
                    data=excel_bytes,
                    file_name="사업체정보_행정정보조회결과.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="btn_biz_download_result",
                )
            except Exception as e:
                st.error(f"엑셀 변환 중 오류: {e}")

        # 결과 초기화 버튼
        if st.button("🔄 결과 초기화", key="btn_biz_reset"):
            for k in ["biz_crawl_result", "biz_crawl_stats", "biz_crawl_selected_nps", "biz_crawl_selected_nhis"]:
                if k in st.session_state:
                    del st.session_state[k]
            st.rerun()




def _auto_detect_col(col_options: list, keywords: list) -> int:
    """컬럼 목록에서 키워드와 매칭되는 컬럼의 인덱스를 반환 (자동 감지)"""
    for i, col in enumerate(col_options):
        for kw in keywords:
            if kw in str(col):
                return i
    return 0  # (선택 안 함)


def show_unit_nonresponse_system():
    """AI 단위 무응답 검토 및 가중치 조정(Weighting) 시스템 UI"""
    st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">AI 단위 무응답 검토</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">표본 편향 교정 솔루션</span>
        <span class="qx-topbar-badge">Raking, Weighting</span>
    </div>
    """, unsafe_allow_html=True)

    # [v4.8 고도화] 실무용 이용 가이드
    with st.expander("📘 AI 단위 무응답(가중치) 검토 - 이용 방법 및 데이터 프로세스 가이드", expanded=False):
        st.markdown("### 🗺️ 데이터 처리 로드맵")
        show_data_roadmap(4)
        st.markdown("""
        ### 🛠️ 데이터 보정 및 가중치(Weighting) 단계
        1. **데이터 및 모집단 정보 입력:** 
           - 분석 대상 데이터와 변수 정의(코드북)를 업로드합니다.
           - 가중치 산출의 기준이 될 인구통계 변수(성별, 연령, 지역 등)를 선택합니다.
        2. **모집단 목표 분포 설정:** 선택한 각 계층별 모집단(Target) 비율(%)을 입력합니다. (합계 100% 필수)
        3. **데이터 증계 (Data Augmentation):** 
           - **세그먼트 타격:** 표본이 극히 부족하여 가중치만으로 보정이 어려운 특정 집단(예: 20대 남성)을 지정합니다.
           - **합성 데이터 생성:** Bootstrap 또는 Noise-added 기법을 통해 통계적 유의성을 가진 가상 표본을 생성하여 쿼터를 충족시킵니다.
        4. **가중치 산출(Raking):** RIM Weighting 알고리즘을 통해 다차원 주변 분포를 모집단 비율에 맞게 반복 보정합니다.
        5. **보정 품질 검정 (Diagnostics):** 
           - **Deff(설계효과):** 가중치 부여로 인한 분산 증가 정도를 확인합니다. (1.5 이하 권장)
           - **ESS(유효표본):** 가중치 적용 후 실제 분석에 유효한 정보량을 확인합니다.
        6. **통합 시트 다운로드:** 가중치 변수가 포함된 최종 분석용 데이터셋을 엑셀로 추출합니다.
        """)

    # [v4.1 고도화] 전문가용 가이드 (수식 포함)
    with st.expander("📘 AI 단위 무응답(가중치) 검토 - 통계적 모델 설계 및 수식 안내", expanded=False):
        st.markdown(r"""
        ### 🔍 통계적 가중치 조정(Weighting) 모델
        단위 무응답으로 인한 표본 편향을 교정하기 위해 본 시스템은 **RIM Weighting (Raking)** 알고리즘을 채택합니다.
        
        #### **1. Raking (Iterative Proportional Fitting) 알고리즘**
        여러 인구통계 변수(성별, 연령 등)의 분포를 모집단(Target) 분포에 맞추기 위해 가중치를 반복적으로 업데이트합니다.
        - **수렴 조건:** 모든 차원에서의 주변 분포(Marginal Distribution)오차가 설정된 허용 오차($\epsilon < 10^{-4}$) 이내일 때 종료됩니다.

        #### **2. 가중치 품질 평가지표 (Diagnostic Metrics)**
        가중치 부여는 추정치의 분산을 증가시킬 수 있으므로, **Kish의 설계효과(Design Effect, Deff)**를 통해 보정의 품질을 평가합니다.
        """)
        st.latex(r"Deff \approx 1 + L = \frac{n \sum_{i=1}^{n} w_i^2}{(\sum_{i=1}^{n} w_i)^2}")
        st.markdown(r"""
        - **유효 표본 크기 (Effective Sample Size, ESS):** 가중치 적용 후의 실제 정보량을 나타내며, $ESS = \frac{n}{Deff}$ 로 산출됩니다. 
        - **해석:** $Deff$ 가 1.5 이하일 경우 통계적으로 모델의 안정성이 높다고 판단하며, 2.0을 초과할 경우 특정 계층에 과도한 가중치가 부여되었음을 시사합니다.
        """)

    # 빈 상태 안내 (2컬럼 업로드 레이아웃)
    st.markdown('<div class="qx-section-label">1. 데이터 및 코드북 업로드 (Survey Data)</div>', unsafe_allow_html=True)
    col_u1, col_u2 = st.columns(2)
    with col_u1:
        df_file = st.file_uploader("가중치 조정을 수행할 데이터 업로드", type=["xlsx", "csv"], key="uploader_unit")
    with col_u2:
        cb_file = st.file_uploader("코드북 업로드 (선택 사항)", type=["xlsx"], key="cb_unit")

    if not df_file:
        st.markdown("""
<div class="qx-card" style="text-align:center; padding:2rem 2rem 2.5rem 2rem; margin-top: 1rem; margin-bottom: 2rem; min-height: 320px;">
    <div style="font-size:3rem; margin-bottom:0.5rem;">⚖️</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">
        무응답 교정을 위한 데이터를 업로드하세요
    </div>
    <div style="font-size:0.87rem; color:#8B96A9; margin-bottom:1.5rem;">
        응답 표본과 모집단 간의 차이를 분석하고 통계적 가중치(Weighting)를 부여합니다.
    </div>
    <div style="display:flex; justify-content:center; gap:1.5rem; flex-wrap:wrap;">
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">📊</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">편향 진단</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🔢</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">Raking 보정</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">📉</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">Deff 평가</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
        st.markdown("<div style='margin-bottom: 5rem;'></div>", unsafe_allow_html=True) 
        return

    # 데이터 로드
    try:
        df = pd.read_csv(df_file) if df_file.name.endswith(".csv") else pd.read_excel(df_file)
        st.success(f"데이터 로드 완료: {len(df)} 행")
    except Exception as e:
        st.error(f"로드 중 오류: {e}")
        return

    # [v4.6] 코드북 파싱 및 매핑
    cb_parser = None
    if cb_file:
        try:
            with st.spinner("코드북 분석 중..."):
                cb_parser = CodebookParser(cb_file)
            st.success("코드북 연동 완료: 변수 및 응답값 라벨이 활성화되었습니다.")
        except Exception as e:
            st.error(f"코드북 데이터 읽기 오류: {e}")

    # 변수 선택
    st.markdown('<div class="qx-section-label">2. 가중치 보정 변수 설정</div>', unsafe_allow_html=True)
    all_col_labels = cb_parser.get_all_var_labels(df.columns) if cb_parser else df.columns.tolist()
    selected_weight_labels = st.multiselect("가중치를 부여할 기준 변수를 선택하세요 (예: 성별, 연령)", options=all_col_labels)
    weight_vars = [cb_parser.get_column_from_label(lb) for lb in selected_weight_labels] if cb_parser else selected_weight_labels

    if not weight_vars:
        st.info("변수를 선택하면 모집단 비율 입력란이 나타납니다.")
        return

    # 목표 비율 입력
    targets = {}
    st.markdown("##### 📍 모집단 목표 분포 입력 (%)")
    for i, var in enumerate(weight_vars):
        display_name = selected_weight_labels[i] if cb_parser else var
        with st.expander(f"변수: {display_name}", expanded=True):
            unique_vals = sorted(df[var].dropna().unique().tolist())
            targets[var] = {}
            
            # 코드북 라벨 적용 (응답값)
            cols = st.columns(len(unique_vals))
            for j, val in enumerate(unique_vals):
                label_val = cb_parser.get_value_label(var, val) if cb_parser else val
                with cols[j]:
                    prop = st.number_input(f"{label_val} (%)", min_value=0.0, max_value=100.0, value=100.0/len(unique_vals), key=f"target_{var}_{val}")
                    targets[var][val] = prop / 100.0
            
            # 합계 체크
            total_p = sum(targets[var].values())
            if abs(total_p - 1.0) > 0.001:
                st.warning(f"합계가 {total_p*100:.1f}%입니다. 100%가 되도록 조정하세요.")

    # [v4.7 추가] 목표 기반 데이터 증계 (Data Augmentation)
    st.markdown('<div class="qx-section-label">3. 목표 기반 데이터 증계 (Target-Driven Augmentation)</div>', unsafe_allow_html=True)
    with st.expander("🛠️ 소수 표본 및 쿼터 미달 세그먼트 데이터 생성", expanded=False):
        st.markdown("""
        표본 수가 부족한 특정 집단(Segment)에 대해 통계적 기법으로 데이터를 생성하여 목표 쿼터를 충족시킵니다.
        """)
        
        aug_cols = st.columns([2, 2, 1])
        with aug_cols[0]:
            sel_aug_vars = st.multiselect("증계 조건 변수 선택", options=all_col_labels, key="aug_vars")
            target_segment_cols = [cb_parser.get_column_from_label(lb) for lb in sel_aug_vars] if cb_parser else sel_aug_vars
        
        filter_dict = {}
        if target_segment_cols:
            st.markdown("##### 📍 세그먼트 조건 설정")
            f_cols = st.columns(len(target_segment_cols))
            for i, col in enumerate(target_segment_cols):
                unique_vals = sorted(df[col].dropna().unique().tolist())
                labels = [cb_parser.get_value_label(col, v) for v in unique_vals] if cb_parser else unique_vals
                with f_cols[i]:
                    sel_labels = st.multiselect(f"{sel_aug_vars[i]} 선택", options=labels, key=f"f_{col}")
                    if sel_labels:
                        if cb_parser:
                            rev_map = {cb_parser.get_value_label(col, v): v for v in unique_vals}
                            filter_dict[col] = [rev_map[lb] for lb in sel_labels]
                        else:
                            filter_dict[col] = sel_labels

        if filter_dict:
            mask = pd.Series([True] * len(df), index=df.index)
            for c, v in filter_dict.items():
                mask &= df[c].isin(v)
            current_count = len(df[mask])
            
            st.info(f"선택한 세그먼트의 현재 표본 수: **{current_count}**개")
            
            c1, c2 = st.columns(2)
            with c1:
                target_count = st.number_input("목표 표본 수", min_value=current_count + 1, value=max(current_count + 1, 20), key="aug_target_count")
            with c2:
                aug_method = st.selectbox("증계 기법", options=["Bootstrap", "Noise-added"], key="aug_method")
            
            if st.button("✨ 데이터 증계 실행", use_container_width=True):
                augmentor = DataAugmentor(df)
                added = augmentor.augment_by_filter(filter_dict, target_count, method=aug_method)
                if added > 0:
                    st.success(f"성공적으로 {added}개의 새로운 레코드가 생성되었습니다. (총 {len(augmentor.df)}명)")
                    st.session_state["aug_df"] = augmentor.df
                    st.session_state["aug_summary"] = augmentor.get_summary()
                    st.rerun() # 전체 UI 갱신 (가중치 섹션에도 반영되도록)
                else:
                    st.warning("증계할 데이터가 없거나 조건이 맞지 않습니다.")

        if "aug_summary" in st.session_state:
            s = st.session_state["aug_summary"]
            st.markdown(f"""
            - **전체 표본:** {s['total_count']}명 (원본: {s['original_count']}명 / 증계: {s['augmented_count']}명)
            - **증계 비율:** {s['augmented_ratio']}%
            """)
            if st.button("🗑️ 증계 취소 (원본 복구)"):
                if "aug_df" in st.session_state: del st.session_state["aug_df"]
                if "aug_summary" in st.session_state: del st.session_state["aug_summary"]
                st.rerun()

    # 증계된 데이터가 있으면 그것을 사용하여 가중치 산출 진행
    active_df = st.session_state["aug_df"] if "aug_df" in st.session_state else df

    # 4. 가중치 산출 실행
    st.markdown('<div class="qx-section-label">4. 가중치(Weighting) 산출 및 검정</div>', unsafe_allow_html=True)
    if st.button("🚀 가중치 산출(Raking) 실행", type="primary", use_container_width=True):
        calculator = WeightCalculator(active_df)
        with st.spinner("RIM Weighting 알고리즘 가동 중..."):
            iters, diff = calculator.apply_raking(targets)
            st.session_state["weighted_df"] = calculator.df
            st.session_state["weight_diag"] = calculator.get_diagnostics()
            st.success(f"가중치 산출 완료! (반복 횟수: {iters}, 최종 수렴 오차: {diff:.6f})")

    # 결과 표시
    if "weighted_df" in st.session_state:
        show_security_notice()
        st.markdown('<div class="qx-section-label">5. 분석 결과 및 다운로드</div>', unsafe_allow_html=True)
        diag = st.session_state["weight_diag"]
        
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("가중치 범위", f"{diag['min']:.2f} ~ {diag['max']:.2f}")
        m2.metric("평균 가중치", f"{diag['mean']:.2f}")
        m3.metric("설계효과 (Deff)", f"{diag['deff']:.3f}")
        m4.metric("유효 표본 (ESS)", f"{int(diag['ess'])}명")

        # 가중치 분포 시각화
        fig = px.histogram(st.session_state["weighted_df"], x="weight", title="가중치 분포 히스토그램",
                          template="plotly_white", nbins=30)
        fig.update_traces(marker_color="#0F6CBD")
        st.plotly_chart(fig, use_container_width=True)

        # [v4.9] 리포트 및 Raw Data 분리 다운로드
        report_output = io.BytesIO()
        with pd.ExcelWriter(report_output, engine='xlsxwriter') as writer:
            st.session_state["weighted_df"].to_excel(writer, index=False, sheet_name='Weighted_Data')
            # 진단 정보 추가 (명사형 요약)
            diag_df = pd.DataFrame([st.session_state["weight_diag"]]).T
            diag_df.columns = ["통계치"]
            diag_df.to_excel(writer, index=True, sheet_name='Diagnostics')
        report_output.seek(0)
        
        raw_output = io.BytesIO()
        with pd.ExcelWriter(raw_output, engine='xlsxwriter') as writer:
            st.session_state["weighted_df"].to_excel(writer, index=False, sheet_name='Raw_Data')
        raw_output.seek(0)

        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button("� 가중치 리포트 다운로드 (Excel)", data=report_output, file_name=f"가중치리포트_{df_file.name}", 
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key="dl_weight_report")
        with col_dl2:
            st.download_button("📥 최종 분석용 Raw Data 다운로드", data=raw_output, file_name=f"최종데이터_{df_file.name}", 
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key="dl_weight_raw")


def show_questionnaire_optimization_system():
    """AI 설문지 최적화 컨설턴트 UI (v5.0)"""
    st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">AI 설문지 최적화</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">설문 설계 전문가 컨설팅</span>
        <span class="qx-topbar-badge">Methodology Optimization</span>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("📘 AI 설문지 최적화 - 이용 방법 및 체크리스트 가이드", expanded=False):
        st.markdown("""
        ### 🛠️ AI 설문지 최적화 프로세스
        1. **설문안 입력:** 현재 작성 중인 설문지 텍스트(문항 및 보기)를 아래 입력창에 붙여넣습니다.
        2. **AI 정밀 진단:** 5대 설계 결함(유도 질문, 이중 질문, 보기 정합성, 모호성, 응답 거부 요소)을 전문가 관점에서 분석합니다.
        3. **개선안 도출:** 분석 결과에 따른 'Before & After' 개선 제언과 기대 효과를 확인합니다.
        
        ### 🔍 주요 체크리스트
        - **유도 질문:** "귀하는 성공적인 정책 A에 찬성하십니까?"와 같이 답을 정해둔 질문인가?
        - **이중 질문:** "가격과 품질에 만족하십니까?"처럼 한 문항에 두 가지를 묻는가?
        - **상호배타성:** 보기 간에 중복이 있거나 빠진 항목이 없는가?
        """)

    st.markdown('<div class="qx-section-label">1. 설문지 업로드 및 텍스트 입력</div>', unsafe_allow_html=True)
    
    q_file = st.file_uploader("설문지 파일 업로드 (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], key="q_opt_file_up")
    
    # 새로운 파일이 업로드된 경우 텍스트 추출
    if q_file and q_file.name != st.session_state.get("q_opt_file_name"):
        with st.spinner("설문지 텍스트를 추출하는 중..."):
            from config import MAX_TEXT_CHARS
            text, _ = extract_text(q_file)
            if text:
                st.session_state["q_opt_input_text"] = truncate_text(text, MAX_TEXT_CHARS)
                st.session_state["q_opt_file_name"] = q_file.name
                st.success(f"'{q_file.name}'에서 텍스트를 성공적으로 가져왔습니다.")

    # [v4.17] 텍스트 영역 배경색 흰색 및 테두리 청색 강제 (Surgical CSS)
    st.markdown("""
        <style>
        /* 특정 플레이스홀더를 가진 텍스트 영역과 그 부모 요소들을 타겟팅 */
        div:has(textarea[placeholder*="Q1. 귀하는"]),
        div:has(textarea[placeholder*="Q1. 귀하는"]) > div,
        div[data-baseweb="textarea"]:has(textarea[placeholder*="Q1. 귀하는"]),
        textarea[placeholder*="Q1. 귀하는"] {
            background-color: #FFFFFF !important;
            background: #FFFFFF !important;
            border-color: #0F6CBD !important;
            color: #1A2237 !important;
        }
        
        /* 테두리 두께 및 색상 상시 강조 */
        div[data-baseweb="textarea"]:has(textarea[placeholder*="Q1. 귀하는"]) {
            border: 1px solid #0F6CBD !important;
            border-radius: 4px !important;
        }
        
        /* 포커스 시 박스 쉐도우 유지 */
        div[data-baseweb="textarea"]:has(textarea[placeholder*="Q1. 귀하는"]):focus-within {
            box-shadow: 0 0 0 1px #0F6CBD !important;
            border-color: #0F6CBD !important;
        }
        </style>
    """, unsafe_allow_html=True)

    q_text = st.text_area(
        "분석할 설문 문항 (직접 입력하거나 위에서 파일을 업로드하세요)", 
        value=st.session_state.get("q_opt_input_text", ""),
        height=175, 
        placeholder="예: Q1. 귀하는 본 서비스에 대해 얼마나 만족하십니까?\n1) 매우 만족  2) 만족  3) 보통  4) 불만족  5) 매우 불만족",
        help="파일을 업로드하면 내용이 자동으로 채워집니다. 직접 수정도 가능합니다."
    )
    # 직접 수정한 내용을 세션 상태에 동기화
    st.session_state["q_opt_input_text"] = q_text

    # [v4.15] 초기 진입 화면 가이드 (Landing Card)
    if not q_file and not q_text.strip() and "q_opt_result" not in st.session_state:
        st.markdown("""
<div class="qx-card" style="text-align:center; padding:2rem 2rem 2.5rem 2rem; margin-top: 1rem; margin-bottom: 2rem; min-height: 320px;">
    <div style="font-size:3rem; margin-bottom:0.5rem;">📋</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">
        설문지 최적화를 위한 텍스트를 입력하세요
    </div>
    <div style="font-size:0.87rem; color:#8B96A9; margin-bottom:1.5rem;">
        문항의 논리적 결함과 응답 편향 요소를 전문가 관점에서 분석하고 개선안을 도출합니다.
    </div>
    <div style="display:flex; justify-content:center; gap:1.5rem; flex-wrap:wrap;">
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🔍</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">설계 결함 탐지</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">💡</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">Before & After</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">📉</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">응답 정합성 확보</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
        st.markdown("<div style='margin-bottom: 5rem;'></div>", unsafe_allow_html=True)

    if st.button("🚀 AI 설문지 최적화 시작", type="primary", use_container_width=True):
        if not q_text.strip():
            st.warning("분석할 설문 텍스트를 입력해 주세요.")
        else:
            from prompts import QUESTIONNAIRE_ANALYSIS_PROMPT
            with st.spinner("설문 설계 전문가가 문항을 심층 분석 중입니다..."):
                prompt = QUESTIONNAIRE_ANALYSIS_PROMPT.format(questionnaire_text=q_text)
                res, err = run_analysis(
                    "{report_text}", 
                    prompt, 
                    model_name=st.session_state["selected_model"], 
                    auto_mode=st.session_state["auto_mode"]
                )
                
                if err:
                    st.error(f"분석 중 오류 발생: {err}")
                else:
                    st.session_state["q_opt_result"] = res
                    st.rerun()

    if "q_opt_result" in st.session_state:
        show_security_notice()
        st.markdown('<div class="qx-section-label">2. 전문가 진단 및 개선 제언</div>', unsafe_allow_html=True)
        st.markdown(st.session_state["q_opt_result"], unsafe_allow_html=True)
        
        # 워드 다운로드 지원 - 한 번 클릭으로 바로 다운로드
        md_content = f"# AI 설문지 최적화 분석 보고서\n\n{st.session_state['q_opt_result']}"
        docx = export_to_docx(md_content)
        # 과업명 기반 파일명 생성
        q_project = st.session_state.get("rfp_project_name", "").strip()
        if not q_project:
            q_base = st.session_state.get("q_opt_file_name", "")
            q_project = q_base.rsplit(".", 1)[0] if q_base else "설문지"
        q_safe_name = q_project.replace("/", "_").replace("\\", "_").replace(":", "_")[:80]
        st.download_button(
            "📝 분석 결과 워드 파일로 다운로드", 
            data=docx, 
            file_name=f"{q_safe_name}_설문지최적화.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


def show_sample_design_system():
    """AI 표본설계 자동화 시스템 UI (v5.0)"""
    import io  # [v6.15] 탭 미선택 시 UnboundLocalError 방지
    st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">AI 표본설계</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">주민등록 인구 기반 표본 배분 솔루션</span>
        <span class="qx-topbar-badge">Sample Allocation Engine</span>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("📘 AI 표본설계 - 주요 할당 방식 및 수식 안내", expanded=False):
        st.markdown(r"""
        ### ⚖️ 표본 할당 방식(Allocation Methods)
        
        1. **인구비례할당 (Proportional Allocation):**
           - 각 층(Strata)의 모집단 크기에 정비례하여 표본을 배분합니다. 가장 일반적인 방식입니다.
           - 산식: $n_h = n \times \frac{P_h}{\sum P_h}$
        
        2. **제곱근 비례 할당 (Square Root Proportional):**
           - 소규모 지역/계층의 대표성이 너무 낮아지는 것을 방지하기 위해 인구수의 제곱근에 비례하여 배분합니다.
           - 산식: $n_h = n \times \frac{\sqrt{P_h}}{\sum \sqrt{P_h}}$
        
        3. **최소표본 할당 후 비례할당 (Min-Proportional):**
           - 모든 계층에 최소한의 표본($Min$)을 먼저 보장하고, 남은 표본을 인구비례로 배분합니다.
           - 산식: $n_h = Min + (n - \sum Min) \times \frac{P_h}{\sum P_h}$

        ---
        ### ℹ️ 주민등록 인구 데이터 획득 방법 (행정안전부)
        정확한 표본 배분을 위해 최신 통계 데이터를 활용하십시오.
        1. **홈페이지 접속:** [행정안전부 주민등록 인구통계](https://jumin.mois.go.kr/#)
        2. **메뉴 선택:** 상단 메뉴의 **'연령별 인구현황'** 클릭
        3. **조건 설정:** 조회 조건에서 **'연령구분 단위'**를 **'1세'** 단위로 설정
        4. **데이터 다운로드:** 조회 후 엑셀 파일을 다운로드하여 활용하십시오.
        """)

    # [v6.20] 표본 설계 이용 방법 및 기능 상세 설명 추가
    with st.expander("📖 표본 설계 이용 방법 및 주요 기능 안내", expanded=False):
        st.markdown("""
        ### 🚀 표본 설계란?
        표본 설계(Sample Design)는 모집단의 특성을 잘 반영할 수 있도록 조사 대상을 할당하는 과정입니다. 본 도구는 **행정안전부의 최신 주민등록 인구 통계**를 기반으로 과학적인 표본 배분안을 자동으로 생성합니다.

        ---

        ### 📋 이용 방법
        1.  **데이터 준비**: 행정안전부 인구통계 사이트에서 받은 엑셀 파일을 업로드하거나, 수동으로 데이터를 입력합니다.
        2.  **모집단 확정**: 지역 레벨(광역/기초)과 연령 범위를 설정하고 '모집단으로 확정' 버튼을 누릅니다.
        3.  **목표 설정**: 조사의 목표 표본 수(n)를 입력하거나, 원하는 **표본오차(±%)**를 설정하여 필요 인원을 산출합니다.
        4.  **배분 방식 선택**: 
            *   **인구비례할당**: 머릿수 비율에 맞춰 정직하게 나눕니다.
            *   **제곱근 비례**: 인구가 적은 지역의 샘플을 통계적 보정을 위해 조금 더 늘려 배분합니다.
            *   **최소표본 할당**: 모든 그룹에 최소 인원(예: 30명)을 먼저 깔고 나머지를 비례 배분합니다.
        5.  **결과 확인 및 조정**: 생성된 결과표를 확인하고, 필요한 경우 **수동 가감(Manual Tuning)** 기능을 통해 숫자를 미세 조정합니다.

        ---

        ### ✨ 주요 특화 기능
        *   **행안부 데이터 자동 감지**: 복잡한 행안부 엑셀 형식을 읽어 성별/연령별로 자동 정리합니다.
        *   **지역 계층 분석**: 광역 시도 단위뿐만 아니라 시/군/구 단위의 세밀한 지역 설계가 가능합니다.
        *   **정밀 배분 옵션**: 10명, 50명 등 특정 단위로 반올림하여 현장 조사가 용이하도록 관리합니다.
        *   **학교급별 구분**: (8-19세 대상) 초/중/고 학생 단위의 특수 표본 설계를 지원합니다.
        *   **18/19세 지능형 통합**: 시작 연령을 18세 또는 19세로 설정 시, 별도 그룹 대신 자동으로 20대와 통합하여 '18~29세' 등으로 깔끔하게 구성합니다.
        """)

    st.markdown('<div class="qx-section-label">1. 모집단 데이터 입력 (인구 현황)</div>', unsafe_allow_html=True)
    
    tab_file, tab_manual = st.tabs(["📁 파일 업로드", "✍️ 직접 입력"])
    df_raw = None



    with tab_file:
        # ── [v6.14] 다운로드 가이드 ──────────────────────────────────
        with st.expander("📥 행정안전부 인구통계 데이터 다운로드 방법", expanded=False):
            g_img_col, g_txt_col = st.columns([0.8, 1.2], gap="large")
            with g_img_col:
                st.image(
                    "mois_guide.png",
                    caption="행정안전부 주민등록 인구통계 설정 예시",
                    use_container_width=True
                )
            with g_txt_col:
                st.markdown("""
**🌐 사이트:** [jumin.mois.go.kr](https://jumin.mois.go.kr) → **행정동별 연령별 인구현황**

| 단계 | 내용 |
|:---:|------|
| ① | 좌측 메뉴 **연령별 인구현황** 클릭 |
| ② | **연령 구분 단위** 선택 (1세 / 5세 / 10세) |
| ③ | **연령 조회 범위** 설정 (예: 0세 ~ 100세) |
| ④ | **검색** 버튼 클릭 |
| ⑤ | 하단 **xlsx 파일 다운로드** 클릭 후 저장 |

> 💡 다운로드한 파일을 그대로 업로드하면 지역·성별·연령대가 자동으로 추출되어 정리됩니다.
                """)

        # ── 파일 업로드 ──────────────────────────────────────────────
        st.markdown("##### 📁 파일 업로드 (Excel, CSV)")
        pop_file = st.file_uploader(
            "인구 통계 파일 업로드 — 행정안전부 xlsx 또는 일반 CSV/Excel",
            type=["xlsx", "csv"], key="sample_pop_file"
        )

        # [v4.15] 초기 진입 화면 가이드 (Landing Card)
        if not pop_file:
            st.markdown("""
<div class="qx-card" style="text-align:center; padding:2rem 2rem 2.5rem 2rem; margin-top: 1rem; margin-bottom: 2rem; min-height: 320px;">
    <div style="font-size:3rem; margin-bottom:0.5rem;">🎯</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">
        표본 설계를 위해 인구 데이터를 연동하세요
    </div>
    <div style="font-size:0.87rem; color:#8B96A9; margin-bottom:1.5rem;">
        행정안전부 데이터를 기반으로 지역·성별·연령대별 최적의 표본 할당안을 자동으로 구성합니다.
    </div>
    <div style="display:flex; justify-content:center; gap:1.5rem; flex-wrap:wrap;">
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🏘️</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">시군구 비례 배분</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🔢</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">표본 오차 산출</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">📊</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">할당표 자동 생성</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
            st.markdown("<div style='margin-bottom: 5rem;'></div>", unsafe_allow_html=True)

        if pop_file:
            # MOIS 엑셀 자동 감지 (skiprows=3)
            df_mois, is_mois = detect_and_load_mois_excel(pop_file)
            is_mois = False
            uploaded_pop = None
            pop_long_df = None

            try:
                if pop_file.name.endswith(".csv"):
                    uploaded_pop = pd.read_csv(pop_file)
                else:
                    uploaded_pop, is_mois = detect_and_load_mois_excel(pop_file)

                if is_mois:
                    st.success(f"✅ 행정안전부 연령별 인구현황 형식 자동 감지! ({len(uploaded_pop)}개 지역 행, {len(uploaded_pop.columns)}개 컬럼)")
                else:
                    st.success(f"'{pop_file.name}' 로드 완료 ({len(uploaded_pop)}개 행)")
            except Exception as e:
                st.error(f"파일 로드 중 오류: {e}")

            if uploaded_pop is not None:
                if is_mois:
                    # ── MOIS 전용 설정 패널 ──────────────────────────
                    with st.expander("🔧 MOIS 데이터 필터 및 집계 설정", expanded=True):
                        st.markdown("**📍 설계 지역 레벨 및 선택**")
                        
                        # [v6.18] 데이터 구조 분석 후 가용 레벨 자동 제안
                        avail_levels = get_mois_region_levels(uploaded_pop)
                        design_level = st.radio(
                            "표본 설계 단위",
                            options=avail_levels,
                            index=len(avail_levels)-1 if avail_levels else 0,
                            horizontal=True,
                            key="mois_design_level"
                        )

                        chk_col0, chk_col1 = st.columns([1, 2])
                        
                        # 전체 선택 해제 시 개별 선택 초기화
                        def on_all_regions_change():
                            if not st.session_state.get("mois_all_regions", True):
                                if "mois_display_list" in st.session_state:
                                    for r in st.session_state["mois_display_list"]:
                                        st.session_state[f"mois_r_{r}"] = False

                        with chk_col0:
                            all_regions = st.checkbox(
                                "전체 지역 포함", 
                                value=True, 
                                key="mois_all_regions",
                                on_change=on_all_regions_change
                            )
                        with chk_col1:
                            sejong_merge = st.checkbox(
                                "세종특별자치시 → 충청남도 합산",
                                value=False, key="mois_sejong_merge"
                            )

                        # [v6.22] 레벨에 따른 전체 목록 미리 계산 (요약 표시용)
                        code_col = uploaded_pop.columns[0]
                        reg_col = uploaded_pop.columns[1]
                        cvals = uploaded_pop[code_col].astype(str).str.strip()
                        
                        if design_level == "광역 시도 단위":
                            mask = cvals.str.endswith("00000000") & (cvals != "0000000000")
                        elif design_level == "기초 시/군/구 단위":
                            # 자식이 있는 부모는 남기고, 자식들은 제외
                            sigungu_mask = cvals.str.endswith("00000") & ~cvals.str.endswith("00000000")
                            sigungu_codes = cvals[sigungu_mask].unique()
                            parents = [c for c in sigungu_codes if c.endswith("000000")]
                            children = [c for c in sigungu_codes if not c.endswith("000000")]
                            to_exclude = []
                            for p in parents:
                                if any(c.startswith(p[:4]) for c in children):
                                    to_exclude.extend([c for c in children if c.startswith(p[:4])])
                            mask = sigungu_mask & ~cvals.isin(to_exclude)
                        else: # 시군구별 상세 단위
                            # 자식이 있는 부모는 제외
                            sigungu_mask = cvals.str.endswith("00000") & ~cvals.str.endswith("00000000")
                            sigungu_codes = cvals[sigungu_mask].unique()
                            parents = [c for c in sigungu_codes if c.endswith("000000")]
                            children = [c for c in sigungu_codes if not c.endswith("000000")]
                            to_exclude = []
                            for p in parents:
                                if any(c.startswith(p[:4]) for c in children):
                                    to_exclude.append(p)
                            mask = sigungu_mask & ~cvals.isin(to_exclude)

                        full_display_list = [r.split('(')[0].strip() for r in uploaded_pop[mask][reg_col].tolist()]
                        
                        if sejong_merge and "세종특별자치시" in full_display_list:
                            full_display_list = [r for r in full_display_list if r != "세종특별자치시"]
                        
                        st.session_state["mois_display_list"] = full_display_list

                        # [v6.22] 현재 선택 현황 요약 표시
                        if all_regions:
                            st.info(f"📍 **{design_level}**로 설정됨: 현재 총 **{len(full_display_list)}개** 지역이 모두 포함되어 있습니다.")
                            with st.expander("🔍 포함된 지역 명단 확인하기", expanded=False):
                                st.caption(", ".join(full_display_list))
                        
                        if not all_regions:
                            st.warning(f"⚠️ 아래 목록에서 분석에 포함할 지역을 직접 선택해 주세요. (현재 레벨: {design_level})")
                            display_list = full_display_list
                            
                            row_cols = 4
                            cols_list = st.columns(row_cols)
                            selected_regions = []
                            for i, region in enumerate(display_list):
                                with cols_list[i % row_cols]:
                                    if st.checkbox(region, value=False, key=f"mois_r_{region}"):
                                        selected_regions.append(region)
                        else:
                            selected_regions = full_display_list

                        st.markdown("**📊 연령 설정**")
                        fa1, fa2 = st.columns(2)
                        with fa1:
                            age_interval = st.selectbox(
                                "연령 구분 단위",
                                options=[1, 5, 10],
                                index=2,
                                format_func=lambda x: f"{x}세 단위",
                                key="mois_age_interval"
                            )
                        with fa2:
                            age_range = st.slider(
                                "연령 범위",
                                min_value=0, max_value=100,
                                value=(19, 100),
                                step=1,
                                key="mois_age_range"
                            )

                        # [v6.16] 학교급별 구분 옵션 (19세 이하 포함 시 노출)
                        school_level_opt = False
                        if age_range[0] <= 19:
                            school_level_opt = st.checkbox(
                                "19세 이하 학교급별 구분 (초등/중등/고등)",
                                value=False,
                                help="8-13세: 초등, 14-16세: 중등, 17-19세: 고등",
                                key="mois_school_level_opt"
                            )

                        # [v6.23] 상위 연령대 통합 옵션
                        upper_age_map = {"선택 안함": None, "60대 이상": 60, "70대 이상": 70, "80대 이상": 80}
                        upper_age_sel = st.selectbox(
                            "상위 연령대 통합",
                            options=list(upper_age_map.keys()),
                            index=0,
                            help="선택 시 해당 연령대 이상을 하나의 그룹으로 통합합니다. 미선택 시 10세 단위로 구분됩니다.",
                            key="mois_upper_age_group"
                        )
                        upper_age_cutoff = upper_age_map[upper_age_sel]

                    # MOIS 데이터 추출
                    try:
                        pop_long_df = parse_mois_excel_with_gender(
                            uploaded_pop,
                            regions=selected_regions,
                            level=design_level if is_mois else "광역 시도 단위",
                            min_age=age_range[0],
                            max_age=age_range[1],
                            interval=age_interval,
                            include_sejong_in_chungnam=sejong_merge,
                            school_level_option=school_level_opt,
                            upper_age_cutoff=upper_age_cutoff
                        )
                        if pop_long_df is not None:
                            st.markdown(f"**📋 데이터 추출 결과 미리보기** — {len(pop_long_df)}개 층, 총인구 {pop_long_df['인구수'].sum():,}명")
                            st.dataframe(pop_long_df.head(15), hide_index=True, use_container_width=True)
                            # 확정 버튼
                            if st.button("✅ 추출된 데이터를 모집단으로 확정", use_container_width=True, key="mois_confirm"):
                                st.session_state["pop_source_df"] = pop_long_df
                                st.session_state["pop_col_hint"] = "인구수"
                                st.session_state["strata_cols_hint"] = ["지역", "성별", "연령대"]
                                st.rerun()
                        else:
                            st.warning("추출된 데이터가 없습니다. 지역 선택 또는 연령 범위를 확인해주세요.")
                    except Exception as e:
                        st.error(f"MOIS 데이터 해석 오류: {e}")

                else:
                    # ── 일반 Excel/CSV 필터 (v6.13 유지) ────────────
                    with st.expander("🔧 데이터 필터 설정 (연령대 / 성별)", expanded=True):
                        fu1, fu2, fu3 = st.columns([1, 1, 1])
                        with fu1:
                            age_col_candidates = [c for c in uploaded_pop.columns if any(kw in str(c) for kw in ["연령", "age", "나이", "Age"])]
                            age_col_file = st.selectbox(
                                "연령대 컬럼 선택",
                                options=["(없음)"] + uploaded_pop.columns.tolist(),
                                index=(1 + uploaded_pop.columns.tolist().index(age_col_candidates[0])) if age_col_candidates else 0,
                                key="age_col_file"
                            )
                        with fu2:
                            gender_col_candidates = [c for c in uploaded_pop.columns if any(kw in str(c) for kw in ["성별", "gender", "sex", "성"])]
                            gender_col_file = st.selectbox(
                                "성별 컬럼 선택",
                                options=["(없음)"] + uploaded_pop.columns.tolist(),
                                index=(1 + uploaded_pop.columns.tolist().index(gender_col_candidates[0])) if gender_col_candidates else 0,
                                key="gender_col_file"
                            )
                        with fu3:
                            if gender_col_file != "(없음)":
                                gender_vals = uploaded_pop[gender_col_file].dropna().unique().tolist()
                                selected_genders = st.multiselect("포함할 성별 값", options=gender_vals, default=gender_vals, key="gender_vals_file")
                            else:
                                selected_genders = []

                    filtered_pop = uploaded_pop.copy()
                    if gender_col_file != "(없음)" and selected_genders:
                        filtered_pop = filtered_pop[filtered_pop[gender_col_file].isin(selected_genders)]

                    st.markdown(f"**미리보기** (필터 후 {len(filtered_pop)}행)")
                    st.dataframe(filtered_pop.head(10), hide_index=True, use_container_width=True)
                    if st.button("✅ 업로드 파일을 모집단으로 확정", use_container_width=True, key="generic_confirm"):
                        st.session_state["pop_source_df"] = filtered_pop
                        st.rerun()


    with tab_manual:
        st.markdown("##### ✍️ 직접 입력 / 예시 데이터")
        default_data = "지역, 성별, 연령대, 인구수\n서울, 남, 20대, 650000\n서울, 여, 20대, 680000\n부산, 남, 20대, 210000\n부산, 여, 20대, 220000\n인천, 남, 20대, 195000\n인천, 여, 20대, 198000"
        pop_input = st.text_area(
            "데이터를 CSV 형식으로 입력하세요", 
            value=default_data,
            height=200,
            help="첫 줄에 헤더(지역, 성별 등)를 포함하여 입력하세요."
        )
        if st.button("✅ 입력 데이터를 모집단으로 확정", use_container_width=True):
            try:
                import io
                st.session_state["pop_source_df"] = pd.read_csv(io.StringIO(pop_input.strip()), skipinitialspace=True)
                st.rerun()
            except Exception as e:
                st.error(f"입력 데이터 분석 오류: {e}")

    # 최종 모집단 데이터 확정 확인
    if "pop_source_df" not in st.session_state:
        st.info("상단 탭 중 하나를 선택하여 모집단 인구 데이터를 로드하고 '확정' 버튼을 클릭해 주세요.")
        return

    df_raw = st.session_state["pop_source_df"]

    # [v6.15] 컬럼 설정 — MOIS 확정 시 저장한 hint 우선 사용, 없으면 키워드 감지
    cols = df_raw.columns.tolist()

    # ① pop_col 결정
    pop_col_hint = st.session_state.get("pop_col_hint")
    if pop_col_hint and pop_col_hint in cols:
        pop_col = pop_col_hint
    else:
        pop_col = None
        for c in cols:
            if any(kw in str(c) for kw in ["인구", "population", "count", "N수"]):
                pop_col = c
                break
        if not pop_col:
            pop_col = cols[-1]

    # ② strata_cols 결정
    strata_hint = st.session_state.get("strata_cols_hint")
    if strata_hint and all(c in cols for c in strata_hint):
        strata_cols = strata_hint
    else:
        strata_cols = [c for c in cols if c != pop_col]

    df_raw[pop_col] = pd.to_numeric(df_raw[pop_col], errors="coerce").fillna(0)

    # 층화 구조 표시
    is_mois_data = strata_hint is not None and "성별" in strata_cols and "연령대" in strata_cols
    layer_label = " > ".join(strata_cols)
    if is_mois_data:
        st.info(f"🧬 **층화 구조:** {layer_label} | **총 인구:** {df_raw[pop_col].sum():,} | 📋 MOIS 행정안전부 데이터")
    else:
        st.info(f"🧬 **층화 구조:** {layer_label} | **총 인구:** {df_raw[pop_col].sum():,}")

    st.markdown('<div class="qx-section-label">2. 표본 설계 설정</div>', unsafe_allow_html=True)
    
    # [v6.19] 정밀 표본 배분 설정 (목표 오차 기준 모드 추가)
    c1, c2, c3, c4 = st.columns([2, 1, 1.5, 1.5])
    with c1:
        alloc_mode = st.radio("n 결정 방식", options=["전체 표본 수 지정", "목표 표본오차 기준"], horizontal=True, key="alloc_mode")
    
    with c2:
        if alloc_mode == "전체 표본 수 지정":
            total_n = st.number_input("전체 표본 수(n)", min_value=1, value=1000, key="total_n_input")
        else:
            # 목표 오차 기반 n 역산 (p=0.5, 95% 신뢰수준)
            target_err = st.number_input("목표 오차(±%)", min_value=0.1, max_value=10.0, value=3.1, step=0.1)
            total_n = int((1.96**2 * 0.25) / ((target_err/100)**2))
            st.caption(f"계산된 n: {total_n}명")

    with c3:
        method = st.selectbox("할당 방식", options=["인구비례할당", "제곱근 비례 할당", "최소표본 할당 후 비례할당"], key="allocation_method")
    
    with c4:
        min_n_val = st.number_input("최소 할당(Min)", min_value=1, value=30, disabled=(method != "최소표본 할당 후 비례할당"), key="min_n_val")

    # [v6.19] 정밀 배분 옵션 (반올림 및 단위 설정)
    with st.expander("🛠️ 정밀 배분 고급 옵션 (단위 반올림 등)", expanded=False):
        fo1, fo2 = st.columns(2)
        with fo1:
            round_unit = st.selectbox("배분 단위 (반올림)", options=[1, 5, 10, 50, 100], index=0, help="할당된 표본을 특정 단위(예: 10명 단위)로 정렬합니다.")
        with fo2:
            st.info("💡 배분 단위를 높이면 목표 n과 최종 합계에 차이가 발생할 수 있으며, 시스템이 이를 자동으로 보정합니다.")

    if st.button("📊 표본 배분 계산 실행", type="primary", use_container_width=True):
        total_pop = df_raw[pop_col].sum()
        df_work = df_raw.copy()
        
        # 1. 기본 할당 계산
        if method == "인구비례할당":
            df_work["allocated"] = (df_work[pop_col] / total_pop * total_n)
        elif method == "제곱근 비례 할당":
            sqrt_sum = df_work[pop_col].apply(np.sqrt).sum()
            df_work["allocated"] = (df_work[pop_col].apply(np.sqrt) / sqrt_sum * total_n)
        elif method == "최소표본 할당 후 비례할당":
            num_groups = len(df_work)
            remaining_n = total_n - (min_n_val * num_groups)
            if remaining_n < 0:
                st.error("최소 할당 합계가 전체 n보다 큽니다.")
                return
            df_work["allocated"] = min_n_val + (df_work[pop_col] / total_pop * remaining_n)
        
        # 2. 정밀 배분 (반올림 단위 적용)
        df_work["final_n"] = (df_work["allocated"] / round_unit).round().fillna(0).astype(int) * round_unit
        
        # 3. 합계 보정 (라운딩으로 인한 차이 조정)
        diff = total_n - df_work["final_n"].sum()
        if diff != 0:
            # 보정 단위 계산
            steps = int(abs(diff) / round_unit)
            if steps > 0:
                df_work["remainder"] = df_work["allocated"] - df_work["final_n"]
                if diff > 0:
                    idx = df_work.nlargest(steps, "remainder").index
                    df_work.loc[idx, "final_n"] += round_unit
                else:
                    idx = df_work.nsmallest(steps, "remainder").index
                    df_work.loc[idx, "final_n"] -= round_unit

        df_work["비율(%)"] = (df_work["final_n"] / df_work["final_n"].sum() * 100).round(1)
        st.session_state["sample_design_df"] = df_work[cols + ["final_n", "비율(%)"]]
        st.session_state["sample_design_meta"] = {"pop_col": pop_col, "strata_cols": strata_cols, "round_unit": round_unit}
        st.rerun()

    if "sample_design_df" in st.session_state:
        st.markdown('<div class="qx-section-label">3. 표본 할당 결과 및 정밀 조정</div>', unsafe_allow_html=True)
        res_df = st.session_state["sample_design_df"]
        meta = st.session_state["sample_design_meta"]
        
        # [v6.19] 수동 정밀 조정 기능 (Data Editor)
        with st.expander("📝 세포별 정밀 수동 가감 (Manual Tuning)", expanded=False):
            st.caption("배분된 수치를 직접 수정할 수 있습니다. 수정 후 '표본 설계 결과 업데이트' 버튼을 클릭하세요.")
            edited_df = st.data_editor(
                res_df,
                column_config={
                    "final_n": st.column_config.NumberColumn("확정 표본(n)", min_value=0, step=st.session_state["sample_design_meta"].get("round_unit", 1)),
                    "비율(%)": st.column_config.NumberColumn("비율(%)", disabled=True)
                },
                disabled=[c for c in res_df.columns if c != "final_n"],
                use_container_width=True,
                key="sample_editor"
            )
            if st.button("🔄 정밀 조정 결과 반영", key="update_manual_tuning"):
                edited_df["비율(%)"] = (edited_df["final_n"] / edited_df["final_n"].sum() * 100).round(1)
                st.session_state["sample_design_df"] = edited_df
                st.rerun()

        sc1, sc2, sc3, sc4 = st.columns(4)
        current_total = res_df['final_n'].sum()
        sc1.metric("최종 인원 (Total)", f"{current_total:,}명")

        # ── [v6.14] 피벗 테이블 (지역×성별×연령대) ─────────────────
        pivot_df = format_sample_pivot_table(res_df)

        if pivot_df is not None:
            st.markdown("##### 📊 표본 배분 결과표")

            # 스타일 함수 — 총계 행: 노란 배경 / 남 상위헤더: 하늘색 / 여: 연한 분홍
            def style_pivot(styler):
                # 총계 행 강조
                styler.apply(
                    lambda row: ["background-color:#FFF176; font-weight:bold" if row.name == "총계" else "" for _ in row],
                    axis=1
                )
                return styler

            pivot_int = pivot_df.fillna(0).astype(int)
            st.dataframe(
                pivot_int.style.pipe(style_pivot),
                use_container_width=True
            )

            # 피벗표 Excel 다운로드
            output_pivot = io.BytesIO()
            with pd.ExcelWriter(output_pivot, engine='xlsxwriter') as writer:
                wb = writer.book

                # ── 공통 포맷 정의 (음영 없음 + 검정 테두리) ─────────────
                border = {'border': 1, 'border_color': '#000000', 'bg_color': '#FFFFFF'}
                fmt_normal   = wb.add_format({**border})
                fmt_header   = wb.add_format({**border, 'bold': True, 'align': 'center', 'valign': 'vcenter'})
                fmt_total_row = wb.add_format({**border, 'bold': True})
                fmt_index    = wb.add_format({**border, 'bold': False})

                # ── 피벗 시트 ──────────────────────────────────────────────
                ws_piv = wb.add_worksheet('표본배분_피벗')

                # 헤더 행 (0행): 인덱스명 + 열명
                ws_piv.write(0, 0, '지역 구분', fmt_header)
                for ci, col in enumerate(pivot_int.columns):
                    ws_piv.write(0, ci + 1, str(col), fmt_header)

                # 데이터 행
                for ri, (idx, row_data) in enumerate(pivot_int.iterrows()):
                    row_num = ri + 1
                    is_total = str(idx) == '총계'
                    idx_fmt  = fmt_total_row if is_total else fmt_index
                    cell_fmt = fmt_total_row if is_total else fmt_normal
                    ws_piv.write(row_num, 0, str(idx), idx_fmt)
                    for ci, val in enumerate(row_data):
                        ws_piv.write(row_num, ci + 1, int(val), cell_fmt)

                # 열 너비 조정
                ws_piv.set_column(0, 0, 18)
                ws_piv.set_column(1, len(pivot_int.columns), 9)

                # ── 상세 시트 ──────────────────────────────────────────────
                ws_det = wb.add_worksheet('표본배분_상세')
                for ci, col in enumerate(res_df.columns):
                    ws_det.write(0, ci, str(col), fmt_header)
                for ri, row_data in enumerate(res_df.itertuples(index=False), start=1):
                    for ci, val in enumerate(row_data):
                        # numpy 타입 → Python 네이티브 변환 (xlsxwriter TypeError 방지)
                        if pd.isna(val):
                            ws_det.write(ri, ci, '', fmt_normal)
                        elif isinstance(val, (int, float)):
                            ws_det.write(ri, ci, val, fmt_normal)
                        elif hasattr(val, 'item'):
                            ws_det.write(ri, ci, val.item(), fmt_normal)
                        else:
                            ws_det.write(ri, ci, str(val), fmt_normal)
                ws_det.set_column(0, len(res_df.columns) - 1, 12)

            output_pivot.seek(0)
            st.download_button(
                "📥 표본 설계 결과표 다운로드 (Excel)",
                data=output_pivot,
                file_name="표본설계_배분결과.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="dl_sample_design_pivot"
            )
        else:
            # 피벗 불가 시 기존 테이블 표시
            st.dataframe(
                res_df,
                hide_index=True,
                column_config={
                    meta["pop_col"]: st.column_config.NumberColumn("모집단(P)", format="%d"),
                    "final_n": st.column_config.NumberColumn("확정 표본(n)", format="%d"),
                },
                use_container_width=True
            )
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                wb2 = writer.book
                border2 = {'border': 1, 'border_color': '#000000', 'bg_color': '#FFFFFF'}
                fmt_h2 = wb2.add_format({**border2, 'bold': True, 'align': 'center'})
                fmt_d2 = wb2.add_format({**border2})
                ws2 = wb2.add_worksheet('표본배분_상세')
                for ci, col in enumerate(res_df.columns):
                    ws2.write(0, ci, str(col), fmt_h2)
                for ri, row_data in enumerate(res_df.itertuples(index=False), start=1):
                    for ci, val in enumerate(row_data):
                        # numpy 타입 → Python 네이티브 변환 (xlsxwriter TypeError 방지)
                        if pd.isna(val):
                            ws2.write(ri, ci, '', fmt_d2)
                        elif isinstance(val, (int, float)):
                            ws2.write(ri, ci, val, fmt_d2)
                        elif hasattr(val, 'item'):
                            ws2.write(ri, ci, val.item(), fmt_d2)
                        else:
                            ws2.write(ri, ci, str(val), fmt_d2)
                ws2.set_column(0, len(res_df.columns) - 1, 12)
            output.seek(0)
            st.download_button(
                "📥 상세 표본 설계 내역 다운로드 (Excel)",
                data=output,
                file_name="표본설계_상세배분안.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="dl_sample_design_final"
            )


        # 시각화
        st.markdown("##### 📈 시각화")
        viz_col = st.selectbox("X축 기준 변수 선택", options=meta["strata_cols"], index=0)
        viz_df = res_df.groupby(viz_col)["final_n"].sum().reset_index()
        fig = px.bar(viz_df, x=viz_col, y="final_n", text="final_n",
                     title=f"{viz_col}별 표본 배분 합계", template="plotly_white")
        fig.update_traces(marker_color="#0F6CBD", textposition="outside")
        st.plotly_chart(fig, use_container_width=True)


def show_outlier_inspection_system(mode="outlier"):
    """AI 이상치/결측치 검토 및 보완 시스템 UI"""
    st.markdown(f"""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">AI {'이상치' if mode == 'outlier' else '결측치'} 검토</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">데이터 품질 진단 솔루션</span>
        <span class="qx-topbar-badge">Call Back, {'Adjustment' if mode == 'outlier' else 'Imputation'}</span>
    </div>
    """, unsafe_allow_html=True)

    # [v4.8 고도화] 실무용 이용 가이드 (명사형 어미)
    with st.expander(f"📘 AI {'이상치' if mode == 'outlier' else '결측치'} 검토 - 이용 방법 및 데이터 프로세스 가이드", expanded=False):
        st.markdown("### 🗺️ 데이터 처리 로드맵")
        show_data_roadmap(3)
        if mode == "outlier":
            st.markdown("""
            ### 🛠️ 데이터 이상치(Outlier) 검토 단계
            1. **데이터 업로드 및 연동:** 원본 시계열/조사 데이터와 코드북을 함께 업로드하여 변수 라벨을 활성화합니다.
            2. **시각적 패턴 분석:** 산점도(Scatter)와 바이올린(Violin) 플롯을 통해 통계적 분포를 벗어난 극단값을 시각적으로 확인합니다.
            3. **이상치 사전 진단:** Z-score > 3 기준의 주요 이상 변수 요약 테이블을 통해 보완 우선순위를 판단합니다.
            4. **보완 전략 수립:** 
               - **AI 추천:** 변수의 특성(수치/범주)에 따른 최적 알고리즘 추천 제안을 확인합니다.
               - **통계적 보정:** 전체/층별 평균, k-NN, MICE 등을 지정하여 이상치를 합리적인 값으로 대체합니다.
               - **재확인(Call-back):** 기입 오류가 아닌 실제 값 확인이 필요한 경우 재조사 대상으로 분류합니다.
            5. **최종 로그 검수:** 대체 전/후 값과 적용 사유가 기록된 감사 로그(Audit Log)를 최종 검토합니다.
            """)
        else:
            st.markdown("""
            ### 🛠️ 데이터 결측치(Missing) 보완 단계
            1. **결측 패턴 정밀 진단:** 
               - '결측 히트맵'을 통해 변수 간 결측 상관관계를 파악합니다.
               - AI 진단을 통해 결측 기제를 판별합니다: **MCAR**(완전 무작위 결측, 보완 적합), **MAR**(조건부 무작위, 모델 기반 보완 필요), **MNAR**(비무작위, 보완 부적합 → 재조사 권장)
            2. **보완 알고리즘 엔진 가동:** 
               - **단일 대체(평균/중앙값/최빈값):** 결측률이 5% 미만이고 MCAR일 때 적합한 단순 보완 방식입니다.
               - **랜덤 대체(Random Imputation):** 관측값 분포에서 무작위 추출하여 분산을 보존합니다. 단일 대체의 분산 축소 문제를 해결합니다.
               - **조건부 평균 대체(Conditional Mean):** 상관관계가 높은 변수를 자동 탐지하여 그룹별 평균으로 대체합니다. 변수 간 관계를 반영한 정교한 평균 대체입니다.
               - **다중 대체(MICE):** 변수 간 회귀 관계를 활용하여 정보 손실을 최소화하는 고난도 보완에 적격입니다. MAR 상황에서 가장 권장됩니다.
               - **최근접 이웃(k-NN):** 유사한 응답 패턴을 가진 다른 사례의 값을 참조하여 정교하게 대체합니다.
            3. **조사 가이드 생성:** 보완이 불가능한 필수 항목 결측(MNAR)이나 핵심 문항 결측에 대해 AI가 재조사(Call-back) 스크립트를 자동 생성합니다.
            4. **통합 데이터 배포:** 보완된 데이터와 원본을 대조할 수 있는 'Imputation Marker'가 포함된 감사 리포트(Audit Report)와 최종 Raw Data를 분리 다운로드합니다.
            """)

    # [v4.5 고도화] 전문가용 가이드 (명사형 어미 및 통계 기법 확장)
    with st.expander(f"📘 AI {'이상치' if mode == 'outlier' else '결측치'} 검토 - 통계적 판별 및 보완 알고리즘 안내", expanded=False):
        if mode == "outlier":
            st.markdown(r"""
            ### 🔍 이상치 탐지 모델 (Outlier Detection)
            데이터의 분포 특성에 따른 두 가지 보편적 통계 기준 적용 방식

            #### **1. 표준점수 (Z-score) 기법**
            데이터가 정규분포를 따른다는 가정 하에 평균으로부터 표준편차($\sigma$)의 3배 이상 이탈한 값을 이상치로 판별하는 방식
            """)
            st.latex(r"z = \frac{x - \mu}{\sigma}")
            st.markdown("""
            - **판정 기준:** $|z| > 3.0$ 인 경우 통계적 유의수준 99.7% 범위를 벗어난 극단치로 간주함

            #### **2. IQR (Interquartile Range) 기법**
            비모수적 분포에서도 강건한(Robust) 탐지가 가능한 사분위수 기반의 격리 방식
            - **Upper Fence:** $Q3 + 1.5 \times (Q3 - Q1)$
            - **Lower Fence:** $Q1 - 1.5 \times (Q3 - Q1)$
            - 바이올린 플롯 내부 박스 구조를 통한 시각적 가중치 확인 기능 포함
            """)
        else:
            st.markdown("""
            ### 🔍 결측치 보완 알고리즘 (Imputation Techniques)
            결측 발생 기제(MCAR, MAR, MNAR)에 따른 최적 알고리즘 선택 및 적용

            #### **1. 평균 대체 (Mean/Median/Mode Imputation)**
            결측값을 해당 변수의 **전체 평균(수치형)**, **중앙값**, 또는 **최빈값(범주형)**으로 대체하는 가장 기본적인 단일 대체법
            """)
            st.latex(r"\hat{x}_{missing} = \bar{x}_{observed} = \frac{1}{n_{obs}} \sum_{i=1}^{n_{obs}} x_i")
            st.markdown(r"""
            - **적용 조건:** MCAR 상황, 결측률 5~10% 이내, 정규분포에 가까운 데이터
            - **한계:** 분산이 과소 추정되며$(\hat{\sigma}^2 \downarrow)$, 변수 간 상관관계가 왜곡될 수 있음

            #### **2. 랜덤 대체 (Random Imputation)**
            결측값을 해당 변수의 **관측값 분포에서 무작위 복원 추출(Random Sampling with Replacement)**하여 대체하는 방식. 단순 평균 대체의 **분산 축소(Variance Underestimation)** 문제를 해결함
            """)
            st.latex(r"\hat{x}_{missing} \sim F_{observed}(x) \quad \text{(관측값의 경험적 분포에서 추출)}")
            st.markdown(r"""
            - **장점:** 원래 데이터의 분포 형태(분산, 왜도)를 보존하여 통계적 추론의 편향을 줄임
            - **적용 조건:** MCAR 상황에서 분포 보존이 중요할 때 권장
            - **한계:** 변수 간 공분산(Covariance) 구조를 반영하지 못하므로, 다변량 분석 시 주의 필요

            #### **3. 조건부 평균 대체 (Conditional Mean Imputation)**
            결측 변수와 **상관관계가 가장 높은 보조 변수**를 자동 감지하여, 해당 변수의 구간(Quintile)별 평균으로 대체하는 방식. 단순 전체 평균보다 변수 간 관계를 보존하는 정교한 대체법
            """)
            st.latex(r"\hat{x}_{missing} = E[X | Z = z_k] = \bar{x}_{\{i: z_i \in Q_k\}}")
            st.markdown(r"""
            - **원리:** 상관계수 $|r|$이 가장 높은 변수 $Z$를 기준으로 5분위(Quintile)로 구간화한 뒤, 해당 구간 내 평균으로 대체
            - **적용 조건:** 수치형 변수 간 상관관계가 존재할 때 (자동 탐지, $|r| \geq 0.1$)
            - **장점:** 전체 평균 대체보다 편향이 적고, 층별 평균 대체보다 자동화된 방식

            #### **4. 회귀 대체 및 MICE (Multivariate Imputation by Chained Equations)**
            다변량 데이터의 상관관계 유지를 위한 최신 기법으로, 변수별 결측치를 타 변수들을 독립변수로 하는 회귀 모델을 통해 반복 예측 보완하는 방식 (연쇄 방정식 기반의 회귀 대체 고도화 모델)
            """)
            st.latex(r"Y_j^{(t)} = f(Y_{-j}^{(t-1)}, X, \hat{\beta}_j) + \epsilon_j")
            st.markdown("""
            - **특이점:** 변수 간 상관성을 유지하며 편향을 최소화하는 하이엔드 통계 기법
            - **적용 조건:** MAR 상황, 결측률 20% 이상의 고결측 데이터, 다변량 분석이 목적인 경우
            - **한계:** 모델 가정(선형성, 정규성)에 민감하며, 소규모 데이터에서 불안정할 수 있음

            #### **5. 최근방 대체 및 k-NN (k-Nearest Neighbors) Imputation**
            유사성이 가장 높은 $k$개의 이웃 사례를 추출하여 해당 관측값들의 가중 평균으로 대체하는 최근방 이웃 방식
            """)
            st.latex(r"\hat{x}_{missing} = \frac{\sum_{j=1}^{k} w_j \cdot x_j}{\sum_{j=1}^{k} w_j}, \quad w_j = \frac{1}{d(i,j)}")
            st.markdown("""
            - **거리 척도:** 유클리드 거리(Euclidean Distance)를 활용한 개체 간 유사도 정밀 측정
            - **적용 조건:** 수치형 변수가 다수이고, 유사한 응답 패턴을 가진 사례가 충분할 때
            - **한계:** 고차원 데이터에서 '차원의 저주(Curse of Dimensionality)'로 정확도가 저하될 수 있음

            #### **6. 기타 실무 대체 기법 (참고)**
            - **유사 사례 기증법 (Hot-deck Imputation):** 현재 조사 내 유사 응답자의 값을 직접 복사하여 기증하는 실무 중심형 방식
            - **고정값 대체 (Cold-deck Imputation):** 과거 조사 결과나 외부 데이터를 기준값으로 활용하는 보수적 대체 방식
            """)

    # [v3.6 추가] 주요 기능 요약 카드 (메인 화면 스타일)

    # [v3.6 추가] 주요 기능 요약 카드 (메인 화면 스타일)
    st.markdown('<div class="qx-section-label">SYSTEM FEATURES</div>', unsafe_allow_html=True)
    if mode == "outlier":
        st.markdown("""
<div class="qx-card" style="padding:1.25rem 1.5rem; margin-bottom:1.5rem;">
    <div style="display:grid; grid-template-columns: 1fr 1fr 1fr; gap:1.5rem;">
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">01</span><span>데이터 분포 시각화 (Plotly)</span>
        </div>
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">02</span><span>통계적 이상치 정밀 탐지</span>
        </div>
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">03</span><span>AI 최적 보완 및 결과 리포트</span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
    else:
        st.markdown("""
<div class="qx-card" style="padding:1.25rem 1.5rem; margin-bottom:1.5rem;">
    <div style="display:grid; grid-template-columns: 1fr 1fr 1fr; gap:1.5rem;">
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">01</span><span>결측 패턴 및 AI 유형 진단</span>
        </div>
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">02</span><span>고급 통계 보완 (MICE/k-NN)</span>
        </div>
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">03</span><span>재확인(Call-back) 대상 관리</span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

    # 1. 파일 업로드 (2컬럼 레이아웃)
    st.markdown('<div class="qx-section-label">1. 데이터 및 코드북 업로드 (Excel/CSV)</div>', unsafe_allow_html=True)
    col_up1, col_up2 = st.columns(2)
    with col_up1:
        df_file = st.file_uploader(f"분석할 데이터 업로드 ({mode})", type=["xlsx", "csv"], key=f"uploader_{mode}")
    with col_up2:
        cb_file = st.file_uploader(f"코드북 업로드 (선택 사항)", type=["xlsx"], key=f"cb_{mode}")

    if not df_file:
        if mode == "outlier":
            st.markdown("""
<div class="qx-card" style="text-align:center; padding:2rem 2rem 2.5rem 2rem; margin-top: 1rem; margin-bottom: 2rem; min-height: 320px;">
    <div style="font-size:3rem; margin-bottom:0.5rem;">📈</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">
        검토할 데이터를 업로드하세요
    </div>
    <div style="font-size:0.87rem; color:#8B96A9; margin-bottom:1.5rem;">
        Excel 또는 CSV 데이터를 업로드하면 AI가 이상치 탐지 및 시각적 진단을 수행합니다.
    </div>
    <div style="display:flex; justify-content:center; gap:1.5rem; flex-wrap:wrap;">
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">📍</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">시각적 진단</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">📏</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">통계적 탐지</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🧠</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">AI 보완 방법</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
        else:
            st.markdown("""
<div class="qx-card" style="text-align:center; padding:2rem 2rem 2.5rem 2rem; margin-top: 1rem; margin-bottom: 2rem; min-height: 320px;">
    <div style="font-size:3rem; margin-bottom:0.5rem;">📊</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">
        분석할 데이터를 업로드하세요
    </div>
    <div style="font-size:0.87rem; color:#8B96A9; margin-bottom:1.5rem;">
        데이터를 업로드하면 AI가 결측 패턴을 분석하고 최적의 통계적 보완을 제안합니다.
    </div>
    <div style="display:flex; justify-content:center; gap:1.5rem; flex-wrap:wrap;">
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🔍</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">패턴 분석</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🧪</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">고급 보완(MICE)</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:0.8rem 1.2rem;min-width:120px;">
            <div style="font-size:1.2rem;">🎙️</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.3rem;">조사 가이드</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
        st.markdown("<div style='margin-bottom: 5rem;'></div>", unsafe_allow_html=True) 
        return

    # 데이터 로드
    try:
        if df_file.name.endswith(".csv"):
            df = pd.read_csv(df_file)
        else:
            df = pd.read_excel(df_file)
            
        # [v4.7] 신규 파일 로드 시 이전 보완 결과 초기화 (KeyError 방지)
        if st.session_state.get(f"last_loaded_file_{mode}") != df_file.name:
            st.session_state[f"last_loaded_file_{mode}"] = df_file.name
            keys_to_clear = [f"imputed_df_{mode}", f"impute_summary_{mode}", f"impute_log_{mode}"]
            for k in keys_to_clear:
                if k in st.session_state:
                    del st.session_state[k]
    except Exception as e:
        st.error(f"데이터 로드 중 오류 발생: {e}")
        return

    # [v4.6] 코드북 파킹 및 매핑
    cb_parser = None
    if cb_file:
        try:
            with st.spinner("코드북 분석 중..."):
                cb_parser = CodebookParser(cb_file)
            st.success("코드북 연동 완료: 변수 설명 및 코드표가 활성화되었습니다.")
            
            # [v4.6 추가] 코드북 미리보기 공간
            with st.expander("📘 연동 코드북 상세 정보 (변수설명 및 코드표)", expanded=False):
                tab_cb1, tab_cb2 = st.tabs(["📋 변수 리스트 및 설명", "🔢 코드/라벨 매핑 테이블"])
                with tab_cb1:
                    if cb_parser.var_map:
                        var_preview_df = pd.DataFrame([
                            {"문번호": v['no'], "변수명": k, "변수설명": v['desc']} 
                            for k, v in cb_parser.var_map.items()
                        ])
                        st.dataframe(var_preview_df, use_container_width=True, hide_index=True)
                    else:
                        st.info("읽어온 변수 설명 정보가 없습니다.")
                with tab_cb2:
                    if cb_parser.code_map:
                        all_codes = []
                        for var, codes in cb_parser.code_map.items():
                            for c, l in codes.items():
                                all_codes.append({"변수명": var, "코드": c, "라벨": l})
                        st.dataframe(pd.DataFrame(all_codes), use_container_width=True, hide_index=True)
                    else:
                        st.info("읽어온 코드표 정보가 없습니다.")
        except Exception as e:
            st.error(f"코드북 데이터 읽기 오류: {e}")

    st.success(f"데이터 로드 완료: {len(df)} 행, {len(df.columns)} 열")
    
    # [v3.5 추가] 시각적 이상치 판별 섹션
    if mode == "outlier":
        with st.expander("📈 시각적 이상치 판별 (Scatter & Violin)", expanded=False):
            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
            if len(numeric_cols) < 1:
                st.warning("시각화할 수 있는 수치형 변수가 없습니다.")
            else:
                viz_tab1, viz_tab2 = st.tabs(["📍 산점도 (Scatter Plot)", "🎻 바이올린 플롯 (Violin Plot)"])
                
                # 라벨 맵핑 리스트
                col_labels = cb_parser.get_all_var_labels(numeric_cols) if cb_parser else numeric_cols

                with viz_tab1:
                    st.markdown("##### 두 변수 간의 관계와 극단값 확인")
                    c1, c2, c3 = st.columns([2, 2, 1])
                    with c1:
                        x_label = st.selectbox("X축 변수", options=col_labels, key="viz_x")
                        x_axis = cb_parser.get_column_from_label(x_label) if cb_parser else x_label
                    with c2:
                        y_label = st.selectbox("Y축 변수", options=col_labels, index=min(1, len(col_labels)-1), key="viz_y")
                        y_axis = cb_parser.get_column_from_label(y_label) if cb_parser else y_label
                    with c3:
                        dot_color = st.color_picker("점 색상", "#0F6CBD", key="viz_color")
                    
                    fig_scatter = px.scatter(df, x=x_axis, y=y_axis, template="plotly_white", 
                                           title=f"{x_label} vs {y_label} 산점도",
                                           labels={x_axis: x_label, y_axis: y_label})
                    fig_scatter.update_traces(marker=dict(color=dot_color, size=8, opacity=0.6))
                    st.plotly_chart(fig_scatter, use_container_width=True)
                
                with viz_tab2:
                    st.markdown("##### 데이터의 분포 밀도와 이상치 범위 확인")
                    v_label = st.selectbox("분석할 변수", options=col_labels, key="viz_v")
                    v_col = cb_parser.get_column_from_label(v_label) if cb_parser else v_label
                    fig_violin = px.violin(df, y=v_col, box=True, points="all", template="plotly_white",
                                         title=f"{v_label} 분포 분석 (Violin & Box)",
                                         labels={v_col: v_label})
                    fig_violin.update_traces(fillcolor="#0F6CBD", opacity=0.6, line=dict(color="black"))
                    st.plotly_chart(fig_violin, use_container_width=True)

        # [v4.6 추가] 데이터 이상치 사전 진단 (결측치와 동일한 테이블 형식)
        with st.expander("📊 데이터 이상치 사전 진단 (v4.6)", expanded=False):
            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
            outlier_stats = []
            for c in numeric_cols:
                m = df[c].mean()
                s = df[c].std()
                cnt = len(df[(df[c] < m - 3*s) | (df[c] > m + 3*s)])
                if cnt > 0:
                    outlier_stats.append({
                        "변수명": cb_parser.get_var_label(c) if cb_parser else c,
                        "원본ID": c,
                        "이상치건수(Z>3)": cnt,
                        "비중(%)": round(cnt / len(df) * 100, 1)
                    })
            
            if not outlier_stats:
                st.info("통계적 이상치(Z-score > 3)가 발견되지 않았습니다. ✨")
            else:
                outlier_df = pd.DataFrame(outlier_stats).sort_values("이상치건수(Z>3)", ascending=False)
                col_o1, col_o2 = st.columns([1, 1])
                with col_o1:
                    st.markdown("### 🔍 주요 이상 변수")
                    st.dataframe(outlier_df.drop(columns=["원본ID"]), hide_index=True, use_container_width=True)
                with col_o2:
                    st.markdown("### 🤖 AI 이상치 원인 추론")
                    if st.button("🧠 AI 이상치 진단 실행", key="diag_btn_out"):
                        diag_prompt = f"다음 변수들의 이상치 정보를 보고 단순 기입 오류 가능성인지, 실제 극단값 사례 가능성인지 추론해줘.\n{outlier_df.to_string()}"
                        res, err = run_analysis(
                            "{report_text}", 
                            diag_prompt, 
                            model_name=st.session_state["selected_model"], 
                            auto_mode=st.session_state["auto_mode"]
                        )
                        if not err: st.info(res)
                        else: st.error(f"진단 오류: {err}")

    # [v3.0 추가] 결측치 패턴 분석 섹션
    if mode == "imputation":
        with st.expander("📊 데이터 결측 패턴 분석 (v3.0)", expanded=False):
            missing_counts = df.isnull().sum()
            
            # [v4.6] 변수명 라벨링 적용
            m_labels = [cb_parser.get_var_label(c) if cb_parser else c for c in missing_counts.index]
            
            missing_df = pd.DataFrame({
                "변수명": m_labels,
                "원본ID": missing_counts.index,
                "결측건수": missing_counts.values,
                "결측비율(%)": (missing_counts.values / len(df) * 100).round(1)
            })
            missing_df = missing_df[missing_df["결측건수"] > 0].sort_values("결측건수", ascending=False)
            
            if missing_df.empty:
                st.info("현재 결측치가 있는 변수가 없습니다. 데이터가 완벽합니다! ✨")
            else:
                col_m1, col_m2 = st.columns([1, 1])
                with col_m1:
                    st.markdown("### 🔍 주요 결측 변수")
                    st.dataframe(missing_df.drop(columns=["원본ID"]), hide_index=True, use_container_width=True)
                with col_m2:
                    st.markdown("### 🤖 AI 결측 유형 진단")
                    if st.button("🧠 AI 패턴 진단 실행", key="diag_btn_ai"):
                        diag_prompt = f"다음 데이터의 결측 현황을 보고 MCAR, MAR, MNAR 중 유형을 진단해줘.\n{missing_df.to_string()}"
                        res, err = run_analysis(
                            "{report_text}", 
                            diag_prompt, 
                            model_name=st.session_state["selected_model"], 
                            auto_mode=st.session_state["auto_mode"]
                        )
                        if not err: st.info(res)
                        else: st.error(f"진단 오류: {err}")

    # 2. 변수 선택
    st.markdown('<div class="qx-section-label">2. 검토 대상 변수 선택</div>', unsafe_allow_html=True)
    all_col_labels = cb_parser.get_all_var_labels(df.columns) if cb_parser else df.columns.tolist()
    selected_labels = st.multiselect("검토가 필요한 변수를 선택하세요", options=all_col_labels, key=f"targets_{mode}")
    target_cols = [cb_parser.get_column_from_label(lb) for lb in selected_labels] if cb_parser else selected_labels

    if not target_cols:
        st.warning("분석할 변수를 최소 하나 이상 선택해 주세요.")
        return

    # 3. 변수별 보완 설정
    st.markdown('<div class="qx-section-label">3. 변수별 보완 방법 설정</div>', unsafe_allow_html=True)
    
    impute_configs = {}
    
    for i, col in enumerate(target_cols):
        display_name = selected_labels[i] if cb_parser else col
        with st.expander(f"📍 변수: {display_name}", expanded=True):
            col_a, col_b = st.columns([2, 1])
            
            with col_b:
                if st.button(f"🪄 AI 추천", key=f"ai_rec_{mode}_{col}"):
                    # AI 추천 로직
                    prompt = f"다음 변수의 데이터 대체 방법을 추천하고 이유를 설명해줘. 변수명: {display_name}, 타입: {df[col].dtype}, 샘플: {df[col].dropna().head(5).tolist()}"
                    res, err = run_analysis(
                        "{report_text}", 
                        prompt, 
                        model_name=st.session_state["selected_model"], 
                        auto_mode=st.session_state["auto_mode"]
                    )
                    if not err: st.session_state[f"rec_{mode}_{col}"] = res
                    else: st.session_state[f"rec_{mode}_{col}"] = f"AI 추천 생성 중 오류: {err}"
                
                if f"rec_{mode}_{col}" in st.session_state:
                    st.caption(st.session_state[f"rec_{mode}_{col}"])

            with col_a:
                methods = ["전체 평균 대체", "조건부 평균 대체", "랜덤 대체", "중앙값 대체", "최빈값 대체", "층별 평균 대체", "MICE 다중 대체", "k-NN 대체", "재확인(Call Back)", "직접 입력"]
                selected_method = st.selectbox(f"보완 방법 선택 ({col})", options=methods, key=f"method_{mode}_{col}")
                
                options = {}
                # [v4.7] 상관관계/연관성 도움말 및 변수 추천 (명목형 지원 추가)
                other_cols = [c for c in df.columns if c != col]
                if other_cols:
                    with st.spinner("연관성 분석 중..."):
                        associations = []
                        for c in other_cols:
                            assoc_val = get_association(df, col, c)
                            associations.append((c, assoc_val))
                        
                        associations.sort(key=lambda x: x[1], reverse=True)
                        top_associations = associations[:5]
                        
                        if top_associations:
                            assoc_items = []
                            for c, val in top_associations:
                                label = cb_parser.get_var_label(c) if cb_parser else c
                                # 타입에 따라 지표 명칭 변경
                                is_num1 = pd.api.types.is_numeric_dtype(df[col])
                                is_num2 = pd.api.types.is_numeric_dtype(df[c])
                                
                                metric_name = "r" # Pearson
                                if not is_num1 and not is_num2: metric_name = "V" # Cramers V
                                elif is_num1 != is_num2: metric_name = "η" # Correlation Ratio
                                
                                assoc_items.append(f"{label}({metric_name}={val:.2f})")
                            
                            st.caption(f"💡 **연관성 상위 변수 (추천 기준)**: {', '.join(assoc_items)}")

                if selected_method == "층별 평균 대체":
                    # [v4.6] 층별 변수도 라벨링 적용
                    st_labels = cb_parser.get_all_var_labels([c for c in df.columns if c != col]) if cb_parser else [c for c in df.columns if c != col]
                    sel_st_labels = st.multiselect(f"층(Strata) 변수 선택 ({display_name})", options=st_labels, key=f"strata_{mode}_{col}")
                    options["strata"] = [cb_parser.get_column_from_label(lb) for lb in sel_st_labels] if cb_parser else sel_st_labels
                elif selected_method == "k-NN 대체":
                    k_val = st.slider(f"k값 설정 ({display_name})", 1, 10, 5, key=f"k_{mode}_{col}")
                    options["k"] = k_val
                    # [v4.7] k-NN 기준 변수(Donors) 선택 - 명목형(명목/서열)도 포함
                    # 내부에서 원핫 인코딩 처리하므로 선택 가능
                    donor_labels = cb_parser.get_all_var_labels(other_cols) if cb_parser else other_cols
                    sel_donor_labels = st.multiselect(f"유사도 측정 기준 변수 선택 (미선택 시 전체 변수 사용)", options=donor_labels, key=f"donors_{mode}_{col}")
                    options["donors"] = [cb_parser.get_column_from_label(lb) for lb in sel_donor_labels] if cb_parser else sel_donor_labels
                elif selected_method == "재확인(Call Back)":
                    st.warning("⚠️ 이 데이터는 보완하지 않고 '재조사 명단'에 포함합니다.")
                
                impute_configs[col] = {"method": selected_method, "options": options}

    # 4. 실행 버튼
    st.markdown("<hr>", unsafe_allow_html=True)
    if st.button("🚀 데이터 보완 실행", type="primary", use_container_width=True, key=f"run_btn_{mode}"):
        imputer = DataImputer(df)
        
        with st.spinner("통계적 알고리즘 처리 중..."):
            for col, config in impute_configs.items():
                if mode == "imputation":
                    missing_idx = df[df[col].isna()].index.tolist()
                else:
                    if pd.api.types.is_numeric_dtype(df[col]):
                        m = df[col].mean()
                        s = df[col].std()
                        missing_idx = df[(df[col] < m - 3*s) | (df[col] > m + 3*s) | df[col].isna()].index.tolist()
                    else:
                        missing_idx = df[df[col].isna()].index.tolist()
                
                if not missing_idx: continue
                
                method = config["method"]
                opts = config["options"]
                
                if method == "전체 평균 대체": imputer.impute_grand_mean(col, missing_idx)
                elif method == "조건부 평균 대체": imputer.impute_conditional_mean(col, missing_idx)
                elif method == "랜덤 대체": imputer.impute_random(col, missing_idx)
                elif method == "중앙값 대체": imputer.impute_median(col, missing_idx)
                elif method == "최빈값 대체": imputer.impute_mode(col, missing_idx)
                elif method == "층별 평균 대체" and opts.get("strata"): imputer.impute_stratified_mean(col, missing_idx, opts["strata"])
                elif method == "k-NN 대체": imputer.impute_knn(col, missing_idx, k=opts.get("k", 5), donor_columns=opts.get("donors"))
                elif method == "MICE 다중 대체": imputer.impute_mice(col, missing_idx)
                elif method == "재확인(Call Back)": imputer._apply_imputation(col, missing_idx, "CALL_BACK", "재확인 대상분류")
                else: imputer.impute_grand_mean(col, missing_idx)
            
            st.session_state[f"imputed_df_{mode}"] = imputer.df
            st.session_state[f"impute_summary_{mode}"] = imputer.get_summary()
            st.session_state[f"impute_log_{mode}"] = imputer.audit_log
            
        st.success("데이터 보완 처리가 완료되었습니다!")

    # 5. 결과 확인 및 다운로드
    if f"imputed_df_{mode}" in st.session_state:
        show_security_notice()
        st.markdown('<div class="qx-section-label">4. 결과 요약 및 다운로드</div>', unsafe_allow_html=True)
        
        summary = st.session_state[f"impute_summary_{mode}"]
        if isinstance(summary, dict) and summary:
            cols_metric = st.columns(min(len(summary), 4))
            for i, (col_name, count) in enumerate(summary.items()):
                idx = i % 4
                display_name = cb_parser.get_var_label(col_name) if cb_parser else col_name
                cols_metric[idx].metric(display_name, f"{count}건")
        
        orig_df = df.copy()
        adj_df = st.session_state[f"imputed_df_{mode}"]
        log_list = st.session_state[f"impute_log_{mode}"]
        log_df = pd.DataFrame(log_list)

        # [v4.6] 로그 프레임에 라벨 추가
        if not log_df.empty:
            log_df["변수설명"] = log_df["변수명"].apply(lambda x: cb_parser.get_var_label(x) if cb_parser else x)
            # 코드값 라벨링 (기존값, 대체값)
            if cb_parser:
                log_df["기존라벨"] = log_df.apply(lambda r: cb_parser.get_value_label(r["변수명"], r["기존값"]), axis=1)
                log_df["대체라벨"] = log_df.apply(lambda r: cb_parser.get_value_label(r["변수명"], r["대체값"]), axis=1)
            
            # 컬럼 순서 조정
            cols_order = ["인덱스", "변수명", "변수설명", "기존값"]
            if "기존라벨" in log_df.columns: cols_order.append("기존라벨")
            cols_order.extend(["대체값"])
            if "대체라벨" in log_df.columns: cols_order.append("대체라벨")
            cols_order.append("적용방법")
            log_df = log_df[cols_order]

        # 결과 엑셀용 DF 구성
        export_df = orig_df.copy()
        for col in target_cols:
            if col in adj_df.columns:
                export_df[f"{col}_보완"] = adj_df[col]
                method_map = {row['인덱스']: row['적용방법'] for row in log_list if row['변수명'] == col}
                export_df[f"{col}_보완방법"] = export_df.index.map(lambda x: method_map.get(x, ""))
            else:
                st.warning(f"변수 '{col}'이(가) 보완 데이터프레임에 존재하지 않아 결과 리포트에서 제외되었습니다. 검토 변수를 다시 선택하거나 보완을 재실행해 주세요.")

        # [v4.9] 리포트 및 Raw Data 분리 다운로드
        report_output = io.BytesIO()
        with pd.ExcelWriter(report_output, engine='xlsxwriter') as writer:
            export_df.to_excel(writer, index=False, sheet_name='Audit_Report')
            if not log_df.empty: 
                log_df.to_excel(writer, index=False, sheet_name='Audit_Log')
        report_output.seek(0)
        
        raw_output = io.BytesIO()
        with pd.ExcelWriter(raw_output, engine='xlsxwriter') as writer:
            adj_df.to_excel(writer, index=False, sheet_name='Raw_Data')
        raw_output.seek(0)

        col_exp1, col_exp2 = st.columns(2)
        with col_exp1:
            st.download_button("📝 감사 리포트 다운로드 (Excel)", data=report_output, file_name=f"보완리포트_{df_file.name}", 
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key=f"dl_report_{mode}")
        with col_exp2:
            st.download_button("📥 최종 분석용 Raw Data 다운로드", data=raw_output, file_name=f"최종데이터_{df_file.name}", 
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key=f"dl_raw_{mode}")
        
        with st.expander("📝 상세 보완 내역 (Log)", expanded=True):
            st.dataframe(log_df, use_container_width=True, hide_index=True)
            
        callback_df = log_df[log_df["적용방법"] == "재확인 대상분류"]
        if not callback_df.empty:
            st.markdown('<div class="qx-section-label" style="color:#d32f2f;">🚨 재조사(Call-back) 필요 명단</div>', unsafe_allow_html=True)
            st.error(f"총 {len(callback_df)}건의 데이터가 재확인 대상으로 분류되었습니다.")
            
            target_callback_labels = callback_df["변수설명"].unique().tolist()
            if st.button("🎙️ AI 재조사 질문 가이드 생성", key="btn_callback_guide"):
                guide_prompt = f"다음 문항들에 대해 전화 재조사를 실시할 때의 스크립트를 작성해줘.\n문항: {', '.join(target_callback_labels)}"
                res, err = run_analysis(
                    "{report_text}", 
                    guide_prompt, 
                    model_name=st.session_state["selected_model"], 
                    auto_mode=st.session_state["auto_mode"]
                )
                if not err: st.info(res)
            
            st.dataframe(callback_df, use_container_width=True, hide_index=True)


# ── 사이드바
with st.sidebar:
    # 로고 추가 (가운데 정렬)
    try:
        _, mid_col, _ = st.columns([1, 3, 1])
        with mid_col:
            st.image("logo.png", width=180)
    except:
        pass
    st.markdown("<h3 style='text-align: center; color: white;'>AI 과업 관리 솔루션</h3>", unsafe_allow_html=True)
    st.markdown("<hr>", unsafe_allow_html=True)

    # 내비게이션
    st.markdown('<div class="qx-section-label">NAVIGATION</div>', unsafe_allow_html=True)
    menu_options = [
        "과업내용 체크 (RFP 분석)", 
        "AI 설문지 최적화",
        "AI 표본설계",
        "AI 이상치 검토 (Call Back, Data Adjustment)", 
        "AI 결측치 검토 (Call Back, Imputation)",
        "AI 단위 무응답 검토",
        "기업체 일반 현황 행정자료 비교",
        "사업체 명부 추출",
        "보고서 검수 AI Tools"
    ]
    
    # [v4.14] 화이트리스트 기반 권한 제어 - 'AI 단위 무응답 검토' 섹션
    # (shjeon, metrix11 아이디만 접근 허용, 비로그인 시 숨김)
    allowed_for_nonresponse = ["shjeon", "metrix11"]
    current_user = st.session_state.get("logged_in_user", "")
    
    if current_user not in allowed_for_nonresponse:
        if "AI 단위 무응답 검토" in menu_options:
            menu_options.remove("AI 단위 무응답 검토")
        if "기업체 일반 현황 행정자료 비교" in menu_options:
            menu_options.remove("기업체 일반 현황 행정자료 비교")
    
    # [v4.12] 관리자(shjeon) 전용 메뉴 추가
    if st.session_state.get("logged_in_user") == "shjeon":
        menu_options.append("⚙️ ADMIN DASHBOARD")
    
    # 세션 상태에 저장된 메뉴가 옵션에 없으면(예: 권한으로 숨겨짐) 기본값(첫 번째) 사용
    if "menu_selection" not in st.session_state or st.session_state["menu_selection"] not in menu_options:
        st.session_state["menu_selection"] = menu_options[0]
        
    try:
        current_idx = menu_options.index(st.session_state["menu_selection"])
    except (ValueError, KeyError):
        current_idx = 0

    menu = st.radio(
        "메뉴를 선택하세요",
        menu_options,
        index=current_idx,
        label_visibility="collapsed",
    )
    
    # [v4.12] 섹션 접근 로깅 및 상태 동기화
    if st.session_state["menu_selection"] != menu:
        st.session_state["menu_selection"] = menu
        tracker.log_event(st.session_state.get("logged_in_user", "unknown"), "ACCESS", menu)
        st.rerun()
    
    st.markdown("<hr>", unsafe_allow_html=True)

    if not st.session_state["is_logged_in"]:
        st.markdown('<div class="qx-section-label">LOGIN</div>', unsafe_allow_html=True)
        login_id = st.text_input(
            "ID",
            placeholder="아이디를 입력하세요",
            label_visibility="collapsed",
            key="login_input",
        )
        if st.button("로그인", use_container_width=True, key="btn_login"):
            do_login(login_id)
            st.rerun()
        if st.session_state["login_error"]:
            st.error(st.session_state["login_error"], icon="⚠️")
        st.markdown("<hr>", unsafe_allow_html=True)
    else:
        st.markdown(
            '<div class="sb-file-card" style="margin-bottom:0.5rem;">'
            '<div class="sb-file-name">&#128100; ' + st.session_state["logged_in_user"] + '</div>'
            '<div class="sb-file-meta">로그인됨</div></div>',
            unsafe_allow_html=True,
        )
        if st.button("로그아웃", use_container_width=True, key="btn_logout"):
            do_logout()
            st.rerun()
        st.markdown("<hr>", unsafe_allow_html=True)

    # [v4.17] 관리자 전용 시스템 도구 (shjeon에게만 노출)
    if st.session_state.get("logged_in_user") == "shjeon":
        # API 키 상태
        _keys = get_api_keys()
        if _keys:
            st.success(f"API 키 {len(_keys)}개 연결됨", icon="✅")
            st.caption(f"요청마다 {len(_keys)}개 키 중 랜덤 선택")
        else:
            st.error("API 키 없음 — Streamlit Secrets에 GEMINI_API_KEYS를 추가하세요.", icon="⚠️")

        st.markdown("<hr>", unsafe_allow_html=True)
        
        # API 할당량 진단 (v2.9)
        st.markdown('<div class="qx-section-label">진단 도구</div>', unsafe_allow_html=True)
        if st.button("🔍 API 할당량 진단 시작", use_container_width=True, key="btn_diag"):
            with st.expander("진단 결과", expanded=True):
                diag_prog = st.progress(0, text="진단 대기 중...")
                diag_status = st.empty()
                
                def diag_callback(curr, total, msg):
                    diag_prog.progress(curr/total, text=f"진단 중... ({curr}/{total})")
                    diag_status.caption(msg)
                    
                with st.spinner("모든 키와 모델 상태를 확인 중입니다..."):
                    results = check_key_quotas(progress_callback=diag_callback)
                    diag_prog.empty()
                    diag_status.empty()
                    
                    if results:
                        import pandas as pd
                        df = pd.DataFrame(results)
                        st.dataframe(
                            df, 
                            hide_index=True,
                            column_config={
                                "상태": st.column_config.TextColumn("상태", width="medium"),
                                "상세 내용": st.column_config.TextColumn("상세 내용", width="large"),
                            },
                            use_container_width=True
                        )
                        
                        ok_count = sum(1 for r in results if "정상" in r["상태"])
                        
                        # INFO 행에서 발견된 모델 목록 추출 및 표시
                        info_row = next((r for r in results if r["순번"] == "INFO"), None)
                        if info_row:
                            st.info(f"🔍 **발견된 가용 모델 ID 목록:**\n\n`{info_row['상세 내용']}`")
                            candidates = [m for m in info_row['상세 내용'].split(", ") if "flash" in m or "pro" in m]
                            st.caption(f"이 중 분석 도구에서 활용 가능한 모델: {', '.join(candidates)}")

                        st.success(f"진단 완료: 총 {len(results)-1 if info_row else len(results)}개 조합 중 {ok_count}개 정상")
                        if ok_count == 0:
                            st.error("사용 가능한 키/모델 조합이 없습니다. API 키를 교체해 주세요.")
                    else:
                        st.warning("진단할 API 키가 없습니다.")
        
        st.markdown("<hr>", unsafe_allow_html=True)

    # 모델 선택
    st.markdown('<div class="qx-section-label">AI 모델 설정</div>', unsafe_allow_html=True)
    auto_mode = st.toggle(
        "🤖 자동 최적화 (권장)",
        value=st.session_state["auto_mode"],
        help="할당량 초과 시 최적 모델로 자동 전환합니다. 우선순위: Gemini 2.5 Flash → 2.0 Flash → ...",
    )
    st.session_state["auto_mode"] = auto_mode

    if auto_mode:
        priority_names = " → ".join(
            MODEL_DISPLAY_NAMES.get(m, m).replace("Gemini ", "") for m in AUTO_MODEL_PRIORITY[:4]
        ) + " → ..."
        st.caption(f"우선순위: {priority_names}")
    else:
        # 수동 선택 영역을 테두리가 있는 컨테이너로 묶음
        with st.container(border=True):
            # 표시 이름 목록 생성
            display_options = [MODEL_DISPLAY_NAMES.get(m, m) for m in AVAILABLE_MODELS]
            display_default = MODEL_DISPLAY_NAMES.get(DEFAULT_MODEL, DEFAULT_MODEL)
            
            selected_display = st.selectbox(
                "모델을 직접 선택하세요",
                display_options,
                index=display_options.index(display_default),
                label_visibility="visible",
                key="sb_model_select"
            )
            
            # 표시 이름 → 실제 모델 ID 역매핑
            reverse_map = {v: k for k, v in MODEL_DISPLAY_NAMES.items()}
            selected_model = reverse_map.get(selected_display, DEFAULT_MODEL)
            st.session_state["selected_model"] = selected_model
            
            # 모델 ID 가시성 극대화 (회색톤 + 화이트 텍스트)
            st.markdown(f"""
                <div style="display: flex; align-items: center; gap: 8px; margin-top: 5px;">
                    <div style="
                        background: #4A5568; 
                        color: #FFFFFF; 
                        padding: 3px 10px; 
                        border-radius: 4px; 
                        font-family: 'Roboto Mono', monospace; 
                        font-size: 0.75rem;
                        font-weight: 700;
                    ">
                        {selected_model}
                    </div>
                    <span style="font-size: 0.7rem; color: #A0AABB;">Active ID</span>
                </div>
            """, unsafe_allow_html=True)


    st.markdown("<hr>", unsafe_allow_html=True)

    # 파일 현황
    st.markdown('<div class="qx-section-label">파일 업로드 현황</div>', unsafe_allow_html=True)
    if st.session_state["file_name"]:
        pages_info = f" · {st.session_state['file_pages']}p" if st.session_state["file_pages"] else ""
        st.markdown(
            '<div class="sb-file-card">'
            '<div class="sb-file-name">&#128196; ' + st.session_state["file_name"] + '</div>'
            '<div class="sb-file-meta">' + f"{len(st.session_state['report_text']):,}자{pages_info}" + '</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        step_done = sum(1 for v in st.session_state["step_results"].values() if v)
        st.progress(step_done / 3, text=f"단계별 분석 {step_done}/3 완료")
    else:
        st.caption("파일을 업로드하면 현황이 표시됩니다.")

    st.markdown("<hr>", unsafe_allow_html=True)

    if st.button("초기화", use_container_width=True):
        for k in ["report_text", "file_name", "full_result", "rfp_curr_text", "rfp_prev_text", "rfp_project_name"]:
            st.session_state[k] = ""
        st.session_state["step_results"] = {1: "", 2: "", 3: ""}
        st.session_state["rfp_results"] = {}
        st.session_state["file_pages"] = 0
        # 기업체 자료 수집 세션 초기화
        for k in ["biz_crawl_result", "biz_crawl_stats", "biz_crawl_selected_nps", "biz_crawl_selected_nhis", "biz_cleaned_df", "biz_clean_stats", "biz_nps_dataset", "biz_nhis_dataset"]:
            if k in st.session_state:
                del st.session_state[k]
        st.rerun()

    st.markdown("<hr>", unsafe_allow_html=True)
    st.caption("개발: ㅈㅅㅎ")
    st.caption("문의: shjeon1@metrix.co.kr")
    st.caption("Powered by Google Gemini · v2.9")

    # [v4.16] 무단 사용 금지 경고 문구 추가 (아이콘/띠 제외)
    st.markdown("""
    <div style="
        background-color: rgba(255, 245, 245, 0.1); 
        padding: 10px; 
        margin-top: 20px; 
        border-radius: 4px;
    ">
        <p style="
            color: #FC8181; 
            font-size: 0.72rem; 
            font-weight: 500; 
            line-height: 1.5; 
            margin: 0;
        ">
            본 솔루션은 회사의 지적 재산이며, 사전 허가 없는 무단 사용 및 배포는 엄격히 금지됩니다.
        </p>
    </div>
    """, unsafe_allow_html=True)



def export_to_docx(markdown_text: str) -> io.BytesIO:
    """마크다운-텍스트를 워드(DOCX) 파일로 변환 (표 객체 및 빨간색 강조 지원, 한국어 폰트 지원)"""
    if Document is None:
        return io.BytesIO(b"python-docx is not installed")
    
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    import re
    
    KOREAN_FONT = '맑은 고딕'
    
    doc = Document()
    
    # ── 기본(Normal) 스타일에 한국어 폰트 설정
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = KOREAN_FONT
    # East-Asian 폰트 명시 (Word가 한글을 해당 폰트로 렌더링하도록 보장)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), KOREAN_FONT)
    
    # ── 제목(Heading) 스타일에도 한국어 폰트 설정
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = KOREAN_FONT
        if heading_style.element.rPr is None:
            heading_style.element.get_or_add_rPr()
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), KOREAN_FONT)
    
    def _set_run_font(run):
        """개별 run에 한국어 폰트를 적용하는 헬퍼"""
        run.font.name = KOREAN_FONT
        run.element.rPr.rFonts.set(qn('w:eastAsia'), KOREAN_FONT)
    
    lines = markdown_text.split('\n')
    table_buffer = []
    
    def flush_table():
        if not table_buffer:
            return
        # 컬럼 수 결정 (가장 긴 행 기준)
        max_cols = max(len(row) for row in table_buffer)
        if max_cols == 0:
            return
        
        table = doc.add_table(rows=0, cols=max_cols)
        table.style = 'Table Grid'
        
        for row_data in table_buffer:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                if i < max_cols:
                    # 셀 내부에도 빨간색 강조 적용 시도
                    p = row_cells[i].paragraphs[0]
                    # 셀 내부에도 볼드 및 빨간색 강조 적용
                    add_formatted_text(p, cell_text)
        table_buffer.clear()
        doc.add_paragraph() # 표 뒤에 공백 추가

    def add_formatted_text(paragraph, text):
        """정규표현식을 사용하여 마크다운 볼드(**), span 태그, blue 태그 처리 (한국어 폰트 적용)"""
        # <br> 태그를 실제 줄바꿈으로 변환 (워드용)
        text = text.replace("<br>", "\n")
        
        # 1. 빨간색 강조(<span...>), 파란색 강조(<blue>), 볼드(**...**)를 식별하기 위한 정규식
        pattern = r"(<span style='color:red'>.*?</span>|<blue>.*?</blue>|\*\*.*?\*\*)"
        parts = re.split(pattern, text)
        
        for part in parts:
            if part.startswith("<span style='color:red'>"):
                inner = re.sub(r"<span style='color:red'>(.*?)</span>", r"\1", part)
                inner = inner.replace("**", "")
                run = paragraph.add_run(inner)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run.bold = True
                _set_run_font(run)
            elif part.startswith("<blue>"):
                inner = re.sub(r"<blue>(.*?)</blue>", r"\1", part)
                inner = inner.replace("**", "")
                run = paragraph.add_run(inner)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                run.bold = True
                _set_run_font(run)
            elif part.startswith("**") and part.endswith("**"):
                inner = part[2:-2]
                run = paragraph.add_run(inner)
                run.bold = True
                _set_run_font(run)
            else:
                if part:
                    run = paragraph.add_run(part)
                    _set_run_font(run)

    for line in lines:
        line_strip = line.strip()
        
        # 표 행 감지 (| 로 시작하거나 포함된 경우)
        if '|' in line_strip:
            if re.match(r'^[|\s\-:]+$', line_strip):
                continue
            parts = [p.strip() for p in line_strip.split('|') if p.strip()]
            if parts:
                table_buffer.append(parts)
                continue
        
        if table_buffer:
            flush_table()
        
        if not line_strip:
            continue
        
        if line_strip.startswith('###'):
            doc.add_heading(line_strip.lstrip('#').strip(), level=3)
        elif line_strip.startswith('##'):
            doc.add_heading(line_strip.lstrip('#').strip(), level=2)
        elif line_strip.startswith('#'):
            doc.add_heading(line_strip.lstrip('#').strip(), level=1)
        else:
            is_bullet = line_strip.startswith('- ') or line_strip.startswith('* ')
            text_content = line_strip[2:] if is_bullet else line_strip
            
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph()
            
            add_formatted_text(p, text_content)
    
    flush_table()
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ── 로그인 가드
if not st.session_state["is_logged_in"]:
    st.markdown("""
<div class="qx-topbar">
    <span class="qx-topbar-logo">보고서 검수 AI Tools</span>
    <span class="qx-topbar-sep"></span>
    <span class="qx-topbar-title">수석 리서치 품질 검수관</span>
</div>
""", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""
<div class="qx-card" style="max-width:480px; margin:0 auto; text-align:center; padding:3rem 2rem;">
    <div style="font-size:2.5rem; margin-bottom:1rem;">&#128272;</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">로그인이 필요합니다</div>
    <div style="font-size:0.85rem; color:#8B96A9;">좌측 사이드바에서 아이디를 입력하고 로그인 버튼을 클릭하세요.</div>
</div>
""", unsafe_allow_html=True)

else:
    if st.session_state["menu_selection"] == "보고서 검수 AI Tools":
        # 상단 헤더 바
        st.markdown("""
    <div class="qx-topbar">
        <span class="qx-topbar-logo">보고서 검수 AI Tools</span>
        <span class="qx-topbar-sep"></span>
        <span class="qx-topbar-title">수석 리서치 품질 검수관</span>
        <span class="qx-topbar-badge">AI-Powered Quality Check</span>
    </div>
    """, unsafe_allow_html=True)
    elif st.session_state["menu_selection"] == "과업내용 체크 (RFP 분석)":
        show_win_strategy_section()
        st.stop()
    elif st.session_state["menu_selection"] == "기업체 일반 현황 행정자료 비교":
        show_business_info_crawling()
        st.stop()
    elif st.session_state["menu_selection"] == "사업체 명부 추출":
        show_unified_business_search()
        st.stop()
    elif st.session_state["menu_selection"] == "AI 설문지 최적화":
        show_questionnaire_optimization_system()
        st.stop()
    elif st.session_state["menu_selection"] == "AI 표본설계":
        show_sample_design_system()
        st.stop()
    elif st.session_state["menu_selection"] == "AI 이상치 검토 (Call Back, Data Adjustment)":
        show_outlier_inspection_system(mode="outlier")
        st.stop()
    elif st.session_state["menu_selection"] == "AI 결측치 검토 (Call Back, Imputation)":
        show_outlier_inspection_system(mode="imputation")
        st.stop()
    elif st.session_state["menu_selection"] == "AI 단위 무응답 검토":
        show_unit_nonresponse_system()
        st.stop()
    elif st.session_state["menu_selection"] == "⚙️ ADMIN DASHBOARD":
        if st.session_state.get("logged_in_user") == "shjeon":
            show_admin_dashboard()
        else:
            st.error("이 페이지에 대한 접근 권한이 없습니다.")
        st.stop()

    # [v4.2 추가] 메인 보고서 검수 가이드
    with st.expander("📘 AI 보고서 검수 서비스 이용 안내", expanded=False):
        st.markdown("""
        ### 🛠️ 이용 방법 안내
        1. **보고서 업로드:** 검수할 리서치 보고서(PDF, DOCX, TXT)를 업로드 영역에 드래그하거나 선택합니다.
        2. **분석 단계 수행:** 하단의 **STEP 1(요약)**, **STEP 2(검수)**, **STEP 3(종합)** 버튼을 순서대로 클릭하여 AI 분석을 생성합니다.
        3. **내용 확인 및 수정:** AI가 생성한 초안을 검토하고 필요한 경우 텍스트 상자에서 직접 수정합니다.
        4. **결과 다운로드:** 모든 분석이 완료되면 상단의 **'검수 결과 다운로드(Word)'** 버튼을 통해 최종 보고서를 저장합니다.
        """)

    # 파일 업로드 + 검수 항목
    col_up, col_items = st.columns([3, 2], gap="large")

    with col_up:
        st.markdown('<div class="qx-section-label">UPLOAD REPORT (보고서 분석용)</div>', unsafe_allow_html=True)
        st.markdown("""
<div class="qx-upload-zone">
    <div class="qx-upload-icon">&#128203;</div>
    <div class="qx-upload-text">보고서 분석용 (PDF · DOCX · TXT) 파일 업로드</div>
    <div class="qx-upload-hint">엑셀 파일은 '기업체 일반 현황...' 메뉴를 이용해 주세요</div>
</div>
""", unsafe_allow_html=True)
        uploaded_file = st.file_uploader(
            "보고서 파일 선택 (PDF, DOCX, TXT)",
            type=["pdf", "docx", "txt"],
            label_visibility="collapsed",
        )

    with col_items:
        st.markdown('<div class="qx-section-label">REVIEW ITEMS</div>', unsafe_allow_html=True)
        st.markdown("""
<div class="qx-card" style="padding:1.25rem 1.5rem;">
    <div style="display:flex;flex-direction:column;gap:0.65rem;">
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">01</span><span>조사 설계 및 요약</span>
        </div>
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">02</span><span>부문별 정밀 검수 및 오류 식별</span>
        </div>
        <div style="display:flex;align-items:center;gap:0.6rem;font-size:0.875rem;color:#2D3A50;">
            <span style="color:#0F6CBD;font-weight:700;">03</span><span>종합 검수 보고서</span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

    # 파일 처리
    if uploaded_file is not None:
        if uploaded_file.name != st.session_state.get("file_name", ""):
            with st.spinner("파일 텍스트 추출 중..."):
                text, pages = extract_text(uploaded_file)
                if text:
                    truncated = truncate_text(text, MAX_TEXT_CHARS)
                    st.session_state["report_text"] = truncated
                    st.session_state["file_name"] = uploaded_file.name
                    st.session_state["file_pages"] = pages
                    st.session_state["step_results"] = {1: "", 2: "", 3: ""}
                    st.session_state["full_result"] = ""
                    if len(text) > MAX_TEXT_CHARS:
                        st.warning(
                            f"텍스트가 너무 깁니다 ({len(text):,}자). "
                            f"{MAX_TEXT_CHARS:,}자로 균형 있게 잘라 분석합니다."
                        )
                    else:
                        pages_str = f", {pages}페이지" if pages else ""
                        st.success(
                            f"파일 로드 완료 — **{uploaded_file.name}** ({len(text):,}자{pages_str})"
                        )
                else:
                    st.error("파일에서 텍스트를 추출할 수 없습니다.")

    # 분석 영역
    if st.session_state["report_text"]:
        st.markdown("<hr>", unsafe_allow_html=True)

        col_btn, col_hint = st.columns([1, 3])
        with col_btn:
            run_full = st.button("⚡ 전체 3단계 병렬 분석", use_container_width=True, type="primary")
        with col_hint:
            st.markdown(
                '<span style="font-size:0.82rem;color:#8B96A9;">'
                '1~3단계를 <b>동시에</b> 병렬 실행합니다. 순차 실행 대비 최대 3배 빠릅니다.'
                '</span>',
                unsafe_allow_html=True,
            )

        if run_full:
            api_key = get_api_key()
            if not api_key:
                st.error("API 키가 설정되지 않았습니다. Streamlit Secrets에 GEMINI_API_KEYS를 추가하세요.")
            else:
                _auto = st.session_state["auto_mode"]
                _model = st.session_state["selected_model"]
                st.session_state["full_result"] = ""
                st.session_state["step_results"] = {1: "", 2: "", 3: ""}

                progress_bar = st.progress(0, text="⚡ 3단계 병렬 분석 실시간 진행 중...")
                status_cols = st.columns(3)
                placeholders = {i+1: status_cols[i].empty() for i in range(3)}
                
                # 초기 상태 표시
                step_names = [
                    "조사 설계 및 요약", 
                    "부문별 정밀 검수 및 오류 식별", 
                    "종합 검수 보고서"
                ]
                for i, name in enumerate(step_names, 1):
                    placeholders[i].markdown(f"**[{i}단계]** {name}...")
                
                st.info("🚀 분석을 시작합니다. 잠시만 기다려 주세요 (약 2~3분 소요)")

                # 실시간 순차 처리 (안정성 및 사용자 피드백 최대화)
                # 데드락 방지를 위해 상위 레벨에서는 순차 실행, 내부(analyzer)에서는 청크별 병렬 실행
                from analyzer import run_step_with_chunks, STEP_PROMPTS
                from file_processor import slice_text_for_step
                from config import STEP_TEXT_RATIO
                
                log_area = st.empty()
                log_content = []
                timer_area = st.empty()
                start_time = datetime.datetime.now()

                def update_logs(msg):
                    import datetime
                    # KST (UTC+9) 설정
                    tz_kst = datetime.timezone(datetime.timedelta(hours=9))
                    now = datetime.datetime.now(tz_kst).strftime("%H:%M:%S")
                    log_content.append(f"[{now}] {msg}")
                    # 서버 콘솔 로깅 (Streamlit Cloud Logs)
                    print(f"[UI LOG] [{now}] {msg}", flush=True)
                    # 최근 5개 로그만 표시
                    display_logs = "\n".join(log_content[-5:])
                    log_area.code(display_logs, language="text")
                    
                    # 타이머 업데이트
                    elapsed = datetime.datetime.now() - start_time
                    timer_area.info(f"⏱️ 분석 시작 후 **{elapsed.seconds}초** 경과 중... (엔진 정상 가동 중)")

                update_logs("🚀 전수 검수 시스템 가동 (병렬 가속 모드 v2.7)")
                
                completed_count = 0
                for s in range(1, 4):
                    update_logs(f"➡ {s}단계 분석 시작: {STEP_LABELS[s]}")
                    
                    # 개별 단계 내 청크 진행 상황을 표시하기 위한 콜백 함수
                    def make_callback(step_idx):
                        def callback(curr, total):
                            is_started = (curr < total)
                            sub_prog = curr / total
                            overall = (step_idx - 1) / 3 + (sub_prog / 3)
                            
                            status_msg = f"⏳ {step_idx}단계 분석 중... (조각 {curr+1 if is_started else curr}/{total})"
                            progress_bar.progress(min(overall, 0.99), text=status_msg)
                            
                            if is_started:
                                update_logs(f"ㄴ {step_idx}단계 조각 {curr+1}/{total} 분석 요청됨...")
                            else:
                                update_logs(f"✅ {step_idx}단계 모든 조각 분석 완료!")
                        return callback

                    # 개별 단계 실행 (v14.1: 3단계는 이전 결과물 요약 방식)
                    if s == 3:
                        # 1, 2단계 결과물 조합하여 3단계의 '분석 대상'으로 전달
                        input_text = f"--- [1단계 조사설계 요약] ---\n{st.session_state['step_results'][1]}\n\n"
                        input_text += f"--- [2단계 정밀검수 결과] ---\n{st.session_state['step_results'][2]}"
                        prompt_ready = STEP_PROMPTS[s]
                    else:
                        prompt_ready = STEP_PROMPTS[s]
                        input_text = slice_text_for_step(st.session_state["report_text"], s, STEP_TEXT_RATIO)

                    text, err = run_step_with_chunks(
                        s, 
                        input_text,
                        prompt_ready,
                        model_name=_model,
                        auto_mode=_auto,
                        progress_callback=make_callback(s)
                    )
                    
                    completed_count += 1
                    
                    if not text and err:
                        placeholders[s].error(f"❌ {s}단계 실패")
                        st.session_state["step_results"][s] = f"실패 사유: {err}"
                        update_logs(f"❌ {s}단계 분석 실패: {err}")
                    elif err:
                        # 일부 조각 실패 (부분 성공)
                        placeholders[s].warning(f"⚠️ {s}단계 부분 성공")
                        st.session_state["step_results"][s] = text
                        update_logs(f"⚠ {s}단계 일부 조각 누락됨 (분석은 계속 진행)")
                    else:
                        placeholders[s].success(f"✅ {s}단계 완료")
                        st.session_state["step_results"][s] = text
                        update_logs(f"✅ {s}단계 분석 성공")
                    
                    progress_bar.progress(completed_count / 3, text=f"⚡ {s}단계 완료! ({completed_count}/3 전체 진행)")

                # 결과 취합
                combined = ""
                for s in range(1, 4):
                    if st.session_state["step_results"][s]:
                        combined += f"\n\n---\n\n## {STEP_LABELS[s]}\n\n{st.session_state['step_results'][s]}"
                
                st.session_state["full_result"] = combined
                if combined:
                    st.success("⚡ 전체 병렬 분석이 완료되었습니다!")
                    st.rerun()
                else:
                    st.error("분석 결과 생성에 실패했습니다.")

        st.markdown("<hr>", unsafe_allow_html=True)

        # 단일 페이지 분석 UI (v16.0: 한 페이지에 모두 표시)
        if any(st.session_state["step_results"].values()) or st.session_state["full_result"]:
            show_security_notice()
            st.markdown("<hr>", unsafe_allow_html=True)
            
        for s in [1, 2, 3]:
            st.markdown(f'<div class="qx-section-label">{STEP_LABELS[s]}</div>', unsafe_allow_html=True)
            
            # 각 단계별 카드 섹션
            with st.container():
                col_hd, col_st = st.columns([3, 1])
                with col_hd:
                    st.markdown(f"### {STEP_LABELS[s]}")
                with col_st:
                    if st.session_state["step_results"][s]:
                        st.markdown('<span class="badge-ok">완료</span>', unsafe_allow_html=True)
                    else:
                        st.markdown('<span class="badge-warn">대기</span>', unsafe_allow_html=True)

                run_btn = st.button(
                    f"{s}단계 분석 실행",
                    key=f"btn_step_{s}",
                )
                
                result_area = st.empty()
                
                # 결과 표시
                if st.session_state["step_results"][s]:
                    result_area.markdown(st.session_state["step_results"][s], unsafe_allow_html=True)
                elif st.session_state["full_result"] and not st.session_state["step_results"][s]:
                    result_area.info(f"{s}단계의 통합 결과가 준비되어 있습니다.")

                if run_btn:
                    if not get_api_keys():
                        st.error("API 키가 설정되지 않았습니다.")
                    else:
                        _auto = st.session_state["auto_mode"]
                        _model = st.session_state["selected_model"]
                        with st.spinner(f"{s}단계 분석 중..."):
                            full_text = ""
                            # 스트리밍 방식 호출 (실시간 피드백)
                            for chunk, is_error, _ in run_analysis_stream(
                                STEP_PROMPTS[s],
                                st.session_state["report_text"],
                                model_name=_model,
                                auto_mode=_auto,
                            ):
                                if is_error:
                                    st.error(chunk)
                                    break
                                full_text += chunk
                                result_area.markdown(full_text + " ▌", unsafe_allow_html=True)
                            
                            if full_text:
                                result_area.markdown(full_text, unsafe_allow_html=True)
                                st.session_state["step_results"][s] = full_text
                                # 전체 결과 업데이트
                                combined = ""
                                for i in range(1, 4):
                                    if st.session_state["step_results"][i]:
                                        combined += f"\n\n---\n\n## {STEP_LABELS[i]}\n\n{st.session_state['step_results'][i]}"
                                st.session_state["full_result"] = combined
                                st.rerun()
            st.markdown("<br>", unsafe_allow_html=True)

        # 다운로드
        # 인쇄 버튼 (JS 기반)
        st.markdown("""
        <div class="no-print">
            <button onclick="window.print()" style="
                width: 100%;
                background-color: #0F6CBD;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 0.75rem;
                font-size: 0.9rem;
                font-weight: 600;
                cursor: pointer;
                margin-top: 1rem;
                margin-bottom: 1.5rem;
                display: flex;
                align-items: center;
                justify-content: center;
                gap: 8px;
            ">
                🖨️ 분석 결과 리포트 인쇄하기
            </button>
        </div>
        <style>
            @media print { .no-print { display: none !important; } }
        </style>
        """, unsafe_allow_html=True)

        # [v4.13 수정] 잘못된 조건부 호출 제거 (이 부분은 하단 다운로드 영역 전 공통 노출로 대체됨)
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="qx-section-label">DOWNLOAD RESULTS</div>', unsafe_allow_html=True)

        def build_full_report() -> str:
            fname = st.session_state.get("file_name", "보고서")
            parts = [f"# 수석 리서치 품질 검수 보고서\n\n**대상 파일:** {fname}\n\n---\n"]
            if st.session_state["full_result"]:
                parts.append("## 전체 3단계 종합 분석\n\n")
                parts.append(st.session_state["full_result"])
            else:
                for step in range(1, 4):
                    result = st.session_state["step_results"].get(step, "")
                    parts.append(f"\n\n## {STEP_LABELS[step]}\n\n")
                    parts.append(result if result else "*아직 분석이 실행되지 않았습니다.*")
            return "\n".join(parts)

        col_dl1, _ = st.columns([2, 2])
        report_md = build_full_report()
        base_name = st.session_state.get("file_name", "report").rsplit(".", 1)[0]
        # 과업명 기반 파일명 생성
        r_project = st.session_state.get("rfp_project_name", "").strip()
        if r_project:
            r_safe_name = r_project.replace("/", "_").replace("\\", "_").replace(":", "_")[:80]
            docx_filename = f"{r_safe_name}_검수보고서.docx"
        else:
            docx_filename = f"{base_name}_검수보고서.docx"

        with col_dl1:
            try:
                docx_file = export_to_docx(report_md)
                st.download_button(
                    label="📝 워드 파일로 다운로드 (.docx)",
                    data=docx_file,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"워드 변환 중 오류 발생: {e}")

    else:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("""
<div class="qx-card" style="text-align:center; padding:3.5rem 2rem;">
    <div style="font-size:3rem; margin-bottom:1rem;">&#128203;</div>
    <div style="font-size:1.1rem; font-weight:600; color:#1A2237; margin-bottom:0.5rem;">
        보고서 파일을 업로드하세요
    </div>
    <div style="font-size:0.87rem; color:#8B96A9; margin-bottom:2rem;">
        PDF, DOCX, TXT 형식의 결과보고서를 업로드하면 AI가 4단계 전문 검수를 자동으로 수행합니다.
    </div>
    <div style="display:flex; justify-content:center; gap:1.5rem; flex-wrap:wrap;">
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:1rem 1.5rem;min-width:130px;">
            <div style="font-size:1.4rem;">&#128203;</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.4rem;">조사 설계 요약</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:1rem 1.5rem;min-width:130px;">
            <div style="font-size:1.4rem;">&#128269;</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.4rem;">부문별 정밀 검수</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:1rem 1.5rem;min-width:130px;">
            <div style="font-size:1.4rem;">&#9888;&#65039;</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.4rem;">5대 오류 기준 점검</div>
        </div>
        <div style="background:#F4F6F9;border:1px solid #E5E9F0;border-radius:8px;padding:1rem 1.5rem;min-width:130px;">
            <div style="font-size:1.4rem;">&#128196;</div>
            <div style="font-size:0.75rem;font-weight:600;color:#3D4F6B;margin-top:0.4rem;">종합 검수 보고서</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# Last forced sync: 2026-02-22 07:10:00
