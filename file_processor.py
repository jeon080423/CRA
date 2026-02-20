"""
파일 텍스트 추출 모듈 (PDF, DOCX, TXT)
토큰 절약을 위한 텍스트 압축 전처리 포함
"""
import re
import io
import streamlit as st


def preprocess_text(text: str) -> str:
    """특수 공백 문자 정제 및 기본 전처리"""
    if not text:
        return ""
    # 비표준 공백 문자 → 일반 공백 치환
    text = re.sub(r'[\u00a0\u200b\u200c\u200d\u2060\ufeff\u3000]', ' ', text)
    # 연속 공백 정리
    text = re.sub(r' {3,}', '  ', text)
    # 연속 줄바꿈 정리 (3개 이상 → 2개)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def compress_text(text: str) -> str:
    """
    토큰 절약을 위한 텍스트 압축.
    분석 품질에 영향 없이 반복 공백/빈 줄을 제거.
    """
    if not text:
        return ""
    # 탭 → 공백 2칸
    text = text.replace('\t', '  ')
    # 줄 단위로 처리
    lines = text.splitlines()
    compressed = []
    prev_blank = False
    for line in lines:
        stripped = line.rstrip()
        is_blank = stripped == ''
        # 연속 빈 줄은 하나로
        if is_blank:
            if not prev_blank:
                compressed.append('')
            prev_blank = True
        else:
            compressed.append(stripped)
            prev_blank = False
    # 과도한 공백 반복 제거 (단어 사이 공백 2개 이상 → 1개)
    result = '\n'.join(compressed)
    result = re.sub(r'  +', ' ', result)
    return result.strip()


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """PDF에서 텍스트 추출 (PyMuPDF 사용). Returns (text, page_count)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text")
            if page_text.strip():
                pages_text.append(f"[페이지 {page_num}]\n{page_text}")
        page_count = len(doc)
        doc.close()
        full_text = "\n\n".join(pages_text)
        return compress_text(preprocess_text(full_text)), page_count
    except ImportError:
        st.error("PyMuPDF 패키지가 설치되지 않았습니다. `pip install PyMuPDF`를 실행하세요.")
        return "", 0
    except Exception as e:
        st.error(f"PDF 텍스트 추출 오류: {e}")
        return "", 0


def extract_text_from_docx(file_bytes: bytes) -> str:
    """DOCX에서 텍스트 추출 (python-docx 사용)"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # 표(table) 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        full_text = "\n".join(paragraphs)
        return compress_text(preprocess_text(full_text))
    except ImportError:
        st.error("python-docx 패키지가 설치되지 않았습니다. `pip install python-docx`를 실행하세요.")
        return ""
    except Exception as e:
        st.error(f"DOCX 텍스트 추출 오류: {e}")
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    """TXT 파일에서 텍스트 추출"""
    try:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("cp949", errors="replace")
        return compress_text(preprocess_text(text))
    except Exception as e:
        st.error(f"TXT 텍스트 추출 오류: {e}")
        return ""


def extract_text(uploaded_file) -> tuple[str, int]:
    """
    업로드된 파일에서 텍스트 추출.
    Returns: (text, page_count_or_0)
    """
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        text, pages = extract_text_from_pdf(file_bytes)
        return text, pages
    elif name.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
        return text, 0
    elif name.endswith(".txt"):
        text = extract_text_from_txt(file_bytes)
        return text, 0
    else:
        st.error(f"지원하지 않는 파일 형식입니다: {name}")
        return "", 0


def truncate_text(text: str, max_chars: int) -> str:
    """텍스트가 너무 길면 앞뒤 균형 있게 자르기"""
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return (
        text[:half]
        + f"\n\n... [중략: {len(text) - max_chars:,}자 생략] ...\n\n"
        + text[-half:]
    )


def slice_text_for_step(text: str, step: int, step_ratio: dict) -> str:
    """
    단계별로 필요한 텍스트 구간만 추출.
    step_ratio = {step: (start_ratio, end_ratio)}
    예) {1: (0.0, 0.4)} → 앞 40%만 전달 → 1단계 토큰 60% 절약
    """
    ratio = step_ratio.get(step, (0.0, 1.0))
    start_ratio, end_ratio = ratio
    total = len(text)
    start_idx = int(total * start_ratio)
    end_idx   = int(total * end_ratio)
    sliced = text[start_idx:end_idx]
    if start_ratio > 0.0:
        sliced = f"... [앞부분 생략] ...\n\n{sliced}"
    if end_ratio < 1.0:
        sliced = f"{sliced}\n\n... [이후 생략] ..."
    return sliced
